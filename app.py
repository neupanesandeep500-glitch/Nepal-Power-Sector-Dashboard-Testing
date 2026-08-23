"""
Nepal Power Plant & Transmission Line License Status Dashboard — WEB EDITION
Author (original desktop app): Er. Sandeep Neupane
Web port: Dash / Plotly, deployable on Render.

FIXED VERSION:
- GIS data bundled inline (no external uploads needed)
- gis-opt-layers moved to main layout (fixes callback error)
- Overview tab shows data immediately (no 4s blank wait)
- f-year defaults to None (no premature filtering)
- District/local cascade refreshes on data load
- All error handling improved with visible messages

ENHANCED VERSION (2026):
- Common background images for Power Plant types and Province cards
- Different colored fonts with effects for status/province labels
- Consistent "... Projects, ... Capacity" / "... KM" pattern
- Ordered stages: Operating, Construction License, ..., Application for Survey
- Ordered provinces: Koshi, Madhesh, Bagmati, Gandaki, Lumbini, Karnali, Sudurpaschim
- KPI summary only on Overview tab
- "Installed Capacity" with operating plants summary at first
- Animated province slides in Power Plants > By Province
- No flipping in Transmission tab when filtered
- Growth tab separated for Transmission vs Power Plants
- Data Table with page size dropdown (10, 25, 50, 100, All)
- Custom tab for chart styling options
- Watermark "Er. Sandeep Neupane" on downloaded charts
"""

import os
import io
import math
import base64
import tempfile
import traceback
import textwrap
import logging
from collections import defaultdict
from datetime import datetime

import pandas as pd
import dash
from dash import dcc, html, Input, Output, State, dash_table, ctx
import dash_bootstrap_components as dbc
import plotly.graph_objects as go
import plotly.express as px
import plotly.io as pio

import data_engine as de
import server_state as ss
import coordinate_transform as ct
import gis_leaflet_map
from admin import admin_bp
import NEA
import transmission_network as tn_net
import visitor_counter
import security

app_logger = logging.getLogger("nepal_power_dashboard")

import matplotlib
matplotlib.use("Agg")  # headless backend — must be set before pyplot is imported anywhere
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages

# ── GLOBAL CHART STYLE STATE (for Custom Tab) ───────────────────────────────

# ── GLOBAL CHART STYLE STATE (for Custom Tab) ───────────────────────────────
CHART_STYLE_STATE = {
    "bar_mode": "group",        # group, stack, relative
    "chart_type": "bar",        # bar, line, area, scatter
    "color_scheme": "default",  # default, pastel, dark, vibrant
    "show_grid": True,
    "font_family": "Arial",
    "title_size": 16,
    "label_size": 12,
    "animation": True,
    "secondary_axis": True,     # show the cumulative line on a secondary y-axis
}


def _secondary_axis_enabled():
    """Whether charts should draw their cumulative-total line on a
    secondary y-axis, per the Custom tab's toggle."""
    return CHART_STYLE_STATE.get("secondary_axis", True)


def _apply_secondary_axis_setting(fig):
    """Call on every figure that has a cumulative line on a secondary
    y-axis, right before returning it. When the Custom tab's "Show
    Secondary (Cumulative) Axis" switch is off, this drops the
    trace(s) plotted on yaxis2 and clears the yaxis2 layout config, so
    the cumulative line and its right-hand axis disappear together
    instead of leaving an axis with nothing plotted on it. When the
    switch is on (the default), this is a no-op."""
    if _secondary_axis_enabled():
        return fig
    fig.data = [tr for tr in fig.data if getattr(tr, "yaxis", None) != "y2"]
    fig.update_layout(yaxis2=None)
    return fig

# ── COLOR SCHEMES ───────────────────────────────────────────────────────────
COLOR_SCHEMES = {
    "default": {
        "Operating": "#2e7d32", "Construction License": "#fb8c00",
        "Application for Construction License": "#ffb300",
        "Survey License": "#42a5f5",
        "Application for Survey License": "#90a4ae",
        "GoN Study Project": "#0277bd", "Cancelled": "#c62828",
        "Technical Clearance": "#9fb3c8",
    },
    "pastel": {
        "Operating": "#81c784", "Construction License": "#ffb74d",
        "Application for Construction License": "#fff176",
        "Survey License": "#64b5f6",
        "Application for Survey License": "#b0bec5",
        "GoN Study Project": "#4fc3f7", "Cancelled": "#e57373",
        "Technical Clearance": "#b0bec5",
    },
    "dark": {
        "Operating": "#1b5e20", "Construction License": "#e65100",
        "Application for Construction License": "#ff6f00",
        "Survey License": "#1565c0",
        "Application for Survey License": "#455a64",
        "GoN Study Project": "#01579b", "Cancelled": "#b71c1c",
        "Technical Clearance": "#455a64",
    },
    "vibrant": {
        "Operating": "#00e676", "Construction License": "#ff9100",
        "Application for Construction License": "#ffea00",
        "Survey License": "#2979ff",
        "Application for Survey License": "#78909c",
        "GoN Study Project": "#00b0ff", "Cancelled": "#ff1744",
        "Technical Clearance": "#78909c",
    },
}

GENERIC_NEA_PALETTE = {
    "default": ["#1565c0", "#c62828", "#2e7d32", "#f9a825", "#6a1b9a", "#00897b", "#ff8f00", "#8e24aa"],
    "pastel":  ["#64b5f6", "#e57373", "#81c784", "#fff176", "#ba68c8", "#4db6ac", "#ffcc80", "#ce93d8"],
    "dark":    ["#1565c0", "#b71c1c", "#1b5e20", "#e65100", "#4a148c", "#004d40", "#bf360c", "#311b92"],
    "vibrant": ["#2979ff", "#ff1744", "#00e676", "#ffea00", "#d500f9", "#1de9b6", "#ff9100", "#651fff"],
}


def _nea_style_payload():
    """Current Custom Style settings, translated into what the NEA
    pages' Chart.js code understands — those pages are static HTML
    (not Plotly/Dash components), so style can't flow through props;
    it's passed as a small JSON blob baked into the page instead."""
    s = CHART_STYLE_STATE
    return {
        "font_family": s.get("font_family", "Arial"),
        "title_size": s.get("title_size", 16),
        "label_size": s.get("label_size", 12),
        "show_grid": bool(s.get("show_grid", True)),
        "animation": bool(s.get("animation", True)),
        "palette": GENERIC_NEA_PALETTE.get(s.get("color_scheme", "default"), GENERIC_NEA_PALETTE["default"]),
    }


PROVINCE_COLOR_SCHEMES = {
    "default": {
        "Koshi": "#00695c", "Madhesh": "#ef6c00", "Bagmati": "#1565c0",
        "Gandaki": "#6a1b9a", "Lumbini": "#2e7d32", "Karnali": "#c62828",
        "Sudurpaschim": "#4527a0", "Unspecified": "#78909c",
    },
    "pastel": {
        "Koshi": "#4db6ac", "Madhesh": "#ff9800", "Bagmati": "#5c6bc0",
        "Gandaki": "#ab47bc", "Lumbini": "#66bb6a", "Karnali": "#ef5350",
        "Sudurpaschim": "#7e57c2", "Unspecified": "#b0bec5",
    },
    "dark": {
        "Koshi": "#004d40", "Madhesh": "#bf360c", "Bagmati": "#0d47a1",
        "Gandaki": "#4a148c", "Lumbini": "#1b5e20", "Karnali": "#b71c1c",
        "Sudurpaschim": "#311b92", "Unspecified": "#37474f",
    },
    "vibrant": {
        "Koshi": "#00bfa5", "Madhesh": "#ff6d00", "Bagmati": "#2962ff",
        "Gandaki": "#aa00ff", "Lumbini": "#00c853", "Karnali": "#ff1744",
        "Sudurpaschim": "#651fff", "Unspecified": "#546e7a",
    },
}

TYPE_COLOR_SCHEMES = {
    "default": {
        "Hydro (>1MW)": "#1565c0", "Hydro (<=1MW)": "#42a5f5", "Solar": "#f9a825",
        "Wind": "#26a69a", "Co-generation": "#8d6e63", "Thermal": "#6d4c41",
        "Biomass": "#558b2f", "Transmission Line": "#6a1b9a", "Other": "#78909c",
    },
    "pastel": {
        "Hydro (>1MW)": "#5c6bc0", "Hydro (<=1MW)": "#90caf9", "Solar": "#fff59d",
        "Wind": "#80cbc4", "Co-generation": "#bcaaa4", "Thermal": "#a1887f",
        "Biomass": "#a5d6a7", "Transmission Line": "#ce93d8", "Other": "#b0bec5",
    },
    "dark": {
        "Hydro (>1MW)": "#0d47a1", "Hydro (<=1MW)": "#1565c0", "Solar": "#f57f17",
        "Wind": "#00695c", "Co-generation": "#4e342e", "Thermal": "#3e2723",
        "Biomass": "#33691e", "Transmission Line": "#4a148c", "Other": "#263238",
    },
    "vibrant": {
        "Hydro (>1MW)": "#2962ff", "Hydro (<=1MW)": "#00b0ff", "Solar": "#ffea00",
        "Wind": "#00bfa5", "Co-generation": "#8d6e63", "Thermal": "#5d4037",
        "Biomass": "#76ff03", "Transmission Line": "#d500f9", "Other": "#607d8b",
    },
}


def get_status_colors():
    return COLOR_SCHEMES.get(CHART_STYLE_STATE["color_scheme"], COLOR_SCHEMES["default"])


def get_province_colors():
    return PROVINCE_COLOR_SCHEMES.get(CHART_STYLE_STATE["color_scheme"], PROVINCE_COLOR_SCHEMES["default"])


def get_type_colors():
    return TYPE_COLOR_SCHEMES.get(CHART_STYLE_STATE["color_scheme"], TYPE_COLOR_SCHEMES["default"])


# ── APP SETUP ────────────────────────────────────────────────────────────────
app = dash.Dash(
    __name__,
    external_stylesheets=[dbc.themes.FLATLY, dbc.icons.BOOTSTRAP],
    title="Nepal Power Plants Dashboard",
    suppress_callback_exceptions=True,
    meta_tags=[
        {"name": "viewport", "content": "width=device-width, initial-scale=1"},
        {
            "name": "description",
            "content": (
                "Nepal Power Plants Dashboard — Live licensing, generation, and "
                "transmission tracker for Nepal Source:www.doed.gov.np & www.nea.org.np."
            ),
        },
        {"name": "theme-color", "content": "#0d47a1"},
        {"name": "color-scheme", "content": "light dark"},
        {"name": "apple-mobile-web-app-capable", "content": "yes"},
        {"name": "apple-mobile-web-app-status-bar-style", "content": "black-translucent"},
    ],
)
server = app.server

TICKER_CSS = """
.ticker-bar { display: flex; align-items: center; overflow: hidden; background: #101726;
  padding: 8px 12px; border-radius: 6px; margin-bottom: 14px; white-space: nowrap; }
.ticker-live-badge { display: flex; align-items: center; flex: 0 0 auto; gap: 6px;
  margin-right: 14px; padding: 3px 10px; border-radius: 4px;
  background: rgba(211,47,47,0.16); border: 1px solid rgba(244,67,54,0.55); }
.ticker-live-dot { width: 9px; height: 9px; border-radius: 50%; background: #ff1744;
  box-shadow: 0 0 6px #ff1744; animation: ticker-live-blink 1.1s ease-in-out infinite; }
.ticker-live-text { color: #ff5252; font-weight: 800; font-size: 12px; letter-spacing: 1px;
  font-family: -apple-system, "Segoe UI", Helvetica, Arial, sans-serif; }
@keyframes ticker-live-blink { 0%, 100% { opacity: 1; } 50% { opacity: 0.15; } }
.ticker-track-wrap { flex: 1 1 auto; overflow: hidden; white-space: nowrap; }
.ticker-track { display: inline-block; white-space: nowrap; padding-left: 100%;
  animation-name: ticker-scroll; animation-timing-function: linear;
  animation-iteration-count: infinite;
  font-family: -apple-system, "Segoe UI", Helvetica, Arial, sans-serif;
  font-weight: 600; font-size: 14px; }
.ticker-bar:hover .ticker-track { animation-play-state: paused; }
@keyframes ticker-scroll { 0% { transform: translateX(0); } 100% { transform: translateX(-50%); } }
.main-tabs-nav .nav-link { font-weight: 600; font-size: 14px; color: #37474f; border: none;
  border-radius: 8px 8px 0 0; padding: 10px 16px; margin-right: 4px;
  transition: transform 0.12s ease, background 0.15s ease; }
.main-tabs-nav .nav-link:hover { background: #eef3fb; transform: translateY(-1px); }
.main-tabs-nav .nav-link.active { color: #fff !important;
  background: linear-gradient(135deg, #1565c0 0%, #0d47a1 100%) !important;
  box-shadow: 0 2px 8px rgba(13,71,161,0.35); }
.live-clock-wrap { background: #0b1730; border: 1px solid #3d5a99; border-radius: 6px;
  padding: 4px 12px; text-align: right; line-height: 1.25; }
.live-clock-date { color: #8fb2ff; font-size: 11px; font-weight: 700; letter-spacing: 0.02em; }
.live-clock-time { color: #ffd166; font-size: 15px; font-weight: 700;
  font-family: Consolas, "Courier New", monospace; }
.site-header { position: relative; overflow: hidden; }
.site-header-title { font-weight: 800; font-size: 20px; line-height: 1.2; }
.site-header-subtitle { font-size: 12px; opacity: 0.85; }
.site-header-flag, .site-header-logo { border-radius: 3px;
  box-shadow: 0 0 0 1px rgba(255,255,255,0.25); }
footer.site-footer { background: #0b1730; color: #b7c4e0; margin-top: 28px;
  padding: 18px 24px; font-size: 13px; }
footer.site-footer a { color: #8fb2ff; text-decoration: none; margin-right: 16px; }
footer.site-footer a:hover { text-decoration: underline; }
.footer-visitor-counter { color: #ffd166; font-size: 15px; font-weight: 700; letter-spacing: 0.02em; }
.footer-last-update { color: #9fd8ff; font-size: 14px; font-weight: 600; margin-top: 4px; }
.footer-disclaimer { border-top: 1px solid rgba(183,196,224,0.25); margin-top: 14px;
  padding-top: 12px; font-size: 12px; line-height: 1.6; color: #9aa8c7; text-align: justify; }
.footer-disclaimer strong { color: #c7d3ec; }
.footer-disclaimer a { color: #8fb2ff; }

/* ── Status label styles with effects ─────────────────────────────────────── */
.status-label-operating { color: #2e7d32; font-weight: 800; text-shadow: 0 0 8px rgba(46,125,50,0.4); }
.status-label-construction { color: #fb8c00; font-weight: 800; text-shadow: 0 0 8px rgba(251,140,0,0.4); }
.status-label-app-construction { color: #ffb300; font-weight: 800; text-shadow: 0 0 8px rgba(255,179,0,0.4); }
.status-label-survey { color: #42a5f5; font-weight: 800; text-shadow: 0 0 8px rgba(66,165,245,0.4); }
.status-label-app-survey { color: #90a4ae; font-weight: 800; text-shadow: 0 0 8px rgba(144,164,174,0.4); }
.status-label-gon { color: #0277bd; font-weight: 800; text-shadow: 0 0 8px rgba(2,119,189,0.4); }
.status-label-cancelled { color: #c62828; font-weight: 800; text-shadow: 0 0 8px rgba(198,40,40,0.4); }
.status-label-tc { color: #9fb3c8; font-weight: 800; text-shadow: 0 0 8px rgba(159,179,200,0.4); }

/* ── Province label styles with effects ──────────────────────────────────── */
.prov-label-koshi { color: #00695c; font-weight: 800; text-shadow: 0 0 8px rgba(0,105,92,0.4); }
.prov-label-madhesh { color: #ef6c00; font-weight: 800; text-shadow: 0 0 8px rgba(239,108,0,0.4); }
.prov-label-bagmati { color: #1565c0; font-weight: 800; text-shadow: 0 0 8px rgba(21,101,192,0.4); }
.prov-label-gandaki { color: #6a1b9a; font-weight: 800; text-shadow: 0 0 8px rgba(106,27,154,0.4); }
.prov-label-lumbini { color: #2e7d32; font-weight: 800; text-shadow: 0 0 8px rgba(46,125,50,0.4); }
.prov-label-karnali { color: #c62828; font-weight: 800; text-shadow: 0 0 8px rgba(198,40,40,0.4); }
.prov-label-sudurpaschim { color: #4527a0; font-weight: 800; text-shadow: 0 0 8px rgba(69,39,160,0.4); }

/* ── Card header gradient overlays ───────────────────────────────────────── */
.card-header-gradient { position: relative; }
.card-header-gradient::after {
  content: "";
  position: absolute; top: 0; left: 0; right: 0; bottom: 0;
  background: linear-gradient(180deg, rgba(0,0,0,0.1) 0%, rgba(0,0,0,0.5) 100%);
  pointer-events: none;
}

/* ── Flip card animation ───────────────────────────────────────────────── */
@keyframes fadeInUp {
  from { opacity: 0; transform: translateY(20px); }
  to { opacity: 1; transform: translateY(0); }
}
.flip-card-animate { animation: fadeInUp 0.6s ease-out; }

/* Every 6s the Overview/category "flip" cards swap their dcc.Graph figure
   via a callback. dash-core-components marks the graph wrapper with
   .dash-graph--pending while that swap is in flight and its own default
   stylesheet dims it (opacity ~0.2-0.7) as a generic "this is loading"
   cue — on a chart that already redraws every 6 seconds, that dimming
   reads as a page/section reload rather than a deliberate flip. The
   fadeInUp animation above already gives the flip its own motion, so the
   pending-dim on top of it is redundant; switch it off here. */
.dash-graph--pending .js-plotly-plot,
.dash-graph--pending .plot-container,
.dash-graph--pending { opacity: 1 !important; transition: none !important; }

/* ── Custom tab styling ──────────────────────────────────────────────────── */
.custom-style-panel { background: #f8f9fa; border-radius: 8px; padding: 20px; }
.custom-style-panel h5 { color: #1565c0; margin-bottom: 16px; }

/* ── REQ 2: all KPI cards in one row, shrinking evenly to fit ─────────────── */
.kpi-flex-row { display: flex; flex-wrap: nowrap; gap: 10px; align-items: stretch;
  overflow-x: auto; padding-bottom: 2px; }
.kpi-flex-item { flex: 1 1 0; min-width: 118px; }
.kpi-flex-item .card { height: 100%; }
.kpi-flex-item .card-body { padding: 8px 10px !important; }
.kpi-flex-item .card-body > div:first-child { font-size: 9px !important; white-space: nowrap; }
.kpi-flex-item .card-body > div:nth-child(2) { font-size: clamp(12px, 1.6vw, 16px) !important;
  white-space: nowrap; overflow: hidden; text-overflow: ellipsis; }
.kpi-flex-item .card-body > div:nth-child(3) { font-size: 9.5px !important;
  white-space: nowrap; overflow: hidden; text-overflow: ellipsis; }
@media (max-width: 991px) { .kpi-flex-row { flex-wrap: wrap; overflow-x: visible; }
  .kpi-flex-item { flex: 1 1 45%; min-width: 130px; } }
@media (max-width: 575px) { .kpi-flex-item { flex: 1 1 100%; } }

/* ── REQ 3/6: every sub-tab bar in the system gets one consistent,       ── */
/* ── clearly-distinct active state (teal/green so it reads as a level    ── */
/* ── below the blue main-tabs-nav, never confused with it or the         ── */
/* ── inactive tabs around it).                                           ── */
.sub-tabs-nav .nav-link, .nea-subtabs-nav .nea-subtab-btn {
  font-weight: 600; font-size: 13px; color: #37474f; border: 1px solid transparent;
  border-radius: 6px; padding: 7px 14px; margin-right: 4px; background: #eef1f6;
  transition: transform 0.12s ease, background 0.15s ease, color 0.15s ease;
}
.sub-tabs-nav .nav-link:hover, .nea-subtabs-nav .nea-subtab-btn:hover {
  background: #dbe6f7; transform: translateY(-1px); color: #0d47a1;
}
.sub-tabs-nav .nav-link.active, .nea-subtabs-nav .nea-subtab-btn.active {
  color: #fff !important; background: linear-gradient(135deg, #00897b 0%, #00695c 100%) !important;
  border-color: transparent !important; box-shadow: 0 2px 7px rgba(0,105,92,0.4);
}
.nea-subtabs-nav { display: flex; flex-wrap: wrap; gap: 4px; margin-bottom: 12px; }
.nea-subtabs-nav .nea-subtab-btn { cursor: pointer; border: 1px solid transparent; }

/* ── Mobile-friendly main tab bar: horizontal scroll instead of wrap/  ── */
/* ── overflow-clip on narrow screens, with a subtle fade+scrollbar     ── */
/* ── hint so it reads as scrollable rather than cut off.               ── */
@media (max-width: 767px) {
  .main-tabs-nav { flex-wrap: nowrap !important; overflow-x: auto; overflow-y: hidden;
    -webkit-overflow-scrolling: touch; scrollbar-width: thin; padding-bottom: 4px; }
  .main-tabs-nav .nav-link { flex: 0 0 auto; font-size: 12.5px; padding: 8px 12px; }
  .kpi-flex-row { gap: 6px; }
}
.main-tabs-nav::-webkit-scrollbar { height: 5px; }
.main-tabs-nav::-webkit-scrollbar-thumb { background: rgba(21,101,192,0.35); border-radius: 4px; }

/* ── Theme toggle button ──────────────────────────────────────────────── */
.theme-toggle-btn { border: 1px solid rgba(255,255,255,0.35); background: rgba(255,255,255,0.08);
  color: inherit; border-radius: 20px; padding: 5px 12px; font-size: 12px; font-weight: 700;
  cursor: pointer; display: inline-flex; align-items: center; gap: 6px;
  transition: background 0.15s ease, transform 0.12s ease; }
.theme-toggle-btn:hover { background: rgba(255,255,255,0.18); transform: translateY(-1px); }

/* ── Back-to-top button ───────────────────────────────────────────────── */
#back-to-top-btn { position: fixed; right: 20px; bottom: 22px; z-index: 1050;
  width: 42px; height: 42px; border-radius: 50%; border: none; cursor: pointer;
  background: linear-gradient(135deg, #1565c0 0%, #0d47a1 100%); color: #fff;
  font-size: 18px; box-shadow: 0 3px 10px rgba(13,71,161,0.45);
  display: flex; align-items: center; justify-content: center;
  opacity: 0; pointer-events: none; transform: translateY(12px);
  transition: opacity 0.2s ease, transform 0.2s ease; }
#back-to-top-btn.visible { opacity: 1; pointer-events: auto; transform: translateY(0); }
#back-to-top-btn:hover { filter: brightness(1.1); }

/* ── Custom dash-loading spinner, matched to the site palette ───────────── */
._dash-loading-callback, .dash-spinner { border-color: #1565c0 !important; }

/* ═══════════════════════════════════════════════════════════════════════
   DARK MODE
   Applied via [data-theme="dark"] on <html>, toggled client-side and
   persisted in localStorage (see THEME_JS below). This re-skins the
   app chrome (page background, cards, sidebar, tables, tabs, footer,
   admin-adjacent surfaces) using CSS variables layered on top of the
   FLATLY Bootstrap theme. Plotly figure canvases keep their own
   white "paper" background by design — full per-chart re-theming
   would mean touching every chart-builder function's color literals,
   so charts render as light cards floating on the dark chrome, which
   is a common and legible pattern rather than a partial/broken dark
   theme on the plots themselves.
   ═══════════════════════════════════════════════════════════════════════ */
:root { --app-bg: #eef1f6; --app-card-bg: #ffffff; --app-text: #1c2230;
  --app-border: rgba(0,0,0,0.08); --app-muted: #5c6b82; }
[data-theme="dark"] { --app-bg: #0b1220; --app-card-bg: #141c30; --app-text: #e6ecfa;
  --app-border: rgba(255,255,255,0.10); --app-muted: #93a2c2; }

[data-theme="dark"] body { background: var(--app-bg); color: var(--app-text); }
[data-theme="dark"] .card { background: var(--app-card-bg) !important; color: var(--app-text);
  border-color: var(--app-border) !important; }
[data-theme="dark"] .card-header { background: rgba(255,255,255,0.04) !important;
  border-color: var(--app-border) !important; }
[data-theme="dark"] .table { color: var(--app-text); }
[data-theme="dark"] .form-control, [data-theme="dark"] .form-select {
  background: #0f1729; color: var(--app-text); border-color: var(--app-border); }
[data-theme="dark"] .custom-style-panel { background: #0f1729 !important; }
[data-theme="dark"] .sub-tabs-nav .nav-link, [data-theme="dark"] .nea-subtabs-nav .nea-subtab-btn {
  background: #182238; color: var(--app-muted); }
[data-theme="dark"] .main-tabs-nav .nav-link { background: transparent; color: var(--app-muted); }
[data-theme="dark"] .main-tabs-nav .nav-link:hover { background: #182238; }
[data-theme="dark"] .dash-table-container, [data-theme="dark"] .dash-spreadsheet-container {
  filter: invert(1) hue-rotate(180deg); }
[data-theme="dark"] .dash-table-container img, [data-theme="dark"] .dash-spreadsheet-container img {
  filter: invert(1) hue-rotate(180deg); }
[data-theme="dark"] ::selection { background: #1565c0; color: #fff; }
"""
CLOCK_JS = """
<script>
function _tickLiveClock() {
  var now = new Date();
  var dateOpts = { weekday: 'long', year: 'numeric', month: 'long', day: '2-digit' };
  var dateStr = now.toLocaleDateString('en-US', dateOpts).toUpperCase();
  var timeStr = '🕐 ' + now.toLocaleTimeString('en-US', { hour12: true });
  document.querySelectorAll('.live-clock-date').forEach(function(el) { el.textContent = dateStr; });
  document.querySelectorAll('.live-clock-time').forEach(function(el) { el.textContent = timeStr; });
}
setInterval(_tickLiveClock, 1000);
document.addEventListener('DOMContentLoaded', _tickLiveClock);
_tickLiveClock();
function _loadVisitorCount() {
  var el = document.getElementById('visitor-counter');
  fetch('/api/visitor-count').then(function(r) { return r.json(); }).then(function(d) {
    if (el) { el.textContent = '👥 ' + d.count.toLocaleString() + ' visitors'; }
  }).catch(function() {
    if (el) { el.textContent = '👥 visitors'; }
  });
}
document.addEventListener('DOMContentLoaded', _loadVisitorCount);
setInterval(function() {
  var el = document.getElementById('visitor-counter');
  if (el && el.textContent.indexOf('\u2026') !== -1) { _loadVisitorCount(); }
}, 2000);

// ── Dark mode: apply saved/preferred theme before paint, wire up toggle ──
(function() {
  function preferredTheme() {
    var saved = localStorage.getItem('npd-theme');
    if (saved === 'dark' || saved === 'light') return saved;
    return window.matchMedia && window.matchMedia('(prefers-color-scheme: dark)').matches
      ? 'dark' : 'light';
  }
  function applyTheme(theme) {
    document.documentElement.setAttribute('data-theme', theme);
    document.querySelectorAll('.theme-toggle-icon').forEach(function(el) {
      el.textContent = theme === 'dark' ? '☀️' : '🌙';
    });
    document.querySelectorAll('.theme-toggle-label').forEach(function(el) {
      el.textContent = theme === 'dark' ? 'Light' : 'Dark';
    });
  }
  applyTheme(preferredTheme());
  document.addEventListener('DOMContentLoaded', function() {
    applyTheme(preferredTheme());
    document.addEventListener('click', function(ev) {
      var btn = ev.target.closest('#theme-toggle-btn');
      if (!btn) return;
      var next = document.documentElement.getAttribute('data-theme') === 'dark' ? 'light' : 'dark';
      localStorage.setItem('npd-theme', next);
      applyTheme(next);
    });
  });
})();

// ── Back-to-top button: fades in after scrolling past one screen height ──
(function() {
  function ensureButton() {
    if (document.getElementById('back-to-top-btn')) return;
    var btn = document.createElement('button');
    btn.id = 'back-to-top-btn';
    btn.setAttribute('aria-label', 'Back to top');
    btn.title = 'Back to top';
    btn.textContent = '↑';
    btn.onclick = function() { window.scrollTo({ top: 0, behavior: 'smooth' }); };
    document.body.appendChild(btn);
  }
  document.addEventListener('DOMContentLoaded', function() {
    ensureButton();
    window.addEventListener('scroll', function() {
      var btn = document.getElementById('back-to-top-btn');
      if (!btn) return;
      if (window.scrollY > window.innerHeight * 0.6) btn.classList.add('visible');
      else btn.classList.remove('visible');
    }, { passive: true });
  });
})();
</script>
"""



GA_MEASUREMENT_ID = os.environ.get("GA_MEASUREMENT_ID", "G-DD12W6FLZ8")

GA_SNIPPET = f"""<script async src="https://www.googletagmanager.com/gtag/js?id={GA_MEASUREMENT_ID}"></script>
<script>
  window.dataLayer = window.dataLayer || [];
  function gtag(){{dataLayer.push(arguments);}}
  gtag('js', new Date());
  gtag('config', '{GA_MEASUREMENT_ID}');
</script>"""

# Sets data-theme on <html> synchronously, before first paint, so
# dark-mode users don't see a flash of the light theme while the main
# CLOCK_JS bundle (loaded at the end of <body>) finishes wiring up.
EARLY_THEME_JS = """<script>
(function() {
  try {
    var saved = localStorage.getItem('npd-theme');
    var theme = (saved === 'dark' || saved === 'light') ? saved :
      ((window.matchMedia && window.matchMedia('(prefers-color-scheme: dark)').matches) ? 'dark' : 'light');
    document.documentElement.setAttribute('data-theme', theme);
  } catch (e) {}
})();
</script>"""

app.index_string = f"""<!DOCTYPE html>
<html><head>{{%metas%}}<title>{{%title%}}</title>{{%favicon%}}{{%css%}}
<style>{TICKER_CSS}</style>
{EARLY_THEME_JS}
{GA_SNIPPET}
</head>
<body>{{%app_entry%}}<footer>{{%config%}}{{%scripts%}}{{%renderer%}}</footer>{CLOCK_JS}</body></html>"""

server.secret_key = os.environ.get("FLASK_SECRET_KEY", os.urandom(32).hex())
if not os.environ.get("FLASK_SECRET_KEY"):
    app_logger.warning(
        "FLASK_SECRET_KEY is not set — using a random key generated at process "
        "startup. Every restart (and every worker, if running with >1 gunicorn "
        "worker) will invalidate existing admin sessions. Set FLASK_SECRET_KEY "
        "for stable sessions in production."
    )
server.register_blueprint(admin_bp)

# Security headers, hardened session cookies, gzip compression (if available),
# friendly 404/500 pages, and a /healthz endpoint for uptime monitors.
security.init_app(server, state_getter=lambda: ss.STATE)

# CRITICAL FIX: Bootstrap GIS immediately (bundled data, always works)
ss.bootstrap_on_startup()
ss.start_background_refresh()

# NEA Operational Data / Forecast Lab: independent of the power-plant
# licensing data above — its own Google Sheet, its own background sync.
NEA.bootstrap()
NEA.start_background_refresh()

# Transmission Network map: independent of both datasets above — its own
# workbook (Substations + Transmission Lines), admin-upload or optional
# Google Sheet sync, own background refresh.
tn_net.bootstrap()
tn_net.start_background_refresh()

# Visitor counter: own Google Sheet (via a service-account), own
# background flush — pulls the last-saved count in, then persists new
# visits back out every VISITOR_FLUSH_INTERVAL seconds so it survives
# redeploys. Never blocks/crashes startup if the sheet is unreachable.
visitor_counter.bootstrap()
visitor_counter.start_background_flush()
visitor_counter.record_deploy()

STATE = ss.STATE


@server.route("/nea-vendor/<path:filename>")
def serve_nea_vendor(filename):
    # Chart.js is bundled locally here (nea_assets/vendor/) instead of being
    # pulled from a CDN, because a blocked/slow/ad-blocked cdnjs request was
    # the actual cause of "Chart is not defined" errors on the NEA charts
    # and Forecast Lab — the KPI cards (plain DOM writes) still rendered
    # fine, but every canvas that needed `new Chart(...)` threw, since the
    # library itself never loaded in the browser.
    from flask import send_from_directory
    vendor_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "nea_assets", "vendor")
    return send_from_directory(vendor_dir, filename)


@server.route("/nea-operational-dashboard")
def nea_operational_dashboard():
    try:
        return NEA.render_dashboard_html(style=_nea_style_payload())
    except Exception:
        app_logger.exception("NEA operational dashboard failed to render")
        return security.render_friendly_error(
            "Dashboard temporarily unavailable",
            "The operational performance dashboard couldn't be rendered right now. "
            "The issue has been logged — please try again shortly.")


@server.route("/nea-forecast-lab")
def nea_forecast_lab_page():
    try:
        return NEA.render_forecast_lab_html(style=_nea_style_payload())
    except Exception:
        app_logger.exception("NEA forecast lab failed to render")
        return security.render_friendly_error(
            "Forecast Lab temporarily unavailable",
            "The forecast lab couldn't be rendered right now. "
            "The issue has been logged — please try again shortly.")


@server.route("/scenario-analysis")
def scenario_analysis_page():
    try:
        return NEA.render_scenario_analysis_html(
            style=_nea_style_payload()
        )

    except Exception as exc:
        app_logger.exception(
            "Scenario Analysis page failed to render"
        )

        return (
            f"<h1>Scenario Analysis Error</h1>"
            f"<pre>{type(exc).__name__}: {exc}</pre>"
        ), 500

@server.route("/api/scenario-baseline")
def api_scenario_baseline():
    from flask import jsonify

    try:
        data = NEA.scenario_baseline_data()
        return jsonify(data)

    except Exception as exc:
        app_logger.exception(
            "Scenario baseline API failed"
        )

        return jsonify({
            "ok": False,
            "error": str(exc)
        }), 500


@server.route("/transmission-network-map")
def transmission_network_map():
    try:
        return tn_net.render_map_html()
    except Exception:
        app_logger.exception("Transmission Network map failed to render")
        return security.render_friendly_error(
            "Map temporarily unavailable",
            "The transmission network map couldn't be rendered right now. "
            "The issue has been logged — please try again shortly.")


@server.route("/api/nea-forecast-params")
def api_nea_forecast_params():
    from flask import jsonify
    try:
        return jsonify(NEA.forecast_param_choices())
    except Exception:
        traceback.print_exc()
        return jsonify([]), 200


@server.route("/api/nea-forecast", methods=["POST"])
def api_nea_forecast():
    from flask import jsonify, request
    try:
        body = request.get_json(force=True) or {}
        fr = NEA.run_forecast(
            param_key=body.get("param_key"),
            model=body.get("model", "linear"),
            n_ahead=body.get("n_ahead", 5),
            monthly=bool(body.get("monthly", False)),
        )
        return jsonify(NEA.forecast_result_to_dict(fr))
    except Exception as e:
        traceback.print_exc()
        return jsonify(error=str(e)), 400


@server.route("/api/nea-forecast-composite-params")
def api_nea_forecast_composite_params():
    from flask import jsonify
    try:
        return jsonify(NEA.composite_param_choices())
    except Exception:
        traceback.print_exc()
        return jsonify([]), 200


@server.route("/api/nea-forecast-composite", methods=["POST"])
def api_nea_forecast_composite():
    from flask import jsonify, request
    try:
        body = request.get_json(force=True) or {}
        result = NEA.run_composite_forecast(
            composite_key=body.get("composite_key"),
            model=body.get("model", "linear"),
            n_ahead=body.get("n_ahead", 5),
        )
        return jsonify(result)
    except Exception as e:
        traceback.print_exc()
        return jsonify(error=str(e)), 400


@server.route("/api/visitor-count")
def api_visitor_count():
    from flask import jsonify, session as flask_session
    try:
        if not flask_session.get("counted_visit"):
            flask_session["counted_visit"] = True
            visitor_counter.increment()
        return jsonify(count=visitor_counter.get_count())
    except Exception as e:
        return jsonify(count=0, error=str(e))


@server.route("/assets-logo")
def serve_logo():
    from flask import send_file
    path = ss.get_logo_path()
    if not path:
        return "No logo uploaded yet.", 404
    return send_file(path)


@server.route("/assets-flag")
def serve_flag():
    from flask import send_file
    path = ss.get_flag_path()
    if not path:
        return "No flag image available.", 404
    return send_file(path)


@server.route("/assets-type-bg/<slug>")
def serve_type_bg(slug):
    from flask import send_file
    fn = (ss.STATE.get("type_bg") or {}).get(slug)
    if not fn:
        return "No background uploaded for this type.", 404
    path = os.path.join(ss.ASSETS_DIR, fn)
    if not os.path.exists(path):
        return "No background uploaded for this type.", 404
    return send_file(path)


@server.route("/assets-status-bg/<slug>")
def serve_status_bg(slug):
    from flask import send_file
    fn = (ss.STATE.get("status_bg") or {}).get(slug)
    if not fn:
        return "No background uploaded for this license stage.", 404
    path = os.path.join(ss.ASSETS_DIR, fn)
    if not os.path.exists(path):
        return "No background uploaded for this license stage.", 404
    return send_file(path)


@server.route("/assets-province-bg/<slug>")
def serve_province_bg(slug):
    from flask import send_file
    fn = (ss.STATE.get("province_bg") or {}).get(slug)
    if not fn:
        return "No background uploaded for this province.", 404
    path = os.path.join(ss.ASSETS_DIR, fn)
    if not os.path.exists(path):
        return "No background uploaded for this province.", 404
    return send_file(path)


@server.route("/assets-background")
def serve_background():
    from flask import send_file
    path = ss.get_background_path()
    if not path:
        return "No background photo uploaded yet.", 404
    return send_file(path)


# ── LAYOUT ─────────────────────────────────────────────────────────────────
CAPACITY_BIN_OPTIONS = [
    {"label": "All capacities", "value": "all"},
    {"label": "Below 1 MW", "value": "lt1"},
    {"label": "1 MW - 10 MW", "value": "1-10"},
    {"label": "10 MW - 25 MW", "value": "10-25"},
    {"label": "25 MW - 50 MW", "value": "25-50"},
    {"label": "50 MW - 100 MW", "value": "50-100"},
    {"label": "Above 100 MW", "value": "gt100"},
]
CAPACITY_BIN_RANGES = {
    "all": (None, None), "lt1": (0, 1), "1-10": (1, 10), "10-25": (10, 25),
    "25-50": (25, 50), "50-100": (50, 100), "gt100": (100, None),
}

TX_LENGTH_BIN_OPTIONS = [
    {"label": "All lengths", "value": "all"},
    {"label": "Short - up to 80 km", "value": "short"},
    {"label": "Medium - 80-200 km", "value": "medium"},
    {"label": "Long - above 200 km", "value": "long"},
]
TX_LENGTH_BIN_RANGES = {
    "all": (None, None), "short": (None, 80), "medium": (80, 200), "long": (200, None),
}


def kpi_card(title, value, sub, color):
    return dbc.Card(
        dbc.CardBody([
            html.Div(title, className="text-muted small fw-semibold text-uppercase"),
            html.H3(value, className="mb-0 fw-bold", style={"color": color}),
            html.Div(sub, className="text-muted small"),
        ]),
        className="shadow-sm h-100",
        style={"borderTop": f"4px solid {color}"},
    )


def kpi_card_compact(title, value, sub, color):
    """Same shape as kpi_card but with smaller fonts, so a denser row
    (e.g. the NEA operational KPIs alongside the license KPIs) fits
    without the card growing or wrapping awkwardly."""
    return dbc.Card(
        dbc.CardBody([
            html.Div(title, className="text-muted fw-semibold text-uppercase",
                      style={"fontSize": "10px"}),
            html.Div(value, className="fw-bold", style={"color": color, "fontSize": "17px"}),
            html.Div(sub, className="text-muted", style={"fontSize": "11px"}),
        ], style={"padding": "10px 12px"}),
        className="shadow-sm h-100",
        style={"borderTop": f"3px solid {color}"},
    )


def nea_kpi_cards():
    """NEA operational KPIs for the Overview row, pulled straight from
    NEA.py's own kpi dict (the same figures the NEA Operational Data tab
    shows) — no separate parsing here. Returns [] if NEA data hasn't
    synced yet, so the Overview row just falls back to the license KPIs
    alone rather than showing empty/broken cards."""
    try:
        data = NEA.get_dashboard_data() or {}
    except Exception:
        data = {}
    k = data.get("kpi") or {}
    if not k:
        return []

    def yoy(v, good_if_up=True):
        if v is None:
            return "—"
        arrow = "▲" if v >= 0 else "▼"
        good = (v >= 0) == good_if_up
        return arrow + f" {abs(v):,.2f}% YoY"

    return [
        kpi_card_compact("Peak Demand (MW)", f"{k.get('latest_peak', 0):,.1f}",
                          yoy(k.get("peak_growth")), "#6a1b9a"),
        kpi_card_compact("System Loss (%)", f"{k.get('latest_system_loss', 0):,.2f}%",
                          (f"▼ {k.get('loss_reduction', 0):,.2f} pts YoY" if (k.get("loss_reduction") or 0) >= 0
                           else f"▲ {abs(k.get('loss_reduction', 0)):,.2f} pts YoY"), "#c62828"),
        kpi_card_compact("Total Availability (MU)", f"{k.get('latest_total_avail', 0):,.1f}",
                          yoy(k.get("avail_growth")), "#00838f"),
        kpi_card_compact("Latest Revenue (Rs. Mn)", f"{k.get('latest_revenue', 0):,.1f}",
                          yoy(k.get("revenue_growth")), "#2e7d32"),
        kpi_card_compact("Total Consumers", f"{k.get('total_consumers', 0):,.0f}",
                          yoy(k.get("consumer_increase_pct")), "#1565c0"),
    ]


def sidebar():
    return dbc.Card(
        dbc.CardBody([
            html.H5([html.I(className="bi bi-sliders me-2"), "Filters"], className="mb-2"),
            dbc.Accordion(id="filter-tree", start_collapsed=False, always_open=True, children=[
                dbc.AccordionItem(title="📍 Location - Province / District", children=[
                    html.Label("Province", className="fw-semibold small"),
                    dcc.Dropdown(id="f-province", multi=True, placeholder="All provinces"),
                    html.Div("↳ narrows the District list below", className="text-muted",
                              style={"fontSize": "11px", "marginLeft": "8px"}),
                    html.Label("District", className="fw-semibold small mt-2"),
                    dcc.Dropdown(id="f-district", multi=True, placeholder="All districts"),
                    html.Div("↳ narrows the Local Body list below", className="text-muted",
                              style={"fontSize": "11px", "marginLeft": "8px"}),
                    html.Label("Local Body (Gaunpalika / Nagarpalika)",
                                className="fw-semibold small mt-2"),
                    dcc.Dropdown(id="f-local", multi=True, placeholder="All local bodies"),
                    # Coordinate System picker turned off here — the GIS Map's
                    # own inbuilt sidebar already has this toggle, so showing
                    # it twice was redundant. Kept mounted (default WGS-84)
                    # so Data Table / GIS callbacks that read f-crs still work.
                    html.Div([
                        html.Label("Coordinate System (GIS Map / Data Table)",
                                    className="fw-semibold small mt-2"),
                        dcc.RadioItems(
                            id="f-crs",
                            options=[{"label": f" {v}", "value": k} for k, v in ct.CRS_LABELS.items()],
                            value=ct.CRS_WGS84, labelStyle={"display": "block", "fontSize": "13px"},
                        ),
                        html.Div(
                            "DoED Lat/Long sheet values are on the Everest 1830 survey datum; "
                            "the GIS boundary layer is WGS-84. Pick WGS-84 to match the map "
                            "(default) or Everest 1830 to match the raw licence sheet.",
                            className="text-muted", style={"fontSize": "11px"},
                        ),
                    ], style={"display": "none"}),
                ], item_id="grp-location"),

                dbc.AccordionItem(title="⚡ Project - Type / Stage / Capacity", children=[
                    html.Label("Project Type", className="fw-semibold small"),
                    dcc.Dropdown(id="f-type", multi=True, placeholder="All types"),
                    html.Div("↳ each type breaks down by stage below", className="text-muted",
                              style={"fontSize": "11px", "marginLeft": "8px"}),
                    html.Label("License Stage", className="fw-semibold small mt-2"),
                    dcc.Dropdown(id="f-status", multi=True, placeholder="All stages"),
                    html.Label("Capacity Range (MW)", className="fw-semibold small mt-2"),
                    dcc.Dropdown(
                        id="f-capacity", options=CAPACITY_BIN_OPTIONS,
                        value="all", clearable=False, placeholder="All capacities",
                    ),
                    html.Label("Transmission Line Length", className="fw-semibold small mt-2"),
                    dcc.Dropdown(
                        id="f-tx-length", options=TX_LENGTH_BIN_OPTIONS,
                        value="all", clearable=False, placeholder="All lengths",
                    ),
                    html.Div("↳ only narrows Transmission Line records - every "
                             "other type is unaffected", className="text-muted",
                              style={"fontSize": "11px", "marginLeft": "8px"}),
                ], item_id="grp-project"),

                dbc.AccordionItem(title="📅 Dates - License Issue / COD", children=[
                    dcc.Store(id="f-year", data=None),
                    html.Label("License Date - exact range (B.S.)", className="fw-semibold small"),
                    html.Div("Any of YYYY, YYYY-MM, or YYYY-MM-DD. Leave both blank "
                             "for all dates.", className="text-muted",
                             style={"fontSize": "11px", "marginLeft": "8px"}),
                    dbc.Row([
                        dbc.Col(dcc.Input(id="f-date-from", type="text",
                                           placeholder="From e.g. 2078-01-01",
                                           className="form-control form-control-sm"), width=6),
                        dbc.Col(dcc.Input(id="f-date-to", type="text",
                                           placeholder="To e.g. 2082-12-30",
                                           className="form-control form-control-sm"), width=6),
                    ], className="g-1"),
                    html.Label("COD Date range (B.S.) - Operating plants",
                                className="fw-semibold small mt-2"),
                    dbc.Row([
                        dbc.Col(dcc.Input(id="f-cod-from", type="text", placeholder="From YYYY-MM-DD",
                                           className="form-control form-control-sm"), width=6),
                        dbc.Col(dcc.Input(id="f-cod-to", type="text", placeholder="To YYYY-MM-DD",
                                           className="form-control form-control-sm"), width=6),
                    ], className="g-1"),
                ], item_id="grp-dates"),

                dbc.AccordionItem(title="🔎 Search", children=[
                    dcc.Input(id="f-search", type="text", placeholder="Type to search…",
                              className="form-control"),
                ], item_id="grp-search"),
            ]),
            html.Hr(),
            html.Div([
                html.I(className="bi bi-file-earmark-pdf me-1"),
                "Need a PDF report? Head over to the ",
                html.Strong("📄 Generate Report"),
                " tab — it picks up whatever filters you've set here.",
            ], className="small text-muted"),
        ]),
        className="shadow-sm",
    )


TAB_DEFAULT_FILTER_GROUP = {
    "overview": "grp-project", "plants": "grp-project", "transmission": "grp-location",
    "gon_study": "grp-project", "cancelled": "grp-project", "growth": "grp-dates",
    "gis": "grp-location", "compare": "grp-project", "table": "grp-search",
}


@app.callback(Output("filter-tree", "active_item"), Input("main-tabs", "active_tab"))
def open_relevant_filter_group(tab):
    default = TAB_DEFAULT_FILTER_GROUP.get(tab, "grp-project")
    others = [g for g in ("grp-location", "grp-project", "grp-dates", "grp-search")
              if g != default]
    return [default] + others[:1]


def settings_panel():
    return html.Div(id="load-status", style={"display": "none"})


# CRITICAL FIX: gis-opt-layers moved to MAIN LAYOUT
app.layout = dbc.Container(fluid=True, children=[
    dcc.Store(id="filtered-data-signal"),
    dcc.Store(id="chart-style-store", data=CHART_STYLE_STATE),

    html.Div(id="site-header", className="site-header p-3 mb-3", children=[
        dbc.Row(align="center", justify="between", className="g-2", children=[
            dbc.Col(width="auto", children=html.Div([
                html.Img(src="/assets-flag", height="40px",
                         alt="Flag of Nepal", title="Nepal",
                         className="me-2 site-header-flag"),
                html.Img(src="/assets-logo", height="34px",
                         alt="Organisation logo", className="me-2 site-header-logo")
                if ss.get_logo_path() else None,
                html.Div([
                    html.Div("Nepal Power Plants Dashboard",
                              className="site-header-title"),
                    html.Div("Source: www.doed.gov.np & www.nea.org.np | Live Licensing, Operational Performance tracker and Estimator for Nepal's Power Sector",
                              className="site-header-subtitle"),
                ]),
            ], className="d-flex align-items-center")),
            dbc.Col(width="auto", children=html.Div([
                html.Button([
                    html.Span("🌙", className="theme-toggle-icon"),
                    html.Span("Dark", className="theme-toggle-label"),
                ], id="theme-toggle-btn", className="theme-toggle-btn me-2",
                    title="Toggle light / dark theme", type="button"),
                html.Div([
                    html.Div(className="live-clock-date"),
                    html.Div(className="live-clock-time"),
                ], className="live-clock-wrap"),
            ], className="d-flex align-items-center")),
        ]),
    ]),

    html.Div(id="ticker-bar"),

    dbc.Tabs(id="main-tabs", active_tab="overview", className="main-tabs-nav", children=[
        dbc.Tab(label="📊 Overview", tab_id="overview"),
        dbc.Tab(label="🪪 License Status", tab_id="license_status"),
        dbc.Tab(label="📈 License Insights", tab_id="license_insights"),
        dbc.Tab(label="🗺️ GIS Map", tab_id="gis"),
        dbc.Tab(label="🔌 Transmission Network", tab_id="transmission_network"),
        dbc.Tab(label="🏭 System Operational Performance", tab_id="nea_operational"),
        dbc.Tab(label="🔬 Forecast Lab", tab_id="nea_forecast"),
        dbc.Tab(label="🧭 Scenario Analysis", tab_id="scenario_analysis"),
        dbc.Tab(label="🎨 Custom Style", tab_id="custom"),
        dbc.Tab(label="📝 Generate Report", tab_id="reports"),
    ]),

    dbc.Row(className="mt-3", children=[
        dbc.Col(id="filter-sidebar-col", md=3,
                children=[sidebar(), html.Div(className="mt-3"), settings_panel()]),
        dbc.Col(id="main-content-col", md=9, children=[
            html.Div(id="kpi-row", className="mb-3"),
            html.Div(id="tab-content"),
        ]),
    ]),

    # FIX: GIS layer toggles are ALWAYS in DOM
    html.Div(id="gis-controls-container", style={"display": "none"}, children=[
        dbc.Checklist(
            id="gis-opt-layers",
            options=[
                {"label": " License Boundary Polygons", "value": "boundary"},
                {"label": " Protected Areas Overlay", "value": "pa"},
            ],
            value=["boundary"],
            inline=True, switch=True, className="mb-2",
        ),
    ]),

    html.Div(id="_init_trigger", style={"display": "none"}),
    dcc.Interval(id="init-once", n_intervals=0, max_intervals=1, interval=500),
    dcc.Interval(id="refresh-poll", n_intervals=0, interval=36000_000),
    dcc.Interval(id="type-flip-interval", n_intervals=0, interval=6_000),
    dcc.Interval(id="province-flip-interval", n_intervals=0, interval=6_000),
    dcc.Interval(id="nea-perf-flip-interval", n_intervals=0, interval=6_000),
   
   html.Footer(className="site-footer", children=[
        dbc.Row([
            dbc.Col(md=8, children=[
                html.Div("Useful links — Nepal Energy Sector", className="fw-semibold mb-1"),
                html.A("Ministry of Energy, Water Resources and Irrigation (MoEWRI)",
                       href="https://moewri.gov.np", target="_blank", className="d-block"),
                html.A("Electricity Regulatory Commission (ERC)",
                       href="https://erc.gov.np", target="_blank", className="d-block"),
                html.A("Water and Energy Commission Secretariat (WECS)",
                       href="https://wecs.gov.np", target="_blank", className="d-block"),
                html.A("Department of Electricity Development (DoED)",
                       href="https://doed.gov.np", target="_blank", className="d-block"),
                html.A("Nepal Electricity Authority (NEA)",
                       href="https://nea.org.np", target="_blank", className="d-block"),
                html.A("Alternative Energy Promotion Center (AEPC)",
                       href="https://aepc.gov.np", target="_blank", className="d-block"),
            ]),
            dbc.Col(md=4, className="text-md-end", children=[
                html.Div("👥 …visitors", id="visitor-counter", className="footer-visitor-counter"),
                html.Div(id="footer-last-update", className="footer-last-update"),
            ]),
        ]),

        # ── Disclaimer ───────────────────────────────────────────────
        html.Div(className="footer-disclaimer", children=[
            html.P([
                html.Strong("Disclaimer: "),
                "The information presented on this website — including power plant status, "
                "capacity, and licensing details — is compiled from data published by the "
                "Department of Electricity Development (DoED), Government of Nepal (",
                html.A("www.doed.gov.np", href="https://www.doed.gov.np", target="_blank"),
                "). This is an independent, unofficial platform intended only to give a "
                "general overview of Nepal's electricity sector; it is not affiliated with "
                "or endorsed by DoED or any government body. While reasonable care has been "
                "taken in processing this data, we do not guarantee its accuracy, "
                "completeness, or currency, and accept no liability for errors, omissions, "
                "or any decisions made based on it.",
            ], className="mb-2"),
            html.P([
                "The GIS map on this site is indicative only, meant to give a general sense "
                "of plant locations.  For authoritative records, please refer "
                "to ",
                html.A("www.doed.gov.np", href="https://www.doed.gov.np", target="_blank"),
                " or contact DoED directly.",
            ], className="mb-2"),
            html.P([
                html.Strong("NEA Operational Data & Forecasts: "),
                "System Operational Performance figures (peak demand, system loss, energy "
                "balance, consumers, revenue, and related metrics) are compiled from data "
                "published by the Nepal Electricity Authority (",
                html.A("www.nea.org.np", href="https://www.nea.org.np", target="_blank"),
                "). The Forecast Lab's projections are generated by statistical models "
                "(Linear Regression, Holt, Moving Average, ARIMA, SARIMA, and a Linear+ARIMA "
                "hybrid) applied to that historical data; they are estimates for general "
                "planning reference only, not official NEA projections, and actual outcomes "
                "may differ materially. As with the licensing data above, this is an "
                "independent, unofficial platform not affiliated with or endorsed by NEA, "
                "and we accept no liability for decisions made based on these figures or "
                "forecasts.",
            ], className="mb-0"),
        ]),

       html.Hr(style={"borderColor": "#3d5a99", "opacity": 0.4, "margin": "10px 0"}),
        html.Div("© 2026 Er. Sandeep Neupane. All rights reserved.",
                  className="small text-center"),
    ]),
])

# ── DATA-SOURCE CALLBACKS ──────────────────────────────────────────────────

# ── DATA-SOURCE CALLBACKS ──────────────────────────────────────────────────
@app.callback(
    Output("load-status", "children"),
    Output("f-type", "options"), Output("f-status", "options"), Output("f-province", "options"),
    Output("f-year", "data"),
    Output("footer-last-update", "children"),
    Input("init-once", "n_intervals"),
    Input("refresh-poll", "n_intervals"),
    prevent_initial_call=False,
)
def handle_data_source(_init, _poll):
    loader = STATE["loader"]
    last_sync = ss.get_last_sync()
    footer_update = f"🕒 Last Update: {last_sync}" if last_sync else "🕒 Last Update: —"

    if loader is None or loader.error:
        detail = (loader.error if (loader and loader.error) else STATE.get("error"))
        msg = (f"⚠️ {detail}" if detail
               else "No data loaded yet. An administrator can add a data source via /admin.")
        return (msg, [], [], [], None, footer_update)

    types = [{"label": t, "value": t} for t in loader.get_types() if t != "Transmission Line"]
    statuses = [{"label": s, "value": s} for s in loader.get_statuses()]
    # Province options come from the permanent GIS boundary layer (all 7 provinces),
    # not just the provinces that happen to have a project in the current data —
    # so the full administrative tree is always selectable.
    if getattr(de.GIS, 'provinces_loaded', False):
        province_names = sorted(de.GIS.provinces.keys())
    else:
        province_names = [p for p in loader.get_provinces() if p != "Unspecified"]
    provinces = [{"label": p, "value": p} for p in province_names]
    y_lo, y_hi = loader.get_license_year_bounds()
    y_lo, y_hi = (y_lo or 2050), (y_hi or 2085)

    status_msg = (f"✅ {len(loader.records)} records loaded — {STATE['source_label']}"
                  + (f" (last sync: {last_sync})" if last_sync else ""))
    return (status_msg, types, statuses, provinces, [y_lo, y_hi], footer_update)


# ── FILTERING HELPER ──────────────────────────────────────────────────────
def get_filtered_records(f_type, f_status, f_province, f_capacity, f_year, f_search,
                          f_date_from=None, f_date_to=None, f_cod_from=None, f_cod_to=None,
                          f_tx_length=None, f_district=None, f_local=None):
    loader = STATE["loader"]
    if loader is None or loader.error or not loader.records:
        return []
    date_from = de.parse_bs_input(f_date_from) if f_date_from else (f_year[0] if f_year else None)
    date_to = de.parse_bs_input(f_date_to, end=True) if f_date_to else (f_year[1] if f_year else None)
    cod_from = de.parse_bs_input(f_cod_from) if f_cod_from else None
    cod_to = de.parse_bs_input(f_cod_to, end=True) if f_cod_to else None
    cap_min, cap_max = CAPACITY_BIN_RANGES.get(f_capacity or "all", (None, None))
    km_min, km_max = TX_LENGTH_BIN_RANGES.get(f_tx_length or "all", (None, None))
    return loader.filter(
        types=f_type or None, statuses=f_status or None, provinces=f_province or None,
        districts=f_district or None, locals_sel=f_local or None,
        cap_min=cap_min, cap_max=cap_max, km_min=km_min, km_max=km_max,
        year_from=date_from, year_to=date_to, cod_from=cod_from, cod_to=cod_to,
        search=f_search or None,
    )


# ── CASCADING FILTER TREE ────────────────────────────────────────────────
@app.callback(
    Output("f-district", "options"),
    Input("f-province", "value"),
    Input("load-status", "children"),
    Input("refresh-poll", "n_intervals"),
)
def update_district_options(f_province, _status, _poll):
    # District options come from the permanent GIS boundary layer (all 77
    # districts, mapped to their real province) rather than only districts
    # that happen to have a project in the current data — so the full
    # Province → District tree is always selectable, and picking a province
    # narrows it to that province's real districts.
    if getattr(de.GIS, 'loaded', False) and de.GIS.district_province:
        dist_prov = dict(de.GIS.district_province)
    else:
        loader = STATE["loader"]
        if loader is None or loader.error or not loader.records:
            return []
        dist_prov = {}
        for r in loader.records:
            d = r.get("district")
            if d and d != "Unspecified" and d not in dist_prov:
                dist_prov[d] = r.get("province")
    all_districts = sorted(dist_prov)
    if not f_province:
        opts = all_districts
    else:
        opts = [d for d in all_districts if dist_prov.get(d) in f_province]
    return [{"label": d, "value": d} for d in opts]


@app.callback(
    Output("f-local", "options"),
    Input("f-district", "value"), Input("f-province", "value"),
    Input("load-status", "children"),
    Input("refresh-poll", "n_intervals"),
)
def update_local_options(f_district, f_province, _status, _poll):
    if not getattr(de.GIS, 'loaded', False):
        return []
    if f_district:
        labels = de.GIS.locals_for_districts(f_district)
    elif f_province:
        dists = [d for d, p in de.GIS.district_province.items() if p in f_province]
        labels = de.GIS.locals_for_districts(dists) if dists else []
    else:
        labels = sorted({L["label"] for L in de.GIS.locals})
    return [{"label": l, "value": l} for l in labels]


# ── KPI ROW ────────────────────────────────────────────────────────────────
@app.callback(
    Output("kpi-row", "children"),
    Input("main-tabs", "active_tab"),
    Input("f-type", "value"), Input("f-status", "value"), Input("f-province", "value"),
    Input("f-capacity", "value"), Input("f-tx-length", "value"), Input("f-year", "data"),
    Input("f-search", "value"),
    Input("f-date-from", "value"), Input("f-date-to", "value"),
    Input("f-cod-from", "value"), Input("f-cod-to", "value"),
    Input("f-district", "value"), Input("f-local", "value"),
    Input("load-status", "children"),
)
def update_kpis(tab, f_type, f_status, f_province, f_capacity, f_tx_length, f_year, f_search,
                 f_date_from, f_date_to, f_cod_from, f_cod_to, f_district, f_local, _status):
    # REQ 6: KPI summary only on Overview tab
    if tab != "overview":
        return []

    loader = STATE["loader"]
    if tab == "overview":
        recs = list(loader.records) if loader and not loader.error else []
    else:
        recs = get_filtered_records(f_type, f_status, f_province, f_capacity, f_year, f_search,
                                     f_date_from, f_date_to, f_cod_from, f_cod_to, f_tx_length,
                                     f_district, f_local)

    active_recs = [r for r in recs if r["status"] not in de.EXTRA_STATUS_ORDER]
    plant_recs = [r for r in active_recs if r["type"] != "Transmission Line"]
    tx_recs = [r for r in active_recs if r["type"] == "Transmission Line"]

    n_plants = len(plant_recs)
    plant_mw = sum(r["capacity_mw"] or 0 for r in plant_recs)
    n_operating = sum(1 for r in plant_recs if r["status"] == "Operating")
    # REQ 7: Installed Capacity with operating plants at first
    op_plants = [r for r in plant_recs if r["status"] == "Operating"]
    op_mw = sum(r["capacity_mw"] or 0 for r in op_plants)
    op_n = len(op_plants)

    n_tx = len(tx_recs)
    tx_mw = sum(r["capacity_mw"] or 0 for r in tx_recs)
    tx_km = sum(r["line_length_km"] or 0 for r in tx_recs)

    cards = [
        # REQ 7: Installed Capacity first
        kpi_card_compact("Installed Capacity", f"{op_mw:,.1f} MW",
                  f"{op_n:,} Operating Plants", "#2e7d32"),
        kpi_card_compact("Active Power Plants", f"{n_plants:,} Projects",
                  f"{plant_mw:,.1f} MW Total • {n_operating:,} operating", "#1565c0"),
        kpi_card_compact("Transmission Lines", f"{n_tx:,} Projects",
                  f"{tx_mw:,.1f} MW • {tx_km:,.1f} km circuit length", "#6a1b9a"),
    ]
    cards += nea_kpi_cards()
    # REQ 2: all KPIs (license + NEA) share one flex row that shrinks each
    # card evenly to fit, instead of a 12-col grid that wraps NEA cards
    # onto a second line once more than a handful of cards are present.
    return html.Div(
        [html.Div(c, className="kpi-flex-item") for c in cards],
        className="kpi-flex-row",
    )


@app.callback(
    Output("site-header", "style"),
    Input("init-once", "n_intervals"), Input("refresh-poll", "n_intervals"),
)
def update_header_bg(_a, _b):
    bg = ss.get_background_path()
    style = {"borderRadius": "10px", "color": "white"}
    if bg:
        style.update({
            "backgroundImage": "linear-gradient(rgba(10,20,40,0.55), rgba(10,20,40,0.55)), "
                                "url('/assets-background')",
            "backgroundSize": "cover", "backgroundPosition": "center",
        })
    else:
        style["background"] = "linear-gradient(135deg, #0b1730 0%, #16325c 100%)"
    return style


@app.callback(
    Output("ticker-bar", "children"),
    Input("load-status", "children"),
    Input("refresh-poll", "n_intervals"),
    Input("f-type", "value"), Input("f-status", "value"), Input("f-province", "value"),
    Input("f-capacity", "value"), Input("f-tx-length", "value"), Input("f-year", "data"),
    Input("f-search", "value"),
    Input("f-date-from", "value"), Input("f-date-to", "value"),
    Input("f-cod-from", "value"), Input("f-cod-to", "value"),
    Input("f-district", "value"), Input("f-local", "value"),
)
def update_ticker(_status, _poll, f_type, f_status, f_province, f_capacity, f_tx_length, f_year, f_search,
                   f_date_from, f_date_to, f_cod_from, f_cod_to, f_district, f_local):
    if not ss.get_marquee_enabled():
        return None
    loader = STATE["loader"]
    if loader is None or loader.error or not loader.records:
        return render_ticker_bar(loader)
    recs = get_filtered_records(f_type, f_status, f_province, f_capacity, f_year, f_search,
                                 f_date_from, f_date_to, f_cod_from, f_cod_to, f_tx_length,
                                 f_district, f_local)
    return render_ticker_bar(loader, recs)


# ── MERGED TAB SHELLS (License Status / License Insights) ──────────────────
def license_status_shell():
    """Power Plants / Transmission Line / GoN Studied Projects / License
    Cancelled, merged under one parent tab as sub-tabs so they share
    one slot in the main nav. Content is filled in by a dedicated
    callback keyed on the sub-tabs' own active_tab (see below) — this
    just lays out the sub-nav + an empty content slot."""
    return html.Div([
        dbc.Tabs(id="license-status-subtabs", active_tab="plants", className="sub-tabs-nav mb-3", children=[
            dbc.Tab(label="⚡ Power Plants", tab_id="plants"),
            dbc.Tab(label="🔌 Transmission Line", tab_id="transmission"),
            dbc.Tab(label="📋 GoN Studied Projects", tab_id="gon_study"),
            dbc.Tab(label="🚫 License Cancelled", tab_id="cancelled"),
        ]),
        dcc.Loading(html.Div(id="license-status-content")),
    ])


def license_insights_shell():
    """Growth Trends + Comparative Charts + Data Table, merged the same way."""
    return html.Div([
        dbc.Tabs(id="license-insights-subtabs", active_tab="growth", className="sub-tabs-nav mb-3", children=[
            dbc.Tab(label="📈 Growth Trends", tab_id="growth"),
            dbc.Tab(label="📉 Comparative Charts", tab_id="compare"),
            dbc.Tab(label="🗂️ Data Table", tab_id="table"),
        ]),
        dcc.Loading(html.Div(id="license-insights-content")),
    ])


# ── TAB CONTENT ────────────────────────────────────────────────────────────
@app.callback(
    Output("tab-content", "children"),
    Output("gis-controls-container", "style"),
    Output("filter-sidebar-col", "style"),
    Output("filter-sidebar-col", "md"),
    Output("main-content-col", "md"),
    Input("main-tabs", "active_tab"),
    Input("f-type", "value"), Input("f-status", "value"), Input("f-province", "value"),
    Input("f-capacity", "value"), Input("f-tx-length", "value"), Input("f-year", "data"),
    Input("f-search", "value"),
    Input("f-date-from", "value"), Input("f-date-to", "value"),
    Input("f-cod-from", "value"), Input("f-cod-to", "value"),
    Input("f-crs", "value"),
    Input("gis-opt-layers", "value"),
    Input("f-district", "value"), Input("f-local", "value"),
    Input("chart-style-store", "data"),
)
def render_tab(tab, f_type, f_status, f_province, f_capacity, f_tx_length, f_year, f_search,
               f_date_from, f_date_to, f_cod_from, f_cod_to, f_crs, gis_layers,
               f_district, f_local, chart_style):
    # Update global chart style state
    global CHART_STYLE_STATE
    if chart_style:
        CHART_STYLE_STATE.update(chart_style)

    loader = STATE["loader"]
    gis_controls_style = {"display": "none"}

    # The filter sidebar is only relevant to License Status / License
    # Insights, which read every f-* Input. GIS Map has its own inbuilt
    # filter scheme (province/type/stage/search) in the Leaflet sidebar,
    # so the app's filter panel stays off there; Overview, the two NEA
    # tabs, and Custom Style don't use filters at all.
    if tab in ("license_status", "license_insights"):
        sidebar_style, sidebar_md, content_md = {"display": "block"}, 3, 9
    else:
        sidebar_style, sidebar_md, content_md = {"display": "none"}, 0, 12

    # Transmission Network map: independent of the power-plant licensing
    # loader (own workbook, own admin upload/sync) — same iframe pattern
    # as the NEA tabs below, and rendered before the data-availability
    # gate so a not-yet-loaded license workbook never blocks it.
    if tab == "transmission_network":
        style_v = abs(hash(tuple(sorted(CHART_STYLE_STATE.items())))) % 100000
        return (html.Iframe(src=f"/transmission-network-map?v={style_v}",
                             style={"width": "100%", "height": "1500px", "border": "none"}),
                gis_controls_style, sidebar_style, sidebar_md, content_md)

    # NEA Operational Data / Forecast Lab: entirely independent of the
    # power-plant licensing loader above (own Google Sheet, own sync) —
    # render them here so a not-yet-loaded power-plant dataset never
    # blocks these two tabs. The `v=` query param is a cache-buster tied
    # to the current Custom Style settings: an iframe with an unchanged
    # `src` never reloads on its own, so without this, changing Custom
    # Style while already on one of these tabs would silently do nothing
    # until the tab was left and reopened.
    style_v = abs(hash(tuple(sorted(CHART_STYLE_STATE.items())))) % 100000
    if tab == "nea_operational":
        return (html.Iframe(src=f"/nea-operational-dashboard?v={style_v}",
                             style={"width": "100%", "height": "2400px", "border": "none"}),
                gis_controls_style, sidebar_style, sidebar_md, content_md)
    if tab == "nea_forecast":
        return (html.Iframe(src=f"/nea-forecast-lab?v={style_v}",
                             style={"width": "100%", "height": "900px", "border": "none"}),
                gis_controls_style, sidebar_style, sidebar_md, content_md)
    if tab == "scenario_analysis":
        return (html.Iframe(src=f"/scenario-analysis?v={style_v}",
                             style={"width": "100%", "height": "2100px", "border": "none"}),
                gis_controls_style, sidebar_style, sidebar_md, content_md)

    # Reports tab: just an options panel + a button — doesn't need license
    # data loaded or a non-empty filtered set to be worth showing, so it's
    # dispatched here rather than after the data-availability gate below
    # (the "Generate Report" button itself still checks for data at click time).
    if tab == "reports":
        return (render_reports_tab(), gis_controls_style, sidebar_style, sidebar_md, content_md)

    if loader is None or loader.error or not loader.records:
        err_detail = (loader.error if (loader and loader.error) else STATE.get("error"))
        detail = f" Details: {err_detail}" if err_detail else ""
        return dbc.Alert([
            html.Div("No project data is loaded yet.", className="fw-semibold"),
            html.Div([
                "Please check back shortly, or refer to the official DoED records at ",
                html.A("www.doed.gov.np", href="https://www.doed.gov.np", target="_blank",
                       className="alert-link"),
                "." + detail,
            ], className="small mt-1"),
        ], color="info", className="mt-3"), gis_controls_style, sidebar_style, sidebar_md, content_md

    if tab == "overview":
        try:
            all_active_recs = [r for r in loader.records if r["status"] not in de.EXTRA_STATUS_ORDER]
            return (render_overview(loader, all_active_recs), gis_controls_style,
                     sidebar_style, sidebar_md, content_md)
        except Exception:
            tb = traceback.format_exc()
            traceback.print_exc()
            return (dbc.Alert([
                html.Div("This tab hit an error while rendering: overview", className="fw-semibold"),
                html.Pre(tb, className="small mt-2", style={"whiteSpace": "pre-wrap"}),
            ], color="danger", className="mt-3"), gis_controls_style, sidebar_style, sidebar_md, content_md)

    if tab == "license_status":
        return (license_status_shell(), gis_controls_style, sidebar_style, sidebar_md, content_md)

    if tab == "license_insights":
        return (license_insights_shell(), gis_controls_style, sidebar_style, sidebar_md, content_md)

    # GIS tab: intentionally bypass the sidebar's own filter panel (f-type,
    # f-status, f-province, f-capacity, f-year, f-search, f-date/cod range,
    # f-district, f-local). That panel is hidden on this tab (see above) but
    # its values were still being fed into get_filtered_records() before,
    # which silently re-applied the "Other tabs'" filters to the map. The
    # GIS Map has its own independent stage/type/province/search filters
    # built into its Leaflet sidebar, so it should always start from the
    # full active record set, same as Overview.
    if tab == "gis":
        recs = [r for r in loader.records if r["status"] not in de.EXTRA_STATUS_ORDER]
    else:
        recs = get_filtered_records(f_type, f_status, f_province, f_capacity, f_year, f_search,
                                     f_date_from, f_date_to, f_cod_from, f_cod_to, f_tx_length,
                                     f_district, f_local)
    if not recs:
        return (dbc.Alert("No projects match the current filters.", color="warning"),
                 gis_controls_style, sidebar_style, sidebar_md, content_md)

    active_recs = [r for r in recs if r["status"] not in de.EXTRA_STATUS_ORDER]

    try:
        if tab == "plants":
            result = render_plants_tab(loader, active_recs)
        elif tab == "transmission":
            result = render_transmission_tab(loader, active_recs)
        elif tab == "gon_study":
            result = render_side_category_tab(loader, recs, "GoN Study Project", "GoN Studied Projects")
        elif tab == "cancelled":
            result = render_side_category_tab(loader, recs, "Cancelled", "License Cancelled")
        elif tab == "growth":
            result = render_growth(loader, active_recs)
        elif tab == "gis":
            gis_layers = gis_layers if gis_layers is not None else ["boundary"]
            result = render_gis_tab(loader, active_recs, f_crs or ct.CRS_WGS84,
                                     show_boundary="boundary" in gis_layers,
                                     show_pa="pa" in gis_layers)
        elif tab == "compare":
            result = render_compare(loader, active_recs)
        elif tab == "table":
            result = render_table(recs, f_crs or ct.CRS_WGS84)
        elif tab == "custom":
            result = render_custom_tab()
        else:
            result = html.Div()
        return result, gis_controls_style, sidebar_style, sidebar_md, content_md
    except Exception:
        tb = traceback.format_exc()
        traceback.print_exc()
        return (dbc.Alert([
            html.Div(f"This tab hit an error while rendering: {tab}", className="fw-semibold"),
            html.Pre(tb, className="small mt-2", style={"whiteSpace": "pre-wrap"}),
        ], color="danger", className="mt-3"), gis_controls_style, sidebar_style, sidebar_md, content_md)


# ── MERGED-TAB SUB-CONTENT CALLBACKS ────────────────────────────────────────
@app.callback(
    Output("license-status-content", "children"),
    Input("license-status-subtabs", "active_tab"),
    Input("f-type", "value"), Input("f-status", "value"), Input("f-province", "value"),
    Input("f-capacity", "value"), Input("f-tx-length", "value"), Input("f-year", "data"),
    Input("f-search", "value"),
    Input("f-date-from", "value"), Input("f-date-to", "value"),
    Input("f-cod-from", "value"), Input("f-cod-to", "value"),
    Input("f-district", "value"), Input("f-local", "value"),
)
def render_license_status_subtab(subtab, f_type, f_status, f_province, f_capacity, f_tx_length, f_year,
                                  f_search, f_date_from, f_date_to, f_cod_from, f_cod_to,
                                  f_district, f_local):
    loader = STATE["loader"]
    if loader is None or loader.error or not loader.records:
        return dbc.Alert("No project data is loaded yet.", color="info", className="mt-2")
    recs = get_filtered_records(f_type, f_status, f_province, f_capacity, f_year, f_search,
                                 f_date_from, f_date_to, f_cod_from, f_cod_to, f_tx_length,
                                 f_district, f_local)
    if not recs:
        return dbc.Alert("No projects match the current filters.", color="warning", className="mt-2")
    active_recs = [r for r in recs if r["status"] not in de.EXTRA_STATUS_ORDER]
    try:
        if subtab == "plants":
            return render_plants_tab(loader, active_recs)
        elif subtab == "transmission":
            return render_transmission_tab(loader, active_recs)
        elif subtab == "gon_study":
            return render_side_category_tab(loader, recs, "GoN Study Project", "GoN Studied Projects")
        elif subtab == "cancelled":
            return render_side_category_tab(loader, recs, "Cancelled", "License Cancelled")
        return html.Div()
    except Exception:
        tb = traceback.format_exc()
        traceback.print_exc()
        return dbc.Alert([
            html.Div(f"This tab hit an error while rendering: {subtab}", className="fw-semibold"),
            html.Pre(tb, className="small mt-2", style={"whiteSpace": "pre-wrap"}),
        ], color="danger", className="mt-3")


@app.callback(
    Output("license-insights-content", "children"),
    Input("license-insights-subtabs", "active_tab"),
    Input("f-type", "value"), Input("f-status", "value"), Input("f-province", "value"),
    Input("f-capacity", "value"), Input("f-tx-length", "value"), Input("f-year", "data"),
    Input("f-search", "value"),
    Input("f-date-from", "value"), Input("f-date-to", "value"),
    Input("f-cod-from", "value"), Input("f-cod-to", "value"),
    Input("f-district", "value"), Input("f-local", "value"),
    Input("f-crs", "value"),
)
def render_license_insights_subtab(subtab, f_type, f_status, f_province, f_capacity, f_tx_length, f_year,
                                    f_search, f_date_from, f_date_to, f_cod_from, f_cod_to,
                                    f_district, f_local, f_crs):
    loader = STATE["loader"]
    if loader is None or loader.error or not loader.records:
        return dbc.Alert("No project data is loaded yet.", color="info", className="mt-2")
    recs = get_filtered_records(f_type, f_status, f_province, f_capacity, f_year, f_search,
                                 f_date_from, f_date_to, f_cod_from, f_cod_to, f_tx_length,
                                 f_district, f_local)
    if not recs:
        return dbc.Alert("No projects match the current filters.", color="warning", className="mt-2")
    active_recs = [r for r in recs if r["status"] not in de.EXTRA_STATUS_ORDER]
    try:
        if subtab == "growth":
            return render_growth(loader, active_recs)
        elif subtab == "compare":
            return render_compare(loader, active_recs)
        elif subtab == "table":
            # Data Table intentionally uses `recs` (not active_recs) — same as
            # its old standalone main-tab behaviour, so cancelled/GoN-study
            # rows still show up in the table view.
            return render_table(recs, f_crs or ct.CRS_WGS84)
        return html.Div()
    except Exception:
        tb = traceback.format_exc()
        traceback.print_exc()
        return dbc.Alert([
            html.Div(f"This tab hit an error while rendering: {subtab}", className="fw-semibold"),
            html.Pre(tb, className="small mt-2", style={"whiteSpace": "pre-wrap"}),
        ], color="danger", className="mt-3")

# ── STATUS / PROVINCE / TYPE COLOR HELPERS ─────────────────────────────────

def get_status_color_class(status):
    """Return CSS class for styled status labels."""
    mapping = {
        "Operating": "status-label-operating",
        "Construction License": "status-label-construction",
        "Application for Construction License": "status-label-app-construction",
        "Survey License": "status-label-survey",
        "Application for Survey License": "status-label-app-survey",
        "GoN Study Project": "status-label-gon",
        "Cancelled": "status-label-cancelled",
        "Technical Clearance": "status-label-tc",
    }
    return mapping.get(status, "")


def get_province_color_class(province):
    """Return CSS class for styled province labels."""
    mapping = {
        "Koshi": "prov-label-koshi",
        "Madhesh": "prov-label-madhesh",
        "Bagmati": "prov-label-bagmati",
        "Gandaki": "prov-label-gandaki",
        "Lumbini": "prov-label-lumbini",
        "Karnali": "prov-label-karnali",
        "Sudurpaschim": "prov-label-sudurpaschim",
    }
    return mapping.get(province, "")


# REQ 4: Ordered stages
STAGE_DISPLAY_ORDER = [
    "Operating",
    "Construction License",
    "Application for Construction License",
    "Survey License",
    "Application for Survey License",
]

# REQ 5: Ordered provinces
PROVINCE_DISPLAY_ORDER = [
    "Koshi", "Madhesh", "Bagmati", "Gandaki", "Lumbini", "Karnali", "Sudurpaschim"
]

_PLACEHOLDER_WORDS = ("load", "tbd", "n/a", "na", "pending", "update", "unknown",
                      "unspecified", "-", "—", "n.a", "to be")


def _looks_like_place(s):
    if not s:
        return False
    low = s.strip().lower()
    if not low or len(low) < 2:
        return False
    return not any(w in low for w in _PLACEHOLDER_WORDS)


def _admin_units_str(r, max_each=3):
    """Provinces/Districts/Local Bodies as per the project's License
    Boundary overlapping the GIS map (province_pct / district_pct /
    local_pct, computed by GISEngine.bbox_overlap_pct against the real
    boundary polygons) — not the sheet's raw address columns. Falls back
    to the sheet address only for records with no GIS overlap on file
    (e.g. no surveyed bbox/coordinates yet)."""
    prov_pct = r.get("province_pct") or {}
    dist_pct = r.get("district_pct") or {}
    local_pct = r.get("local_pct") or []

    if prov_pct or dist_pct or local_pct:
        provs = [p for p, _ in sorted(prov_pct.items(), key=lambda kv: -kv[1])][:max_each]
        dists = [d for d, _ in sorted(dist_pct.items(), key=lambda kv: -kv[1])][:max_each]
        locals_ = [lb["name"] for lb in
                   sorted(local_pct, key=lambda lb: -(lb.get("pct") or 0))][:max_each]
        parts = []
        if provs:
            parts.append(("Province" if len(provs) == 1 else "Provinces") + ": " + ", ".join(provs))
        if dists:
            parts.append(("District" if len(dists) == 1 else "Districts") + ": " + ", ".join(dists))
        if locals_:
            parts.append(("Local Body" if len(locals_) == 1 else "Local Bodies") + ": " + ", ".join(locals_))
        if parts:
            return " • ".join(parts)

    # Fallback: no GIS boundary-overlap data for this record — use the
    # sheet's own address columns instead.
    provs = [p.strip() for p in (r.get("province") or "").split("/") if _looks_like_place(p)]
    dists = [d.strip() for d in (r.get("district") or "").split("/") if _looks_like_place(d)]
    provs = list(dict.fromkeys(provs))[:max_each]
    dists = list(dict.fromkeys(dists))[:max_each]
    local = de.record_local(r)
    local_str = local if _looks_like_place(local) else None
    parts = []
    if provs:
        parts.append(("Province" if len(provs) == 1 else "Provinces") + ": " + ", ".join(provs))
    if dists:
        parts.append(("District" if len(dists) == 1 else "Districts") + ": " + ", ".join(dists))
    if local_str:
        parts.append("Local Body: " + local_str)
    return " • ".join(parts) if parts else "Province/District: not yet resolved"


def _cat_segment(label, n, mw, extra=None):
    s = f"{label} — {mw:,.0f} MW • {n:,} Projects"
    if extra:
        s += f" • {extra}"
    return s


def _category_admin_totals(sel):
    """Aggregate capacity-weighted GIS boundary-overlap across many
    records — the same province_pct/district_pct/local_pct overlay data
    _admin_units_str() uses for a single record, combined here over a
    whole category (e.g. every operating Hydro project) so the marquee
    can list every Province/District/Local Body the category actually
    touches, not just one "largest" pick. Falls back to each record's
    own GIS-resolved province/district/local_body field when it has no
    surveyed bbox/overlap on file."""
    prov_totals = defaultdict(float)
    dist_totals = defaultdict(float)
    local_totals = defaultdict(float)
    for r in sel:
        mw = r["capacity_mw"] or 0.0
        prov_pct = r.get("province_pct") or {}
        dist_pct = r.get("district_pct") or {}
        local_pct = r.get("local_pct") or []
        if prov_pct or dist_pct or local_pct:
            for p, pct in prov_pct.items():
                if _looks_like_place(p):
                    prov_totals[p] += mw * (pct / 100.0)
            for d, pct in dist_pct.items():
                if _looks_like_place(d):
                    dist_totals[d] += mw * (pct / 100.0)
            for lb in local_pct:
                name = lb.get("name")
                pct = lb.get("pct") or 0
                if _looks_like_place(name):
                    local_totals[name] += mw * (pct / 100.0)
        else:
            p = r["province"]
            if _looks_like_place(p) and p != "Unspecified":
                prov_totals[p] += mw
            d = (r["district"] or "").split("/")[0].split("(")[0].strip()
            if _looks_like_place(d) and d != "Unspecified":
                dist_totals[d] += mw
            lb = de.record_local(r)
            if _looks_like_place(lb):
                local_totals[lb] += mw
    return prov_totals, dist_totals, local_totals


def _fmt_admin_all(totals, top_n=None):
    """Every admin unit present, ordered by capacity — the marquee's
    scroll speed already scales to content length, so listing all of
    them (rather than only the top one) doesn't need a hard cap.
    When top_n is set, only the top N entries are returned."""
    if not totals:
        return None
    ordered = sorted(totals.items(), key=lambda kv: -kv[1])
    if top_n:
        ordered = ordered[:top_n]
    return ", ".join(f"{name} ({mw:,.0f} MW)" for name, mw in ordered)
def _nea_ticker_segments():
    """NEA operational insights for the marquee — pulls straight from
    NEA.py's own kpi dict and unit_economics() (same figures shown on
    the NEA Operational Data tab), so nothing here is a separate parse
    of the sheet. Returns [] if NEA data hasn't synced yet."""
    try:
        data = NEA.get_dashboard_data() or {}
    except Exception:
        return []
    k = data.get("kpi") or {}
    if not k:
        return []
    fin = data.get("financial") or {}
    segs = []

    def yoy_txt(v):
        if v is None:
            return "n/a"
        return f"{'+' if v >= 0 else ''}{v:,.2f}% YoY"

    segs.append((f"🏭 NEA PEAK DEMAND: {k.get('latest_peak', 0):,.1f} MW ({yoy_txt(k.get('peak_growth'))})",
                 "#9fd8ff"))
    segs.append((f"📉 NEA SYSTEM LOSS: {k.get('latest_system_loss', 0):,.2f}% "
                 f"({'down' if (k.get('loss_reduction') or 0) >= 0 else 'up'} "
                 f"{abs(k.get('loss_reduction', 0)):,.2f} pts vs last year)", "#ffb4a2"))
    segs.append((f"💰 NEA REVENUE: Rs. {k.get('latest_revenue', 0):,.1f} Mn ({yoy_txt(k.get('revenue_growth'))})",
                 "#b9f6ca"))
    segs.append((f"⚡ NEA AVAILABILITY: {k.get('latest_total_avail', 0):,.1f} MU ({yoy_txt(k.get('avail_growth'))})",
                 "#ffe082"))

    # Import / export summary for the latest year, with YoY change
    imp_mu, exp_mu = fin.get("import_mu") or [], fin.get("export_mu") or []
    if imp_mu and exp_mu and len(imp_mu) == len(exp_mu):
        def latest_and_yoy(series):
            vals = [v for v in series if v is not None]
            if not vals:
                return None, None
            if len(vals) < 2 or not vals[-2]:
                return vals[-1], None
            return vals[-1], round((vals[-1] - vals[-2]) / abs(vals[-2]) * 100, 2)
        imp_latest, imp_yoy = latest_and_yoy(imp_mu)
        exp_latest, exp_yoy = latest_and_yoy(exp_mu)
        if imp_latest is not None or exp_latest is not None:
            segs.append((f"🔁 IMPORT/EXPORT (latest FY): Import {imp_latest or 0:,.1f} MU "
                         f"({yoy_txt(imp_yoy)}) | Export {exp_latest or 0:,.1f} MU ({yoy_txt(exp_yoy)})",
                         "#d0bfff"))

    try:
        econ = NEA.unit_economics()
        for key, label in [("import_rate_rs_per_unit", "Avg. Import Rate"),
                            ("export_rate_rs_per_unit", "Avg. Export Rate"),
                            ("avg_revenue_rate_rs_per_unit", "Avg. Revenue Rate")]:
            vals = [v for v in econ.get(key, []) if v is not None]
            if vals:
                segs.append((f"💵 {label}: Rs. {vals[-1]:,.2f}/kWh", "#e0e0e0"))
    except Exception:
        pass

    # Top 3 consumer classes by count, with each one's % share of the
    # latest year's total consumers.
    cg = data.get("consumers") or {}
    if cg.get("categories") and cg.get("total"):
        totals = cg["total"]
        idx = next((i for i in range(len(totals) - 1, -1, -1) if totals[i]), None)
        if idx is not None and totals[idx]:
            grand = totals[idx]
            rows = sorted(
                ((name, s[idx]) for name, s in cg["categories"].items() if idx < len(s) and s[idx] is not None),
                key=lambda kv: -kv[1],
            )[:3]
            if rows:
                parts = " | ".join(f"{name} {v:,.0f} ({v / grand * 100:,.1f}%)" for name, v in rows)
                segs.append((f"👥 TOP CONSUMER CLASSES: {parts}", "#a5d8ff"))

    # Top revenue-contributing consumer classes, with each one's % share
    # of the latest year's total revenue collection.
    sr = data.get("sales") or {}
    if sr.get("categories") and sr.get("total"):
        totals = sr["total"]
        idx = next((i for i in range(len(totals) - 1, -1, -1) if totals[i]), None)
        if idx is not None and totals[idx]:
            grand = totals[idx]
            rows = sorted(
                ((name, s[idx]) for name, s in sr["categories"].items() if idx < len(s) and s[idx] is not None),
                key=lambda kv: -kv[1],
            )[:3]
            if rows:
                parts = " | ".join(f"{name} Rs. {v:,.1f} Mn ({v / grand * 100:,.1f}%)" for name, v in rows)
                segs.append((f"💰 TOP REVENUE CONTRIBUTORS: {parts}", "#ffd6a5"))

    # Generation-source contribution — Energy (summed across the latest
    # complete fiscal year) and Power (latest month with real data) —
    # highlighting NEA RoR/PROR, NEA Storage and IPP specifically.
    src_defs = [("nea_ror", "NEA RoR/PROR"), ("nea_storage", "NEA Storage"), ("ipp", "IPP")]
    eb = data.get("energyBalanceMonthly") or {}
    if eb:
        latest_fy = sorted(eb.keys())[-1]
        entry = eb[latest_fy]
        energy_keys = ["ipp", "nea_sub", "nea_ror", "nea_storage", "nea_solar", "thermal", "import"]
        by_src = {k: sum(v for v in (entry.get(k) or []) if v is not None) for k in energy_keys}
        grand = sum(by_src.values())
        if grand:
            parts = " | ".join(f"{label} {by_src.get(key, 0):,.0f} GWh ({by_src.get(key, 0) / grand * 100:,.1f}%)"
                                for key, label in src_defs)
            segs.append((f"⚡ ENERGY MIX ({latest_fy}): {parts}", "#c3f0ca"))

    cb = data.get("capacityBalanceMonthly") or {}
    if cb:
        latest_fy = sorted(cb.keys())[-1]
        entry = cb[latest_fy]
        power_keys = ["ipp", "nea_sub", "nea_ror", "nea_storage", "import"]
        # Average across every month with real data in the latest fiscal
        # year (not just the single latest month) — NEA data is updated
        # roughly once a year, so "latest month" always landed on the
        # same year-end month rather than giving a representative
        # picture of the year's power mix.
        by_src = {}
        for k in power_keys:
            vals = [v for v in (entry.get(k) or []) if v is not None]
            by_src[k] = (sum(vals) / len(vals)) if vals else 0.0
        grand = sum(by_src.values())
        if grand:
            parts = " | ".join(f"{label} {by_src.get(key, 0):,.0f} MW ({by_src.get(key, 0) / grand * 100:,.1f}%)"
                                for key, label in src_defs)
            segs.append((f"🔌 POWER MIX (avg. FY {latest_fy}): {parts}", "#e0c3fc"))

    return segs


def build_ticker_segments(loader, recs=None):
    all_recs = recs if recs is not None else loader.records
    plants = [r for r in all_recs if r["type"] != "Transmission Line" and r["status"] in de.STATUS_ORDER]
    txs = [r for r in all_recs if r["type"] == "Transmission Line" and r["status"] in de.STATUS_ORDER]
    canc = [r for r in all_recs if r["status"] == "Cancelled"]
    gons = [r for r in all_recs if r["status"] == "GoN Study Project"]
    tcs = [r for r in all_recs if r["status"] == "Technical Clearance"]

    segs = [(_cat_segment("⚡ ACTIVE POWER PLANTS", len(plants),
                          sum(r['capacity_mw'] or 0 for r in plants)), "#ffd166")]
    for st in STAGE_DISPLAY_ORDER:
        if st not in de.STATUS_ORDER:
            continue
        sel = [r for r in plants if r["status"] == st]
        if sel:
            segs.append((_cat_segment(st, len(sel),
                                       sum(r['capacity_mw'] or 0 for r in sel)),
                         get_status_colors().get(st, "#c8d3e8")))
    km_all = sum(r["line_length_km"] or 0 for r in txs)
    segs.append((_cat_segment("🔌 TRANSMISSION", len(txs),
                              sum(r['capacity_mw'] or 0 for r in txs),
                              extra=f"{km_all:,.0f} KM"), "#7fd1ff"))
    segs.append((_cat_segment("🏛 GoN STUDY PROJECTS", len(gons),
                              sum(r['capacity_mw'] or 0 for r in gons)), "#f4b860"))
    if tcs:
        segs.append((_cat_segment("Technical Clearance", len(tcs),
                                  sum(r['capacity_mw'] or 0 for r in tcs)), "#9fb3c8"))
    segs.append((_cat_segment("🚫 LICENCE CANCELLED", len(canc),
                              sum(r['capacity_mw'] or 0 for r in canc)), "#ff8a80"))

    op = [r for r in plants if r["status"] == "Operating"]

    for tlabel, icon, sel in (
            ("HYDRO", "💧", [r for r in op if str(r["type"]).startswith("Hydro")]),
            ("SOLAR", "☀", [r for r in op if r["type"] == "Solar"])):
        if not sel:
            continue
        segs.append((_cat_segment(f"{icon} {tlabel} IN OPERATION", len(sel),
                                  sum(r['capacity_mw'] or 0 for r in sel)), "#a5f3c4"))
        prov_t, dist_t, local_t = _category_admin_totals(sel)
        for lab, totals in (("Top 5 Provinces", prov_t), ("Top 5 Districts", dist_t),
                            ("Top 5 Local Bodies", local_t)):
            txt = _fmt_admin_all(totals, top_n=5)
            if txt:
                segs.append((f"{icon} {tlabel} {lab}: {txt}", "#7be3a2"))

    ty_, tm_, td_ = de.today_bs()

    def _cod_key(r):
        t = r.get("cod_bs")
        if not t:
            return None
        return (t[0], t[1] if len(t) > 1 and t[1] else 1, t[2] if len(t) > 2 and t[2] else 1)

    def _added(year, until):
        sel = [r for r in op if _cod_key(r) and (year, 1, 1) <= _cod_key(r) <= until]
        return sel

    cur_sel = _added(ty_, (ty_, tm_, td_))
    prv_sel = _added(ty_ - 1, (ty_ - 1, tm_, td_))
    n_cur = len(cur_sel); mw_cur = sum(r["capacity_mw"] or 0 for r in cur_sel)
    n_prv = len(prv_sel); mw_prv = sum(r["capacity_mw"] or 0 for r in prv_sel)
    d_mw = mw_cur - mw_prv
    pct = (d_mw / mw_prv * 100.0) if mw_prv else (100.0 if mw_cur else 0.0)
    arrow, acol = ("▲", "#2ecc71") if d_mw >= 0 else ("▼", "#ff6b6b")
    segs.append((f"📈 Capacity added this year {ty_}: "
                 f"{mw_cur:,.0f} MW ({n_cur:,} Projects)  vs  same period {ty_-1}: "
                 f"{mw_prv:,.0f} MW ({n_prv:,})  →  {arrow} {abs(d_mw):,.0f} MW "
                 f"({pct:+.1f}%)", acol))

    last_full_year = _added(ty_ - 1, (ty_ - 1, 12, 32))
    segs.append((_cat_segment(f"📅 Connected in {ty_-1} (full year)", len(last_full_year),
                              sum(r['capacity_mw'] or 0 for r in last_full_year)), "#ffe08a"))

    yr_sel = cur_sel
    segs.append((_cat_segment(f"🆕 In operation this year {ty_}", len(yr_sel),
                              sum(r['capacity_mw'] or 0 for r in yr_sel)), "#ffe08a"))

    largest_this_year = sorted(cur_sel, key=lambda r: r["capacity_mw"] or 0, reverse=True)[:1]
    largest_ids = {id(r) for r in largest_this_year}
    for r in largest_this_year:
        prov_pct = r.get("province_pct") or {}
        dist_pct = r.get("district_pct") or {}
        local_pct = r.get("local_pct") or []
        provs = [p for p, _ in sorted(prov_pct.items(), key=lambda kv: -kv[1])] if prov_pct else [r.get("province")]
        dists = [d for d, _ in sorted(dist_pct.items(), key=lambda kv: -kv[1])] if dist_pct else [r.get("district")]
        locals_ = [lb["name"] for lb in sorted(local_pct, key=lambda lb: -(lb.get("pct") or 0))] if local_pct else [de.record_local(r)]
        prov_str = ", ".join(p for p in provs if p and p != "Unspecified") or "—"
        dist_str = ", ".join(d for d in dists if d and d != "Unspecified") or "—"
        local_str = ", ".join(l for l in locals_ if l and l != "Unspecified") or "—"
        segs.append((f"🏆 LARGEST CONNECTED {ty_}: {r['project'][:40]} | "
                     f"Capacity: {de.fmt_mw(r['capacity_mw'])} MW | "
                     f"Promoter: {textwrap.shorten(r['promoter'] or '—', 30)} | "
                     f"Province: {prov_str} | District: {dist_str} | Local Body: {local_str} | "
                     f"Connected: {de.bs_str(r['cod_bs'])}", "#7be3a2"))

    latest_candidates = sorted([r for r in cur_sel if _cod_key(r)], key=_cod_key, reverse=True)
    latest = [r for r in latest_candidates if id(r) not in largest_ids][:1]
    for r in latest:
        segs.append((f"🔌 Latest plant connected: {r['project'][:34]} — "
                     f"{de.fmt_mw(r['capacity_mw'])} MW • {_admin_units_str(r, max_each=None)} • "
                     f"{textwrap.shorten(r['promoter'] or '—', 26)} • "
                     f"COD {de.bs_str(r['cod_bs'])}", "#c9b6ff"))

    segs.extend(_nea_ticker_segments())
    return segs
    
_TICKER_BG_RGB = (0x10, 0x17, 0x26)  # matches .ticker-bar background: #101726


def _ticker_safe_color(hex_color):
    """The marquee's colored titles are drawn straight from the status/
    province/type color palettes, which were designed for chart bars on
    a light background — some are dark enough to nearly disappear
    against the ticker bar's own dark navy background (near-matching
    font/background color). Lighten any color that doesn't have enough
    contrast against that background, leaving already-legible colors
    untouched."""
    try:
        h = hex_color.lstrip("#")
        r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    except (ValueError, AttributeError, IndexError):
        return hex_color
    bg_r, bg_g, bg_b = _TICKER_BG_RGB
    # Simple perceptual distance check against the ticker background.
    dist = ((r - bg_r) ** 2 + (g - bg_g) ** 2 + (b - bg_b) ** 2) ** 0.5
    if dist >= 140:
        return hex_color
    # Blend toward white until it's clearly distinguishable from the
    # dark background, preserving the color's hue rather than replacing it.
    for step in (0.35, 0.55, 0.75):
        nr = round(r + (255 - r) * step)
        ng = round(g + (255 - g) * step)
        nb = round(b + (255 - b) * step)
        if ((nr - bg_r) ** 2 + (ng - bg_g) ** 2 + (nb - bg_b) ** 2) ** 0.5 >= 140:
            return f"#{nr:02x}{ng:02x}{nb:02x}"
    return "#e7edf3"  # safe fallback: near-white


def render_ticker_bar(loader, recs=None):
    if loader is None or not loader.records:
        return None
    try:
        segs = build_ticker_segments(loader, recs)
    except Exception:
        traceback.print_exc()
        return None
    if not segs:
        return None
    spans = []
    for text, color in segs:
        spans.append(html.Span(text, style={"color": _ticker_safe_color(color), "marginRight": "48px"}))
    track_children = spans + spans
    total_chars = sum(len(t) for t, _ in segs)
    duration = max(60, round(total_chars / 9))
    live_badge = html.Div([
        html.Span(className="ticker-live-dot"),
        html.Span("LIVE", className="ticker-live-text"),
    ], className="ticker-live-badge")
    return html.Div([
        live_badge,
        html.Div(
            html.Div(track_children, className="ticker-track",
                      style={"animationDuration": f"{duration}s"}),
            className="ticker-track-wrap",
        ),
    ], className="ticker-bar")


# ── WATERMARK HELPER ───────────────────────────────────────────────────────
def _cumsum(values):
    total, out = 0.0, []
    for v in values:
        total += v
        out.append(total)
    return out


def add_watermark(fig):
    """'Er. Sandeep Neupane' watermark as a footer caption below the
    legend (and any axis labels), not floating inside the plotted
    chart area itself."""
    legend = fig.layout.legend
    legend_at_bottom = bool(legend and legend.orientation == "h"
                             and legend.y is not None and legend.y < 0)
    wm_y = (legend.y - 0.14) if legend_at_bottom else -0.16
    current_b = fig.layout.margin.b if fig.layout.margin else None
    needed_b = 110 if legend_at_bottom else 70
    fig.update_layout(margin=dict(b=max(current_b or 0, needed_b)))
    fig.add_annotation(
        text="Er. Sandeep Neupane",
        xref="paper", yref="paper",
        x=1.0, y=wm_y, xanchor="right", yanchor="top",
        showarrow=False,
        font=dict(size=10, color="rgba(100,100,100,0.55)", family="Arial"),
    )
    return fig


def add_watermark_matplotlib(fig):
    """Add watermark BELOW the legend, as a figure-level footer caption."""
    fig.subplots_adjust(bottom=0.16)
    fig.text(0.98, 0.02, "Er. Sandeep Neupane",
             fontsize=8, color='gray', ha='right', va='bottom', alpha=0.5)
    return fig


# ── CATEGORY CARD WITH COMMON BACKGROUND ────────────────────────────────────
def flip_frame_style(bg_url=None):
    """Plain style for the frame that wraps a flip card + its chart.
    The background photo is NOT applied here — it lives only on the
    heading above the frame (see flip_heading_style). bg_url is kept as
    an accepted (ignored) argument so existing call sites don't need to
    change their signatures."""
    return {
        "borderRadius": "12px", "padding": "16px", "position": "relative",
        "backgroundColor": "#f5f6f8",
    }


def flip_heading_style(bg_url=None):
    """Plain section-title bar. bg_url is accepted (ignored) for
    call-site compatibility — the category photo now lives on the flip
    card itself (see render_category_card's photo-frame zone), matching
    the requested photo-on-top / solid-color-block infographic layout,
    rather than as a full-width photo strip above the whole section."""
    return {
        "borderRadius": "8px", "padding": "10px 16px", "marginBottom": "12px",
        "color": "#fff", "backgroundColor": "#37474f",
    }


# Solid panel look for charts/cards — no background photo is drawn behind
# them anymore (REQ: image restricted to the flip card's own photo zone).
# NOTE: no `transition` here on purpose. Plotly's built-in transition
# animates the axis range itself whenever it changes between flips, which
# is what made the chart look like it was "wobbling"/rescaling every few
# seconds. The chart should just swap to its new values cleanly, the same
# way a table's cells update — the gentle/smooth part comes from the
# card's own CSS crossfade (see _photo_frame_with_label), not from
# animating the plot itself.
def _legend_below_xaxis():
    """Shared horizontal-legend position for every chart in the app.

    Plotly's default (or a small y like -0.15/-0.18/-0.2) anchors the
    legend by its own center, so as the legend box grows (long trace
    names, secondary-axis entries) it creeps upward and sits on top of
    the x-axis tick labels / axis title instead of below them. Anchoring
    the TOP of the legend box at a fixed, more negative y instead means
    the legend always starts below the axis, however many entries it
    has. Pair this with a wide-enough bottom margin (see
    _BOTTOM_MARGIN_FOR_LEGEND) so Plotly actually reserves the room."""
    return dict(orientation="h", yanchor="top", y=-0.32, xanchor="left", x=0.0)


# Bottom margin (px) to reserve so a legend placed via
# _legend_below_xaxis() has room to sit under the x-axis without being
# clipped or overlapping the axis title/tick labels.
_BOTTOM_MARGIN_FOR_LEGEND = 100

_FLIP_PANEL_CHART_KWARGS = dict(plot_bgcolor="#ffffff", paper_bgcolor="#ffffff")


def _photo_frame_with_label(label, bg_url, height=150):
    """Photo frame (uploaded category image, or a plain placeholder
    when none has been uploaded) with the Power/Stage/Province name
    overlaid directly on top of the picture — bottom-left, on a dark
    gradient scrim so it stays readable on any photo — instead of
    being written on a separate line underneath the image."""
    photo_inner_style = {"height": "100%", "borderRadius": "10px",
                          "backgroundColor": "#e3e6ea", "position": "relative"}
    if bg_url:
        photo_inner_style.update({
            "backgroundImage": f'url("{bg_url}")',
            "backgroundSize": "cover", "backgroundPosition": "center",
            "transition": "background-image 0.6s ease-in-out",
        })
    label_overlay = html.Div(label, style={
        "position": "absolute", "left": 0, "right": 0, "bottom": 0,
        "padding": "8px 12px 6px", "color": "#fff", "fontWeight": "700",
        "fontSize": "18px", "textShadow": "0 1px 3px rgba(0,0,0,0.8)",
        "background": "linear-gradient(to top, rgba(0,0,0,0.68), rgba(0,0,0,0))",
        "borderRadius": "0 0 10px 10px",
    })
    return html.Div(
        html.Div(label_overlay, style=photo_inner_style),
        style={"backgroundColor": "#e9ebee", "borderRadius": "14px 14px 0 0",
               "padding": "10px", "height": f"{height}px", "flex": "0 0 auto"},
    )


def render_category_card(label, stage_map, total_n, total_mw, bg_url, base_color, total_km=0.0,
                          stage_order=None, is_transmission=False):
    """Two-zone infographic card: a light rounded photo frame on top
    (the category's admin-uploaded background image, or a plain
    placeholder frame when none has been uploaded) with the name
    written inside the picture itself, and a solid-color stat block
    underneath carrying the stage breakdown."""
    stage_order = stage_order or STAGE_DISPLAY_ORDER

    photo_frame = _photo_frame_with_label(label, bg_url)

    stage_rows = []
    for st in stage_order:
        if st not in stage_map:
            continue
        n, mw, km = stage_map[st]
        # REQ 3: Consistent pattern
        if is_transmission:
            detail = f"{n:,} Projects · {mw:,.1f} MW · {km:,.1f} KM"
        else:
            detail = f"{n:,} Projects · {mw:,.1f} MW"
        stage_rows.append(html.Div([
            html.Span(st, className="small", style={"color": "rgba(255,255,255,0.85)"}),
            html.Span(detail, className="small fw-semibold float-end", style={"color": "#fff"}),
        ], className="d-flex justify-content-between py-1",
           style={"borderBottom": "1px solid rgba(255,255,255,0.25)"}))

    # REQ 3: Consistent pattern for totals
    if is_transmission:
        totals_line = f"{total_n:,} Projects · {total_mw:,.1f} MW · {total_km:,.1f} KM"
    else:
        totals_line = f"{total_n:,} Projects · {total_mw:,.1f} MW"

    stat_block = html.Div([
        html.Div(totals_line, className="small text-center mb-2 fw-semibold",
                  style={"color": "rgba(255,255,255,0.9)"}),
        html.Div(stage_rows or [html.Div("No records", className="small text-center",
                                          style={"color": "rgba(255,255,255,0.75)"})]),
    ], style={
        "backgroundColor": base_color, "borderRadius": "0 0 14px 14px",
        "padding": "14px 16px", "flex": "1 1 auto",
    })

    # No fixed/limited height and no overflowY scroll here — the card
    # grows to fit every stage row so the full summary is visible by
    # default; minHeight only keeps it visually aligned with the chart
    # beside it when there's little content.
    return dbc.Card([photo_frame, stat_block],
                     key=f"cat-{label}", className="mb-3 shadow-sm flip-card-animate",
                     style={"minHeight": "360px", "height": "auto",
                            "display": "flex", "flexDirection": "column"})


def compute_breakdown(recs, key_field):
    totals = defaultdict(lambda: [0, 0.0, 0.0])
    stages = defaultdict(dict)
    for r in recs:
        k = r[key_field] or "Unspecified"
        km = r["line_length_km"] or 0.0
        totals[k][0] += 1
        totals[k][1] += r["capacity_mw"] or 0.0
        totals[k][2] += km
        entry = stages[k].setdefault(r["status"], [0, 0.0, 0.0])
        entry[0] += 1
        entry[1] += r["capacity_mw"] or 0.0
        entry[2] += km
    return totals, stages


def status_pie(recs, title):
    by_status = defaultdict(int)
    for r in recs:
        by_status[r["status"]] += 1
    colors = [get_status_colors().get(s, "#90a4ae") for s in by_status.keys()]
    fig = go.Figure(go.Pie(
        labels=list(by_status.keys()), values=list(by_status.values()), hole=0.45,
        marker_colors=colors,
    ))
    fig.update_layout(title=title, height=380, margin=dict(l=10, r=10, t=40, b=10))
    add_watermark(fig)
    return fig


# ── OVERVIEW TAB ────────────────────────────────────────────────────────────
def type_flip_chart_figure(t, stage_map, bg_url=None, y_max=None, cum_max=None):
    """Chart figure for the type flip card — shows THIS type's own
    stage breakdown, so it flips in sync with the card instead of
    staying constant. bg_url is accepted for call-site compatibility
    but not drawn here.

    Stable-axis fix: every flip used to redraw with only the stages
    that type actually has, so the x categories themselves jumped
    around (bars appearing/disappearing/reordering) on every tick.
    The x-axis now always carries the full STAGE_DISPLAY_ORDER
    (0-value bars for stages this type has none of), so that part
    stays fixed across flips.

    y_max/cum_max: optional override to force a specific axis range;
    when not supplied (the default), the axis is sized to THIS type's
    own data, so a small type isn't flattened by a much larger type's
    scale. Includes cumulative line on secondary axis."""
    use_km = (t == "Transmission Line")
    idx = 2 if use_km else 1
    unit = "KM" if use_km else "MW"
    colors = [get_status_colors().get(s, "#90a4ae") for s in STAGE_DISPLAY_ORDER]
    yvals = [stage_map.get(s, [0, 0.0, 0.0])[idx] for s in STAGE_DISPLAY_ORDER]
    cum_vals = _cumsum(yvals)

    fig = go.Figure()
    fig.add_trace(go.Bar(
        x=STAGE_DISPLAY_ORDER, y=yvals,
        marker_color=colors,
        text=[f"{v:,.1f} {unit}" if v else "" for v in yvals], textposition="outside",
        name=unit, width=0.5,
    ))
    fig.add_trace(go.Scatter(
        x=STAGE_DISPLAY_ORDER, y=cum_vals, mode="lines+markers",
        name=f"Cumulative {unit}", yaxis="y2",
        line=dict(color="#37474f", width=3, dash="dot"),
        marker=dict(size=7),
    ))
    axis_max = y_max if y_max else (max(yvals, default=0) or 1)
    axis_cum_max = cum_max if cum_max else (max(cum_vals, default=0) or 1)
    layout_kwargs = dict(
        title=f"{t} — {'Length (KM)' if use_km else 'Capacity (MW)'} by License Stage",
        height=360, yaxis_title=unit,
        yaxis2=dict(title=f"Cumulative {unit}", overlaying="y", side="right",
                     showgrid=False, range=[0, axis_cum_max * 1.15]),
        margin=dict(l=10, r=10, t=40, b=_BOTTOM_MARGIN_FOR_LEGEND),
        xaxis=dict(categoryorder="array", categoryarray=STAGE_DISPLAY_ORDER),
        yaxis=dict(range=[0, axis_max * 1.15], autorange=False),
        legend=_legend_below_xaxis(),
        **_FLIP_PANEL_CHART_KWARGS,
    )
    fig.update_layout(**layout_kwargs)
    fig = _apply_secondary_axis_setting(fig)
    add_watermark(fig)
    return fig

def province_flip_chart_figure(p, stage_map, y_max=None):
    """Chart figure for the province flip card — shows THIS province's
    own stage breakdown, flipping in sync with the card. Same fixed
    x-categories as type_flip_chart_figure.

    y_max: optional override to force a specific axis range; when not
    supplied (the default), the axis is sized to THIS province's own
    data."""
    colors = [get_status_colors().get(s, "#90a4ae") for s in STAGE_DISPLAY_ORDER]
    yvals = [stage_map.get(s, [0, 0.0, 0.0])[1] for s in STAGE_DISPLAY_ORDER]
    fig = go.Figure(go.Bar(
        x=STAGE_DISPLAY_ORDER, y=yvals,
        marker_color=colors,
        text=[f"{v:,.1f} MW" if v else "" for v in yvals], textposition="outside",
    ))
    axis_max = y_max if y_max else (max(yvals, default=0) or 1)
    fig.update_layout(
        title=f"{p} — Capacity (MW) by License Stage", height=360,
        yaxis_title="MW", margin=dict(l=10, r=10, t=40, b=10),
        xaxis=dict(categoryorder="array", categoryarray=STAGE_DISPLAY_ORDER),
        yaxis=dict(range=[0, axis_max * 1.15], autorange=False),
        **_FLIP_PANEL_CHART_KWARGS,
    )
    add_watermark(fig)
    return fig



def render_overview(loader, recs):
    """Overview: a small card flips through Power Plant types, then
    through provinces, every few seconds — and the chart beside each
    card flips WITH it, showing that type's/province's own stage
    breakdown rather than staying on one constant full-comparison
    chart. Transmission Line is excluded here so it never gets mixed
    into the Power Plants by Type card/chart (it gets its own tab).

    A System Performance Insights flip section (National vs System
    Peak Demand / Energy Import / Energy Export, sourced from the NEA
    operational data) sits at the bottom, below both Power Plants
    sections — those KPIs used to live only on the System Operational
    Performance tab; this surfaces the headline trend right on
    Overview instead."""
    card, bg_url, fig_type = _flip_card_only(0)
    prov_card, prov_bg_url, fig_prov = _overview_province_flip_card_only(0)
    perf_title, perf_card, perf_fig = _nea_perf_flip_frame_only(0)

    return html.Div([
        html.Div(html.H5("⚡ Power Plants by Type", className="m-0"),
                 id="type-flip-heading", style=flip_heading_style()),
        html.Div(
            id="type-flip-frame",
            style=flip_frame_style(),
            children=dbc.Row([
                dbc.Col(html.Div(id="type-flip-card", children=card,
                                  style={"height": "auto", "minHeight": "360px"}), md=5),
                dbc.Col(dcc.Graph(id="type-flip-chart", figure=fig_type, style={"height": "360px"}), md=7),
            ]),
        ),
        html.Hr(),
        html.Div(html.H5("🗺️ Power Plants by Province", className="m-0"),
                 id="overview-province-flip-heading", style=flip_heading_style()),
        html.Div(
            id="overview-province-flip-frame",
            style=flip_frame_style(),
            children=dbc.Row([
                dbc.Col(html.Div(id="overview-province-flip-card", children=prov_card,
                                  style={"height": "auto", "minHeight": "360px"}), md=5),
                dbc.Col(dcc.Graph(id="overview-province-flip-chart", figure=fig_prov,
                                   style={"height": "360px"}), md=7),
            ]),
        ),
        html.Hr(),
        html.Div(html.H5(perf_title, className="m-0", id="nea-perf-flip-heading-text"),
                 style=flip_heading_style()),
        html.Div(
            style=flip_frame_style(),
            children=dbc.Row([
                dbc.Col(html.Div(id="nea-perf-flip-card", children=perf_card,
                                  style={"height": "auto", "minHeight": "360px"}), md=5),
                dbc.Col(dcc.Graph(id="nea-perf-flip-chart", figure=perf_fig, style={"height": "360px"}), md=7),
            ]),
        ),
    ])


def _overview_province_flip_card_only(n):
    """Card AND chart for the current flip tick — Power Plants only,
    Transmission Line excluded (province breakdown belongs to plants)."""
    loader = STATE["loader"]
    if loader is None or loader.error or not loader.records:
        return None, None, go.Figure()
    try:
        recs = [r for r in loader.records
                if r["status"] not in de.EXTRA_STATUS_ORDER and r["type"] != "Transmission Line"]
        if not recs:
            return None, None, go.Figure()

        prov_totals, prov_stages = compute_breakdown(recs, "province")
        provinces_present = [p for p in PROVINCE_DISPLAY_ORDER if p in prov_totals] + \
                            [p for p in prov_totals if p not in PROVINCE_DISPLAY_ORDER]
        if not provinces_present:
            return None, None, go.Figure()

        p = provinces_present[n % len(provinces_present)]
        bg_url = ss.get_province_bg_url(p)
        color = get_province_colors().get(p, "#455a64")

        card = render_category_card(
            p, prov_stages[p], prov_totals[p][0], prov_totals[p][1],
            bg_url, color, stage_order=STAGE_DISPLAY_ORDER
        )
        # Axis is scaled to THIS province's own data (not a shared max
        # across every province) so each province's bars use its own
        # headroom. No animation is used, so this no longer causes any
        # rescale-wobble — the chart just renders each province's own
        # scale directly.
        fig = province_flip_chart_figure(p, prov_stages[p])
        return card, bg_url, fig
    except Exception:
        tb = traceback.format_exc()
        traceback.print_exc()
        err_card = dbc.Alert([
            html.Div("The Overview province card hit an error while rendering.",
                      className="fw-semibold small"),
            html.Pre(tb, className="small mt-1",
                      style={"whiteSpace": "pre-wrap", "maxHeight": "280px", "overflowY": "auto"}),
        ], color="danger")
        return err_card, None, go.Figure()


@app.callback(
    Output("overview-province-flip-card", "children"),
    Output("overview-province-flip-frame", "style"),
    Output("overview-province-flip-heading", "style"),
    Output("overview-province-flip-chart", "figure"),
    Input("province-flip-interval", "n_intervals"),
)
def flip_overview_province_card(n):
    card, bg_url, fig = _overview_province_flip_card_only(n)
    return card, flip_frame_style(), flip_heading_style(), fig


def _flip_card_only(n):
    """Card AND chart for the current flip tick. Transmission Line is
    filtered out here so it never mixes into the "Power Plants by Type"
    card/chart — it has its own tab (and its own flip card there)."""
    loader = STATE["loader"]
    if loader is None or loader.error or not loader.records:
        return None, None, go.Figure()
    try:
        recs = [r for r in loader.records
                if r["status"] not in de.EXTRA_STATUS_ORDER and r["type"] != "Transmission Line"]
        if not recs:
            return None, None, go.Figure()
        totals, stages = compute_breakdown(recs, "type")
        types = [t for t in de.TYPE_ORDER if t in totals and t != "Transmission Line"] + \
                [t for t in totals if t not in de.TYPE_ORDER and t != "Transmission Line"]
        if not types:
            return None, None, go.Figure()
        t = types[n % len(types)]
        bg_url = ss.get_type_bg_url(t)
        card = render_category_card(t, stages[t], totals[t][0], totals[t][1],
                                     bg_url, get_type_colors().get(t, "#607d8b"),
                                     total_km=totals[t][2], stage_order=STAGE_DISPLAY_ORDER)
        # Axis is scaled to THIS type's own data (not a shared max across
        # every type) so a small type like Solar isn't flattened by
        # Hydro's much larger scale. No animation is used, so this no
        # longer causes any rescale-wobble between flips.
        fig = type_flip_chart_figure(t, stages[t], bg_url)
        return card, bg_url, fig
    except Exception:
        tb = traceback.format_exc()
        traceback.print_exc()
        err_card = dbc.Alert([
            html.Div("The Overview activity card hit an error while rendering.",
                      className="fw-semibold small"),
            html.Pre(tb, className="small mt-1",
                      style={"whiteSpace": "pre-wrap", "maxHeight": "280px", "overflowY": "auto"}),
        ], color="danger")
        return err_card, None, go.Figure()


@app.callback(
    Output("type-flip-card", "children"),
    Output("type-flip-frame", "style"),
    Output("type-flip-heading", "style"),
    Output("type-flip-chart", "figure"),
    Input("type-flip-interval", "n_intervals"),
)
def flip_type_card(n):
    card, bg_url, fig = _flip_card_only(n)
    return card, flip_frame_style(), flip_heading_style(), fig


# ── NEA SYSTEM PERFORMANCE FLIP (Overview, top section) ─────────────────────
def _nea_perf_flip_frames():
    """Frames for the Overview's top 'System Performance Insights'
    flip section — National vs System Peak Demand, Energy Import,
    Energy Export — each pulled straight from NEA.get_dashboard_data(),
    the same live-synced cache the System Operational Performance tab
    reads. Returns [] (frames skipped) for any series that hasn't
    synced yet, rather than showing an empty/broken chart."""
    try:
        data = NEA.get_dashboard_data() or {}
    except Exception:
        data = {}
    ae = data.get("annualEnergy") or {}
    fin = data.get("financial") or {}
    frames = []

    def _last_real(seq):
        for v in reversed(seq or []):
            if v is not None:
                return v
        return None

    years_ae = ae.get("years") or []
    nat_peak = ae.get("national_peak") or []
    sys_peak = ae.get("system_peak") or []
    if years_ae and (nat_peak or sys_peak):
        fig = go.Figure()
        if nat_peak:
            fig.add_trace(go.Scatter(x=years_ae, y=nat_peak, name="National Peak (MW)",
                                      mode="lines+markers", line=dict(color="#1565c0", width=3)))
        if sys_peak:
            fig.add_trace(go.Scatter(x=years_ae, y=sys_peak, name="System Peak (MW)",
                                      mode="lines+markers", line=dict(color="#c62828", width=3)))
        fig.update_layout(title="National vs System Peak Demand (MW)", template="plotly_white",
                           height=360, legend=dict(orientation="h", y=-0.22),
                           margin=dict(t=50, b=60, l=40, r=20))
        add_watermark(fig)
        last_nat, last_sys = _last_real(nat_peak), _last_real(sys_peak)
        card = kpi_card("National Peak Demand",
                          f"{last_nat:,.1f} MW" if last_nat is not None else "—",
                          f"System Peak: {last_sys:,.1f} MW" if last_sys is not None else "System Peak: —",
                          "#1565c0")
        frames.append(("⚡ System Performance — National vs System Peak Demand", card, fig))

    years_fin = fin.get("years") or []
    imp = fin.get("import_mu") or []
    if years_fin and imp and any(v is not None for v in imp):
        fig = go.Figure([go.Bar(x=years_fin, y=imp, marker_color="#6a1b9a", name="Import (MU)")])
        fig.update_layout(title="Energy Import from India (MU)", template="plotly_white",
                           height=360, margin=dict(t=50, b=40, l=40, r=20))
        add_watermark(fig)
        last_imp = _last_real(imp)
        card = kpi_card("Latest Energy Import",
                          f"{last_imp:,.1f} MU" if last_imp is not None else "—",
                          "From India, latest fiscal year", "#6a1b9a")
        frames.append(("🔁 System Performance — Energy Import", card, fig))

    exp = fin.get("export_mu") or []
    if years_fin and exp and any(v is not None for v in exp):
        fig = go.Figure([go.Bar(x=years_fin, y=exp, marker_color="#2e7d32", name="Export (MU)")])
        fig.update_layout(title="Energy Export (MU)", template="plotly_white",
                           height=360, margin=dict(t=50, b=40, l=40, r=20))
        add_watermark(fig)
        last_exp = _last_real(exp)
        card = kpi_card("Latest Energy Export",
                          f"{last_exp:,.1f} MU" if last_exp is not None else "—",
                          "Latest fiscal year", "#2e7d32")
        frames.append(("🔁 System Performance — Energy Export", card, fig))

    return frames


def _nea_perf_flip_frame_only(n):
    frames = _nea_perf_flip_frames()
    if not frames:
        return "🏭 System Performance Insights", None, go.Figure()
    title, card, fig = frames[n % len(frames)]
    return title, card, fig


@app.callback(
    Output("nea-perf-flip-heading-text", "children"),
    Output("nea-perf-flip-card", "children"),
    Output("nea-perf-flip-chart", "figure"),
    Input("nea-perf-flip-interval", "n_intervals"),
)
def flip_nea_perf_card(n):
    return _nea_perf_flip_frame_only(n)


# ── STAGE FLIP CARD (for Plants tab) ───────────────────────────────────────
def render_single_stage_card(stage, sel_recs, bg_url, base_color, is_transmission=False):
    n = len(sel_recs)
    mw = sum(r["capacity_mw"] or 0 for r in sel_recs)
    km = sum(r["line_length_km"] or 0 for r in sel_recs) if is_transmission else 0.0

    photo_frame = _photo_frame_with_label(stage, bg_url)

    rows = []
    if is_transmission:
        # Transmission Line stage cards break down by voltage class
        # instead of province — province detail belongs to the Power
        # Plants side of the app, not here.
        volt_totals = defaultdict(lambda: [0, 0.0, 0.0])
        for r in sel_recs:
            v = r["voltage_kv"]
            key = f"{v:.0f} kV" if v else "Unspecified kV"
            volt_totals[key][0] += 1
            volt_totals[key][1] += r["capacity_mw"] or 0.0
            volt_totals[key][2] += r["line_length_km"] or 0.0
        for key in sorted(volt_totals, key=lambda k: (k == "Unspecified kV", k)):
            n_, mw_, km_ = volt_totals[key]
            rows.append(html.Div([
                html.Span(key, className="small", style={"color": "rgba(255,255,255,0.85)"}),
                html.Span(f"{n_:,} Projects · {km_:,.1f} KM · {mw_:,.1f} MW",
                          className="small fw-semibold float-end", style={"color": "#fff"}),
            ], className="d-flex justify-content-between py-1",
               style={"borderBottom": "1px solid rgba(255,255,255,0.25)"}))
    else:
        prov_totals = defaultdict(lambda: [0, 0.0])
        for r in sel_recs:
            p = r["province"] or "Unspecified"
            prov_totals[p][0] += 1
            prov_totals[p][1] += r["capacity_mw"] or 0.0
        # All 7 canonical provinces, ordered — not just the top 5 by capacity.
        ordered_provs = [p for p in PROVINCE_DISPLAY_ORDER] + \
                        [p for p in prov_totals if p not in PROVINCE_DISPLAY_ORDER]
        top_provs = [(p, prov_totals.get(p, [0, 0.0])) for p in ordered_provs]
        for p, v in top_provs:
            rows.append(html.Div([
                html.Span(p, className="small", style={"color": "rgba(255,255,255,0.85)"}),
                html.Span(f"{v[0]:,} Projects · {v[1]:,.1f} MW", className="small fw-semibold float-end",
                          style={"color": "#fff"}),
            ], className="d-flex justify-content-between py-1",
               style={"borderBottom": "1px solid rgba(255,255,255,0.25)"}))

    # REQ 3: Consistent pattern
    if is_transmission:
        totals_line = f"{n:,} Projects · {mw:,.1f} MW · {km:,.1f} KM"
    else:
        totals_line = f"{n:,} Projects · {mw:,.1f} MW"

    stat_block = html.Div([
        html.Div(totals_line, className="small text-center mb-2 fw-semibold", style={"color": "#fff"}),
        html.Div(rows or [html.Div("No records for this stage yet", className="small text-center",
                                    style={"color": "rgba(255,255,255,0.75)"})]),
    ], style={
        "backgroundColor": base_color, "borderRadius": "0 0 14px 14px",
        "padding": "14px 16px", "flex": "1 1 auto",
    })

    # No fixed height / overflowY scroll on the body — the card grows to
    # fit every row so the full stage summary shows by default.
    return dbc.Card([photo_frame, stat_block],
                     key=f"stage-{stage}", className="mb-3 shadow-sm flip-card-animate",
                     style={"minHeight": "360px", "height": "auto",
                            "display": "flex", "flexDirection": "column"})


def stage_province_chart_figure(stage, sel_recs, is_transmission=False, bg_url=None):
    """bg_url is accepted for call-site compatibility but no longer drawn
    here — the shared flip frame around the card+chart pair carries it.
    Includes cumulative line on secondary axis."""
    prov_totals = defaultdict(lambda: [0, 0.0, 0.0])
    for r in sel_recs:
        p = r["province"] or "Unspecified"
        prov_totals[p][0] += 1
        prov_totals[p][1] += r["capacity_mw"] or 0.0
        prov_totals[p][2] += r.get("line_length_km") or 0.0
    provinces_present = [p for p in PROVINCE_DISPLAY_ORDER if p in prov_totals] + \
                        [p for p in prov_totals if p not in PROVINCE_DISPLAY_ORDER]
    idx = 2 if is_transmission else 1
    unit = "KM" if is_transmission else "MW"
    yvals = [prov_totals[p][idx] for p in provinces_present]
    cum_vals = _cumsum(yvals)
    colors = [get_province_colors().get(p, "#455a64") for p in provinces_present]

    fig = go.Figure()
    fig.add_trace(go.Bar(
        x=provinces_present, y=yvals,
        marker_color=colors,
        text=[f"{v:,.1f} {unit}" for v in yvals], textposition="outside",
        name=f"{unit}", width=0.5,
    ))
    fig.add_trace(go.Scatter(
        x=provinces_present, y=cum_vals, mode="lines+markers",
        name=f"Cumulative {unit}", yaxis="y2",
        line=dict(color="#37474f", width=3, dash="dot"),
        marker=dict(size=8),
    ))
    layout_kwargs = dict(
        title=f"{stage} — {'Length (KM)' if is_transmission else 'Capacity (MW)'} by Province",
        height=360, yaxis_title=unit,
        yaxis2=dict(title=f"Cumulative {unit}", overlaying="y", side="right",
                     showgrid=False, range=[0, max(cum_vals, default=0) * 1.15 if cum_vals else 1]),
        margin=dict(l=10, r=10, t=40, b=_BOTTOM_MARGIN_FOR_LEGEND),
        legend=_legend_below_xaxis(),
        **_FLIP_PANEL_CHART_KWARGS,
    )
    fig.update_layout(**layout_kwargs)
    fig = _apply_secondary_axis_setting(fig)
    add_watermark(fig)
    return fig

def _stage_flip_card_and_chart(n, recs, is_transmission=False):
    """Returns (card, fig, bg_url) — bg_url is applied once by the caller's
    shared flip frame, not drawn again inside the card or the chart."""
    empty_fig = go.Figure()
    try:
        stages_present = [s for s in STAGE_DISPLAY_ORDER if any(r["status"] == s for r in recs)]
        if not stages_present:
            return None, empty_fig, None
        st = stages_present[n % len(stages_present)]
        sel = [r for r in recs if r["status"] == st]
        bg_url = ss.get_status_bg_url(st)
        card = render_single_stage_card(st, sel, bg_url, get_status_colors().get(st, "#90a4ae"),
                                         is_transmission=is_transmission)
        fig = stage_province_chart_figure(st, sel, is_transmission=is_transmission, bg_url=bg_url)
        return card, fig, bg_url
    except Exception:
        tb = traceback.format_exc()
        traceback.print_exc()
        err_card = dbc.Alert([
            html.Div("This stage card hit an error while rendering.", className="fw-semibold small"),
            html.Pre(tb, className="small mt-1",
                      style={"whiteSpace": "pre-wrap", "maxHeight": "280px", "overflowY": "auto"}),
        ], color="danger")
        return err_card, empty_fig, None


def _stage_flip_card_only(n, recs, is_transmission=False):
    """Card-only version — the chart next to it is now the tab's static
    full stage-breakdown chart (fig_stage), not a per-tick chart."""
    try:
        stages_present = [s for s in STAGE_DISPLAY_ORDER if any(r["status"] == s for r in recs)]
        if not stages_present:
            return None, None
        st = stages_present[n % len(stages_present)]
        sel = [r for r in recs if r["status"] == st]
        bg_url = ss.get_status_bg_url(st)
        card = render_single_stage_card(st, sel, bg_url, get_status_colors().get(st, "#90a4ae"),
                                         is_transmission=is_transmission)
        return card, bg_url
    except Exception:
        tb = traceback.format_exc()
        traceback.print_exc()
        err_card = dbc.Alert([
            html.Div("This stage card hit an error while rendering.", className="fw-semibold small"),
            html.Pre(tb, className="small mt-1",
                      style={"whiteSpace": "pre-wrap", "maxHeight": "280px", "overflowY": "auto"}),
        ], color="danger")
        return err_card, None


@app.callback(
    Output("plants-stage-flip-card", "children"),
    Output("plants-stage-flip-frame", "style"),
    Output("plants-stage-flip-heading", "style"),
    Input("plants-stage-next-btn", "n_clicks"),
    Input("f-type", "value"), Input("f-status", "value"), Input("f-province", "value"),
    Input("f-capacity", "value"), Input("f-tx-length", "value"), Input("f-year", "data"),
    Input("f-search", "value"),
    Input("f-date-from", "value"), Input("f-date-to", "value"),
    Input("f-cod-from", "value"), Input("f-cod-to", "value"),
    Input("f-district", "value"), Input("f-local", "value"),
)
def flip_plants_stage_card(n_clicks, f_type, f_status, f_province, f_capacity, f_tx_length, f_year,
                            f_search, f_date_from, f_date_to, f_cod_from, f_cod_to,
                            f_district, f_local):
    loader = STATE["loader"]
    if loader is None or loader.error or not loader.records:
        return None, flip_frame_style(), flip_heading_style(None)
    recs = get_filtered_records(f_type, f_status, f_province, f_capacity, f_year, f_search,
                                 f_date_from, f_date_to, f_cod_from, f_cod_to, f_tx_length,
                                 f_district, f_local)
    plant_recs = [r for r in recs if r["type"] != "Transmission Line"
                  and r["status"] not in de.EXTRA_STATUS_ORDER]
    # Only the "Next →" button advances the step; a filter change re-renders
    # at the same step (safely wrapped by len() inside _stage_flip_card_only)
    # so the card always matches the current filter selection.
    card, bg_url = _stage_flip_card_only(n_clicks or 0, plant_recs)
    return card, flip_frame_style(), flip_heading_style(bg_url)



@app.callback(
    Output("tx-stage-flip-card", "children"),
    Output("tx-stage-flip-frame", "style"),
    Output("tx-stage-flip-heading", "style"),
    Input("tx-stage-next-btn", "n_clicks"),
    Input("f-type", "value"), Input("f-status", "value"), Input("f-province", "value"),
    Input("f-capacity", "value"), Input("f-tx-length", "value"), Input("f-year", "data"),
    Input("f-search", "value"),
    Input("f-date-from", "value"), Input("f-date-to", "value"),
    Input("f-cod-from", "value"), Input("f-cod-to", "value"),
    Input("f-district", "value"), Input("f-local", "value"),
)
def flip_tx_stage_card(n_clicks, f_type, f_status, f_province, f_capacity, f_tx_length, f_year,
                        f_search, f_date_from, f_date_to, f_cod_from, f_cod_to,
                        f_district, f_local):
    loader = STATE["loader"]
    if loader is None or loader.error or not loader.records:
        return None, flip_frame_style(), flip_heading_style(None)
    recs = get_filtered_records(f_type, f_status, f_province, f_capacity, f_year, f_search,
                                 f_date_from, f_date_to, f_cod_from, f_cod_to, f_tx_length,
                                 f_district, f_local)
    tx_recs = [r for r in recs if r["type"] == "Transmission Line"
               and r["status"] not in de.EXTRA_STATUS_ORDER]
    card, bg_url = _stage_flip_card_only(n_clicks or 0, tx_recs, is_transmission=True)
    return card, flip_frame_style(), flip_heading_style(bg_url)


# ── POWER PLANTS TAB ────────────────────────────────────────────────────────
def render_plants_tab(loader, recs):
    plant_recs = [r for r in recs if r["type"] != "Transmission Line"]
    if not plant_recs:
        return dbc.Alert("No power-plant records match the current filters.", color="info")

    stage_totals, _ = compute_breakdown(plant_recs, "status")
    stages_present = [s for s in STAGE_DISPLAY_ORDER if s in stage_totals]

    # REQ 3: Consistent pattern with styled status labels
    # CSS Grid with fixed column widths (same pattern as the Transmission
    # tab's equivalent list below) — plain flexbox `justify-content-between`
    # only pins the first/last span and free-floats everything in between,
    # so with labels of different lengths per stage the numbers land in a
    # different horizontal position on every row ("zigzag"). A grid keeps
    # every column's data lined up under its own header regardless of how
    # long the stage name or the numbers happen to be.
    stage_rows = []
    for st in stages_present:
        color_cls = get_status_color_class(st)
        stage_rows.append(html.Div([
            html.Span(st, className=f"fw-semibold {color_cls}"),
            html.Span(f"{stage_totals[st][0]:,} Projects", className="text-muted",
                      style={"textAlign": "center"}),
            html.Span(f"{stage_totals[st][1]:,.1f} MW", className="fw-semibold",
                      style={"textAlign": "right"}),
        ], className="border-bottom py-2", style={
            "display": "grid",
            "gridTemplateColumns": "2fr 1fr 1fr",
            "alignItems": "center",
        }))

    colors = [get_status_colors().get(s, "#90a4ae") for s in stages_present]
    mw_values = [stage_totals[s][1] for s in stages_present]
    cum_mw = _cumsum(mw_values)
    fig_stage = go.Figure()
    fig_stage.add_trace(go.Bar(
        x=stages_present, y=mw_values,
        marker_color=colors,
        text=[f"{v:,.0f} MW" for v in mw_values], textposition="outside",
        name="Capacity (MW)", width=0.5,
    ))
    fig_stage.add_trace(go.Scatter(
        x=stages_present, y=cum_mw, mode="lines+markers",
        name="Cumulative Capacity", yaxis="y2",
        line=dict(color="#37474f", width=3, dash="dot"),
        marker=dict(size=8),
    ))
    fig_stage.update_layout(
        title="Power Plants — Capacity (MW) by License Stage",
        height=420, yaxis_title="Capacity (MW)",
        yaxis2=dict(title="Cumulative Capacity (MW)", overlaying="y", side="right",
                     showgrid=False, range=[0, max(cum_mw) * 1.15 if cum_mw else 1]),
        margin=dict(l=10, r=10, t=40, b=_BOTTOM_MARGIN_FOR_LEGEND),
        legend=_legend_below_xaxis(),
    )
    fig_stage = _apply_secondary_axis_setting(fig_stage)
    add_watermark(fig_stage)

    # Stage flip card (animated card only) + static full stage-breakdown
    # chart alongside it. The chart no longer flips/rescales every tick —
    # it's the same stable fig_stage comparison used below.
    #
    # NOTE: unlike the Overview tab, this card does NOT auto-flip on a
    # timer — it only advances when the "Next →" button below it is
    # clicked, and it re-renders immediately (frozen on the current
    # step) whenever a filter changes, so what's shown always matches
    # the current filter selection. See flip_plants_stage_card().
    stage_card0, stage_bg0 = _stage_flip_card_only(0, plant_recs)
    stage_flip_row = html.Div([
        html.Div(html.H5("⚡ License Stage ", className="m-0"),
                 id="plants-stage-flip-heading", style=flip_heading_style(stage_bg0)),
        html.Div(
            id="plants-stage-flip-frame",
            style=flip_frame_style(),
            children=dbc.Row([
                dbc.Col(html.Div(id="plants-stage-flip-card", children=stage_card0,
                                  style={"height": "auto", "minHeight": "360px"}), md=5),
                dbc.Col(dcc.Graph(figure=fig_stage, style={"height": "360px"}), md=7),
            ]),
        ),
        html.Div(
            dbc.Button("Next →", id="plants-stage-next-btn", color="primary",
                       outline=True, size="sm", n_clicks=0),
            className="text-center mt-2",
        ),
    ])

    stage_section = dbc.Row([
        dbc.Col(html.Div([html.H5("All License Stages")] + stage_rows), md=12),
    ], className="mb-4")
    stage_section = html.Div([stage_flip_row, html.Hr(), stage_section])

    # REQ 8: Animated province slides in By Province sub-tab
    prov_totals, prov_stages = compute_breakdown(plant_recs, "province")
    provinces_present = [p for p in PROVINCE_DISPLAY_ORDER if p in prov_totals] + \
                        [p for p in prov_totals if p not in PROVINCE_DISPLAY_ORDER]

    prov_colors = [get_province_colors().get(p, "#455a64") for p in provinces_present]
    prov_mw_values = [prov_totals[p][1] for p in provinces_present]
    prov_cum_mw = _cumsum(prov_mw_values)
    fig_prov = go.Figure()
    fig_prov.add_trace(go.Bar(
        x=provinces_present, y=prov_mw_values,
        marker_color=prov_colors,
        text=[f"{prov_totals[p][0]:,} Projects" for p in provinces_present], textposition="outside",
        name="Capacity (MW)", width=0.5,
    ))
    fig_prov.add_trace(go.Scatter(
        x=provinces_present, y=prov_cum_mw, mode="lines+markers",
        name="Cumulative Capacity", yaxis="y2",
        line=dict(color="#37474f", width=3, dash="dot"),
        marker=dict(size=8),
    ))
    fig_prov.update_layout(
        title="Power Plant Capacity by Province", height=460,
        yaxis_title="Capacity (MW)",
        yaxis2=dict(title="Cumulative Capacity (MW)", overlaying="y", side="right",
                     showgrid=False, range=[0, max(prov_cum_mw) * 1.15 if prov_cum_mw else 1]),
        margin=dict(l=10, r=10, t=40, b=_BOTTOM_MARGIN_FOR_LEGEND),
        legend=_legend_below_xaxis(),
    )
    fig_prov = _apply_secondary_axis_setting(fig_prov)
    add_watermark(fig_prov)

    # Province flip card (animated card only) — the chart alongside it is
    # the same stable fig_prov comparison, not a per-tick chart.
    #
    # NOTE: same manual-advance-only behavior as the License Stage card
    # above — no timer, "Next →" click only, immediate re-render on
    # filter change. See flip_province_card().
    prov_card, prov_bg_url = _province_flip_card_only(0, plant_recs)

    # REQ 8: Animated province slide section — background photo on heading only
    province_slide_section = html.Div([
        html.Div(html.H5("🗺️ Province Overview ", className="m-0"),
                 id="province-flip-heading", style=flip_heading_style(prov_bg_url)),
        html.Div(
            id="province-flip-frame",
            style=flip_frame_style(),
            children=dbc.Row([
                dbc.Col(html.Div(id="province-flip-card", children=prov_card,
                                  style={"height": "auto", "minHeight": "360px"}), md=5),
                dbc.Col(dcc.Graph(figure=fig_prov, style={"height": "360px"}), md=7),
            ]),
        ),
        html.Div(
            dbc.Button("Next →", id="province-next-btn", color="primary",
                       outline=True, size="sm", n_clicks=0),
            className="text-center mt-2",
        ),
    ])

    return dbc.Tabs(id="plants-subtabs", active_tab="stage", className="sub-tabs-nav", children=[
        dbc.Tab(stage_section, label="License Stage", tab_id="stage",
                tab_style={"marginTop": "10px"}),
        dbc.Tab(province_slide_section,
                label="By Province", tab_id="by-province", tab_style={"marginTop": "10px"}),
    ])


def _province_flip_card_only(n, recs):
    """Card-only version for Power Plants > By Province — the chart next
    to it is now the tab's static full province-breakdown chart (fig_prov)."""
    try:
        prov_totals, prov_stages = compute_breakdown(recs, "province")
        provinces_present = [p for p in PROVINCE_DISPLAY_ORDER if p in prov_totals] + \
                            [p for p in prov_totals if p not in PROVINCE_DISPLAY_ORDER]
        if not provinces_present:
            return None, None

        p = provinces_present[n % len(provinces_present)]
        bg_url = ss.get_province_bg_url(p)
        color = get_province_colors().get(p, "#455a64")

        card = render_category_card(
            p, prov_stages[p], prov_totals[p][0], prov_totals[p][1],
            bg_url, color, stage_order=STAGE_DISPLAY_ORDER
        )
        return card, bg_url
    except Exception:
        tb = traceback.format_exc()
        traceback.print_exc()
        err_card = dbc.Alert([
            html.Div("Province card hit an error while rendering.", className="fw-semibold small"),
            html.Pre(tb, className="small mt-1",
                      style={"whiteSpace": "pre-wrap", "maxHeight": "280px", "overflowY": "auto"}),
        ], color="danger")
        return err_card, None


@app.callback(
    Output("province-flip-card", "children"),
    Output("province-flip-frame", "style"),
    Output("province-flip-heading", "style"),
    Input("province-next-btn", "n_clicks"),
    Input("f-type", "value"), Input("f-status", "value"), Input("f-province", "value"),
    Input("f-capacity", "value"), Input("f-tx-length", "value"), Input("f-year", "data"),
    Input("f-search", "value"),
    Input("f-date-from", "value"), Input("f-date-to", "value"),
    Input("f-cod-from", "value"), Input("f-cod-to", "value"),
    Input("f-district", "value"), Input("f-local", "value"),
)
def flip_province_card(n_clicks, f_type, f_status, f_province, f_capacity, f_tx_length, f_year,
                        f_search, f_date_from, f_date_to, f_cod_from, f_cod_to,
                        f_district, f_local):
    loader = STATE["loader"]
    if loader is None or loader.error or not loader.records:
        return None, flip_frame_style(), flip_heading_style(None)
    recs = get_filtered_records(f_type, f_status, f_province, f_capacity, f_year, f_search,
                                 f_date_from, f_date_to, f_cod_from, f_cod_to, f_tx_length,
                                 f_district, f_local)
    plant_recs = [r for r in recs if r["type"] != "Transmission Line"
                  and r["status"] not in de.EXTRA_STATUS_ORDER]
    card, bg_url = _province_flip_card_only(n_clicks or 0, plant_recs)
    return card, flip_frame_style(), flip_heading_style(bg_url)


# ── TRANSMISSION TAB ────────────────────────────────────────────────────────
def render_transmission_tab(loader, recs):
    tx_recs = [r for r in recs if r["type"] == "Transmission Line"]
    if not tx_recs:
        return dbc.Alert("No transmission-line records match the current filters.", color="info")

    total_n = len(tx_recs)
    total_km = sum(r["line_length_km"] or 0 for r in tx_recs)
    total_mw = sum(r["capacity_mw"] or 0 for r in tx_recs)
    n_volt_classes = len({r["voltage_kv"] for r in tx_recs if r["voltage_kv"]})

    # REQ 3: Consistent KM pattern
    kpis = dbc.Row([
        dbc.Col(kpi_card("Total Lines", f"{total_n:,} Projects", "matching current filters", "#6a1b9a"), md=3),
        dbc.Col(kpi_card("Total Length", f"{total_km:,.0f} KM", "circuit length", "#1565c0"), md=3),
        dbc.Col(kpi_card("Total Capacity", f"{total_mw:,.1f} MW", "transfer capacity", "#2e7d32"), md=3),
        dbc.Col(kpi_card("Voltage Classes", f"{n_volt_classes}", "distinct kV levels", "#e65100"), md=3),
    ], className="g-3 mb-4")

    stage_totals = defaultdict(lambda: [0, 0.0, 0.0])
    for r in tx_recs:
        s = stage_totals[r["status"]]
        s[0] += 1
        s[1] += r["line_length_km"] or 0
        s[2] += r["capacity_mw"] or 0
    stages_present = [s for s in STAGE_DISPLAY_ORDER if s in stage_totals]

    # REQ 3: Consistent pattern with KM
    stage_rows = []
    for st in stages_present:
        color_cls = get_status_color_class(st)
        stage_rows.append(html.Div([
            html.Span(st, className=f"fw-semibold {color_cls}"),
            html.Span(f"{stage_totals[st][0]:,} Projects", className="text-muted",
                      style={"textAlign": "center"}),
            html.Span(f"{stage_totals[st][1]:,.0f} KM", className="text-muted",
                      style={"textAlign": "center"}),
            html.Span(f"{stage_totals[st][2]:,.1f} MW", className="fw-semibold",
                      style={"textAlign": "right"}),
        ], className="border-bottom py-2", style={
            "display": "grid",
            "gridTemplateColumns": "2fr 1fr 1fr 1fr",
            "alignItems": "center",
        }))

    colors = [get_status_colors().get(s, "#90a4ae") for s in stages_present]
    km_values = [stage_totals[s][1] for s in stages_present]
    cum_km = _cumsum(km_values)
    fig_stage = go.Figure()
    fig_stage.add_trace(go.Bar(
        x=stages_present, y=km_values,
        marker_color=colors,
        text=[f"{v:,.0f} KM" for v in km_values], textposition="outside",
        name="Length (KM)", width=0.5,
    ))
    fig_stage.add_trace(go.Scatter(
        x=stages_present, y=cum_km, mode="lines+markers",
        name="Cumulative Length", yaxis="y2",
        line=dict(color="#37474f", width=3, dash="dot"),
        marker=dict(size=8),
    ))
    fig_stage.update_layout(
        title="Transmission Lines — Length (KM) by License Stage",
        height=420, yaxis_title="Length (KM)",
        yaxis2=dict(title="Cumulative Length (KM)", overlaying="y", side="right",
                     showgrid=False, range=[0, max(cum_km) * 1.15 if cum_km else 1]),
        margin=dict(l=10, r=10, t=40, b=_BOTTOM_MARGIN_FOR_LEGEND),
        legend=_legend_below_xaxis(),
    )
    fig_stage = _apply_secondary_axis_setting(fig_stage)
    add_watermark(fig_stage)

    # No auto-flip on this tab — the card only advances via the "Next →"
    # button below it, and re-renders immediately on any filter change
    # (see flip_tx_stage_card()), so it always reflects the current
    # filter selection rather than cycling on a timer.
    tx_card0, tx_bg0 = _stage_flip_card_only(0, tx_recs, is_transmission=True)
    stage_flip_row = html.Div([
        html.Div(html.H5("🔌 License Stage ", className="m-0"),
                 id="tx-stage-flip-heading", style=flip_heading_style(tx_bg0)),
        html.Div(
            id="tx-stage-flip-frame",
            style=flip_frame_style(),
            children=dbc.Row([
                dbc.Col(html.Div(id="tx-stage-flip-card", children=tx_card0,
                                  style={"height": "auto", "minHeight": "360px"}), md=5),
                dbc.Col(dcc.Graph(figure=fig_stage, style={"height": "360px"}), md=7),
            ]),
        ),
        html.Div(
            dbc.Button("Next →", id="tx-stage-next-btn", color="primary",
                       outline=True, size="sm", n_clicks=0),
            className="text-center mt-2",
        ),
    ])

    stage_section = dbc.Row([
        dbc.Col(html.Div([html.H5("All License Stages")] + stage_rows), md=12),
    ], className="mb-4")
    stage_section = html.Div([stage_flip_row, html.Hr(), stage_section])

    by_volt = defaultdict(lambda: [0, 0.0, 0.0])
    for r in tx_recs:
        if r["voltage_kv"]:
            v = by_volt[r["voltage_kv"]]
            v[0] += 1
            v[1] += r["line_length_km"] or 0
            v[2] += r["capacity_mw"] or 0
    volts = sorted(by_volt.keys())

    # REQ 3: Consistent KM pattern — CSS Grid, same reasoning as the
    # License Stage list above: fixed columns keep every row's Projects/
    # KM/MW numbers lined up under each other instead of drifting
    # left/right depending on how long the voltage-class label is.
    volt_rows = []
    for v in volts:
        volt_rows.append(html.Div([
            html.Span(f"{v:.0f} kV", className="fw-semibold"),
            html.Span(f"{by_volt[v][0]:,} Projects", className="text-muted",
                      style={"textAlign": "center"}),
            html.Span(f"{by_volt[v][1]:,.0f} KM", className="text-muted",
                      style={"textAlign": "center"}),
            html.Span(f"{by_volt[v][2]:,.1f} MW", className="fw-semibold",
                      style={"textAlign": "right"}),
        ], className="border-bottom py-2", style={
            "display": "grid",
            "gridTemplateColumns": "1.3fr 1fr 1fr 1fr",
            "alignItems": "center",
        }))

    volt_km_values = [by_volt[v][1] for v in volts]
    volt_cum_km = _cumsum(volt_km_values)
    fig_volt = go.Figure()
    fig_volt.add_trace(go.Bar(
        x=[f"{v:.0f} kV" for v in volts], y=volt_km_values,
        marker_color="#6a1b9a", text=[by_volt[v][0] for v in volts], textposition="outside",
        name="Length (KM)", width=0.5,
    ))
    fig_volt.add_trace(go.Scatter(
        x=[f"{v:.0f} kV" for v in volts], y=volt_cum_km, mode="lines+markers",
        name="Cumulative Length", yaxis="y2",
        line=dict(color="#37474f", width=3, dash="dot"),
        marker=dict(size=8),
    ))
    fig_volt.update_layout(
        title="Length (KM) by Voltage Class", height=420,
        yaxis_title="Length (KM)",
        yaxis2=dict(title="Cumulative Length (KM)", overlaying="y", side="right",
                     showgrid=False, range=[0, max(volt_cum_km) * 1.15 if volt_cum_km else 1]),
        margin=dict(l=10, r=10, t=40, b=_BOTTOM_MARGIN_FOR_LEGEND),
        legend=_legend_below_xaxis(),
    )
    fig_volt = _apply_secondary_axis_setting(fig_volt)
    add_watermark(fig_volt)

    volt_section = dbc.Row([
        dbc.Col(html.Div([html.H5("By Voltage Class")] + volt_rows), md=5),
        dbc.Col(dcc.Graph(figure=fig_volt), md=7),
    ])

    return html.Div([
        kpis,
        dbc.Tabs(id="tx-subtabs", active_tab="stage", className="sub-tabs-nav", children=[
            dbc.Tab(stage_section, label="License Stage", tab_id="stage",
                    tab_style={"marginTop": "10px"}),
            dbc.Tab(volt_section, label="By Voltage Class", tab_id="by-voltage",
                    tab_style={"marginTop": "10px"}),
        ]),
    ])



# ── SIDE CATEGORY TABS (GoN Study, Cancelled) ────────────────────────────
def render_side_category_tab(loader, recs, status_value, page_title):
    side_recs = [r for r in recs if r["status"] == status_value]
    color = de.EXTRA_STATUS_COLORS.get(status_value, "#455a64")
    if not side_recs:
        return dbc.Alert(f"No {page_title.lower()} records match the current filters.",
                          color="info")

    plant_recs = [r for r in side_recs if r["type"] != "Transmission Line"]
    tx_recs = [r for r in side_recs if r["type"] == "Transmission Line"]
    total_mw = sum(r["capacity_mw"] or 0 for r in plant_recs)
    total_km = sum(r["line_length_km"] or 0 for r in tx_recs)

    # REQ 3: Consistent pattern
    kpis = dbc.Row([
        dbc.Col(kpi_card("Total Records", f"{len(side_recs):,} Projects", page_title, color), md=3),
        dbc.Col(kpi_card("Power Plants", f"{len(plant_recs):,} Projects", f"{total_mw:,.1f} MW", color), md=3),
        dbc.Col(kpi_card("Transmission Lines", f"{len(tx_recs):,} Projects", f"{total_km:,.0f} KM", color), md=3),
    ], className="g-3 mb-4")

    by_type, _ = compute_breakdown(side_recs, "type")
    types = [t for t in de.TYPE_ORDER if t in by_type] + [t for t in by_type if t not in de.TYPE_ORDER]
    type_colors = [get_type_colors().get(t, "#607d8b") for t in types]
    type_counts = [by_type[t][0] for t in types]
    cum_type_counts = _cumsum([float(c) for c in type_counts])
    fig_type = go.Figure()
    fig_type.add_trace(go.Bar(
        x=types, y=type_counts, marker_color=type_colors,
        text=[f"{c:,} Projects" for c in type_counts], textposition="outside",
        name="Projects", width=0.5,
    ))
    fig_type.add_trace(go.Scatter(
        x=types, y=cum_type_counts, mode="lines+markers",
        name="Cumulative Projects", yaxis="y2",
        line=dict(color="#37474f", width=3, dash="dot"),
        marker=dict(size=8),
    ))
    fig_type.update_layout(
        title=f"{page_title} — Count by Project Type", height=380,
        yaxis_title="Number of projects",
        yaxis2=dict(title="Cumulative Projects", overlaying="y", side="right",
                     showgrid=False, range=[0, max(cum_type_counts) * 1.15 if cum_type_counts else 1]),
        margin=dict(l=10, r=10, t=40, b=_BOTTOM_MARGIN_FOR_LEGEND),
        legend=_legend_below_xaxis(),
    )
    fig_type = _apply_secondary_axis_setting(fig_type)
    add_watermark(fig_type)

    by_prov, _ = compute_breakdown(side_recs, "province")
    provs = [p for p in PROVINCE_DISPLAY_ORDER if p in by_prov] + [p for p in by_prov if p not in PROVINCE_DISPLAY_ORDER]
    prov_colors = [get_province_colors().get(p, "#455a64") for p in provs]
    prov_counts = [by_prov[p][0] for p in provs]
    cum_prov_counts = _cumsum([float(c) for c in prov_counts])
    fig_prov = go.Figure()
    fig_prov.add_trace(go.Bar(
        x=provs, y=prov_counts, marker_color=prov_colors,
        text=[f"{c:,} Projects" for c in prov_counts], textposition="outside",
        name="Projects", width=0.5,
    ))
    fig_prov.add_trace(go.Scatter(
        x=provs, y=cum_prov_counts, mode="lines+markers",
        name="Cumulative Projects", yaxis="y2",
        line=dict(color="#37474f", width=3, dash="dot"),
        marker=dict(size=8),
    ))
    fig_prov.update_layout(
        title=f"{page_title} — Count by Province", height=380,
        yaxis_title="Number of projects",
        yaxis2=dict(title="Cumulative Projects", overlaying="y", side="right",
                     showgrid=False, range=[0, max(cum_prov_counts) * 1.15 if cum_prov_counts else 1]),
        margin=dict(l=10, r=10, t=40, b=_BOTTOM_MARGIN_FOR_LEGEND),
        legend=_legend_below_xaxis(),
    )
    fig_prov = _apply_secondary_axis_setting(fig_prov)
    add_watermark(fig_prov)

    return html.Div([
        kpis,
        dbc.Row([dbc.Col(dcc.Graph(figure=fig_type), md=6),
                 dbc.Col(dcc.Graph(figure=fig_prov), md=6)]),
        html.Hr(),
        render_table(side_recs, ct.CRS_WGS84),
    ])


# ── GROWTH TAB ──────────────────────────────────────────────────────────────
def render_growth(loader, recs):
    """REQ 10: Separate charts for Transmission Lines and active Power Plant stages."""
    plants = [r for r in recs if r["type"] != "Transmission Line"]
    tx_lines = [r for r in recs if r["type"] == "Transmission Line"]

    # Power Plants Growth
    plant_series = loader.yearly_series(plants, key_field="type")
    plant_years = sorted(plant_series.keys())
    all_plant_types = sorted({k for y in plant_years for k in plant_series[y].keys()})

    fig_plant_cap = go.Figure()
    for t in all_plant_types:
        fig_plant_cap.add_trace(go.Scatter(
            x=[str(y) for y in plant_years],
            y=[plant_series[y].get(t, [0, 0])[1] for y in plant_years],
            mode="lines+markers", name=t,
            line=dict(color=get_type_colors().get(t, "#607d8b")),
        ))
    # Secondary axis: cumulative installed capacity of Operating plants only
    # (running total of capacity that has actually come online, year by
    # year — distinct from the per-year "licensed capacity by type" lines
    # on the primary axis).
    operating_plants = [r for r in plants if r["status"] == "Operating"]
    op_series = loader.yearly_series(operating_plants, key_field="status")
    op_by_year = {y: op_series.get(y, {}).get("Operating", [0, 0.0])[1] for y in plant_years}
    cum = 0.0
    cum_capacity = []
    for y in plant_years:
        cum += op_by_year.get(y, 0.0)
        cum_capacity.append(cum)
    fig_plant_cap.add_trace(go.Scatter(
        x=[str(y) for y in plant_years], y=cum_capacity,
        mode="lines+markers", name="Cumulative Installed Capacity (Operating)",
        line=dict(color="#2e7d32", width=3, dash="dot"), yaxis="y2",
    ))
    fig_plant_cap.update_layout(
        title="Power Plants — Licensed Capacity by Year (B.S.)",
        xaxis_title="B.S. Year", yaxis_title="Capacity Licensed This Year (MW)",
        yaxis2=dict(title="Cumulative Installed Capacity — Operating (MW)",
                     overlaying="y", side="right", showgrid=False),
        height=480, legend=_legend_below_xaxis(),
        margin=dict(b=_BOTTOM_MARGIN_FOR_LEGEND),
    )
    fig_plant_cap = _apply_secondary_axis_setting(fig_plant_cap)
    add_watermark(fig_plant_cap)

    fig_plant_count = go.Figure()
    for t in all_plant_types:
        fig_plant_count.add_trace(go.Bar(
            x=[str(y) for y in plant_years],
            y=[plant_series[y].get(t, [0, 0])[0] for y in plant_years],
            name=t, marker_color=get_type_colors().get(t, "#607d8b"),
        ))
    fig_plant_count.update_layout(
        barmode="stack",
        title="Power Plants — Project Count by Year",
        height=420, xaxis_title="B.S. Year", yaxis_title="Number of projects",
    )
    add_watermark(fig_plant_count)

    # Transmission Lines Growth
    tx_series = loader.yearly_series(tx_lines, key_field="status")
    tx_years = sorted(tx_series.keys())
    all_tx_statuses = sorted({k for y in tx_years for k in tx_series[y].keys()})

    fig_tx_cap = go.Figure()
    for st in all_tx_statuses:
        fig_tx_cap.add_trace(go.Scatter(
            x=[str(y) for y in tx_years],
            y=[tx_series[y].get(st, [0, 0])[1] for y in tx_years],
            mode="lines+markers", name=st,
            line=dict(color=get_status_colors().get(st, "#90a4ae")),
        ))
    # Cumulative total transmission capacity across all statuses
    tx_totals_by_year = [sum(tx_series[y].get(st, [0, 0.0])[1] for st in all_tx_statuses) for y in tx_years]
    tx_cum_capacity = _cumsum(tx_totals_by_year)
    fig_tx_cap.add_trace(go.Scatter(
        x=[str(y) for y in tx_years], y=tx_cum_capacity,
        mode="lines+markers", name="Cumulative Total Capacity",
        line=dict(color="#37474f", width=3, dash="dot"), yaxis="y2",
    ))
    fig_tx_cap.update_layout(
        title="Transmission Lines — Licensed Capacity by Year (B.S.)",
        xaxis_title="B.S. Year", yaxis_title="Capacity (MW)",
        yaxis2=dict(title="Cumulative Total Capacity (MW)", overlaying="y", side="right",
                     showgrid=False, range=[0, max(tx_cum_capacity) * 1.15 if tx_cum_capacity else 1]),
        height=480, legend=_legend_below_xaxis(),
        margin=dict(b=_BOTTOM_MARGIN_FOR_LEGEND),
    )
    fig_tx_cap = _apply_secondary_axis_setting(fig_tx_cap)
    add_watermark(fig_tx_cap)

    fig_tx_count = go.Figure()
    for st in all_tx_statuses:
        fig_tx_count.add_trace(go.Bar(
            x=[str(y) for y in tx_years],
            y=[tx_series[y].get(st, [0, 0])[0] for y in tx_years],
            name=st, marker_color=get_status_colors().get(st, "#90a4ae"),
        ))
    fig_tx_count.update_layout(
        barmode="stack",
        title="Transmission Lines — Project Count by Year",
        height=420, xaxis_title="B.S. Year", yaxis_title="Number of projects",
    )
    add_watermark(fig_tx_count)

    return html.Div([
        html.H4("⚡ Power Plants Growth", className="mt-3 mb-3"),
        dcc.Graph(figure=fig_plant_cap),
        dcc.Graph(figure=fig_plant_count),
        html.Hr(),
        html.H4("🔌 Transmission Lines Growth", className="mt-3 mb-3"),
        dcc.Graph(figure=fig_tx_cap),
        dcc.Graph(figure=fig_tx_count),
    ])


# ── GIS TAB ─────────────────────────────────────────────────────────────────
def render_gis_tab(loader, recs, f_crs, show_boundary=True, show_pa=False):
    gis_loaded = getattr(de.GIS, 'loaded', False)
    plant_recs = [r for r in recs if r["lat"] and r["lon"]]

    if not gis_loaded and not plant_recs:
        map_view = dbc.Alert(
            "No map data available yet — neither the district/province "
            "boundary package nor any licensed-project coordinates have "
            "been loaded. An administrator can add these at /admin (sync "
            "the workbook and the GIS package, or set DEFAULT_SHEET_URL / "
            "DEFAULT_GIS_DRIVE_URL on the server).",
            color="info",
        )
    else:
        html_str = gis_leaflet_map.build_gis_map_html(
            recs, get_status_colors(), get_type_colors(), get_province_colors(),
        )
        iframe = html.Iframe(
            srcDoc=html_str,
            style={"width": "100%", "height": "690px", "border": "none", "borderRadius": "6px"},
        )
        map_view = html.Div([
            iframe,
            html.Div(
                "Filter by license stage, project type, province, or search by name/promoter/"
                "district/local body directly in the map's own sidebar. Toggle district/local-body/"
                "protected-area layers on and off, and switch the emphasized coordinate "
                "system (WGS-84 vs. Everest 1830) from the sidebar as well.",
                className="text-muted small mt-1",
            ),
        ])
        if not gis_loaded:
            map_view = html.Div([
                dbc.Alert(
                    "District/province boundary shading isn't loaded yet — "
                    "showing project locations only. An administrator can add "
                    "the GIS package at /admin.",
                    color="warning", className="mb-2", dismissable=True,
                ),
                map_view,
            ])

    return map_view



# ── COMPARE TAB ────────────────────────────────────────────────────────────
def render_compare(loader, recs):
    plants = [r for r in recs if r["type"] != "Transmission Line"]
    lines = [r for r in recs if r["type"] == "Transmission Line"]

    by_status_mw = defaultdict(float)
    for r in plants:
        by_status_mw[r["status"]] += r["capacity_mw"] or 0
    stages_present = [s for s in STAGE_DISPLAY_ORDER if s in by_status_mw] + \
                     [s for s in by_status_mw if s not in STAGE_DISPLAY_ORDER]
    mw_values = [by_status_mw[s] for s in stages_present]
    colors = [get_status_colors().get(s, "#90a4ae") for s in stages_present]
    cum_mw = _cumsum(mw_values)
    fig_plants = go.Figure()
    fig_plants.add_trace(go.Bar(x=stages_present, y=mw_values, marker_color=colors,
                                 name="Capacity (MW)", width=0.45))
    fig_plants.add_trace(go.Scatter(x=stages_present, y=cum_mw, mode="lines+markers",
                                     name="Cumulative Capacity", yaxis="y2",
                                     line=dict(color="#37474f", width=3, dash="dot")))
    fig_plants.update_layout(
        title="Power Plants — Capacity (MW) by License Stage",
        height=560, bargap=0.45, yaxis_title="MW",
        yaxis2=dict(title="Cumulative Capacity (MW)", overlaying="y", side="right",
                     showgrid=False, range=[0, max(cum_mw) * 1.1 if cum_mw else 1]),
        legend=_legend_below_xaxis(),
        margin=dict(b=_BOTTOM_MARGIN_FOR_LEGEND),
    )
    fig_plants = _apply_secondary_axis_setting(fig_plants)
    add_watermark(fig_plants)

    by_status_km = defaultdict(float)
    for r in lines:
        by_status_km[r["status"]] += r["line_length_km"] or 0
    tx_stages_present = [s for s in STAGE_DISPLAY_ORDER if s in by_status_km] + \
                        [s for s in by_status_km if s not in STAGE_DISPLAY_ORDER]
    km_values = [by_status_km[s] for s in tx_stages_present]
    colors_tx = [get_status_colors().get(s, "#90a4ae") for s in tx_stages_present]
    cum_km = _cumsum(km_values)
    fig_lines = go.Figure()
    fig_lines.add_trace(go.Bar(x=tx_stages_present, y=km_values, marker_color=colors_tx,
                                name="Length (KM)", width=0.45))
    fig_lines.add_trace(go.Scatter(x=tx_stages_present, y=cum_km, mode="lines+markers",
                                    name="Cumulative Length", yaxis="y2",
                                    line=dict(color="#37474f", width=3, dash="dot")))
    fig_lines.update_layout(
        title="Transmission Lines — Length (KM) by License Stage",
        height=560, bargap=0.45, yaxis_title="KM",
        yaxis2=dict(title="Cumulative Length (KM)", overlaying="y", side="right",
                     showgrid=False, range=[0, max(cum_km) * 1.1 if cum_km else 1]),
        legend=_legend_below_xaxis(),
        margin=dict(b=_BOTTOM_MARGIN_FOR_LEGEND),
    )
    fig_lines = _apply_secondary_axis_setting(fig_lines)
    add_watermark(fig_lines)

    by_volt = defaultdict(int)
    for r in lines:
        if r["voltage_kv"]:
            by_volt[r["voltage_kv"]] += 1
    volt_keys = sorted(by_volt)
    volt_vals = [by_volt[v] for v in volt_keys]
    cum_volt_vals = _cumsum([float(c) for c in volt_vals])
    fig_volt = go.Figure()
    fig_volt.add_trace(go.Bar(
        x=[f"{v:.0f} kV" for v in volt_keys], y=volt_vals,
        marker_color="#6a1b9a", width=0.45,
        name="Projects",
    ))
    fig_volt.add_trace(go.Scatter(
        x=[f"{v:.0f} kV" for v in volt_keys], y=cum_volt_vals, mode="lines+markers",
        name="Cumulative Projects", yaxis="y2",
        line=dict(color="#37474f", width=3, dash="dot"),
        marker=dict(size=8),
    ))
    fig_volt.update_layout(
        title="Transmission Lines by Voltage Class", height=480, bargap=0.45,
        yaxis2=dict(title="Cumulative Projects", overlaying="y", side="right",
                     showgrid=False, range=[0, max(cum_volt_vals) * 1.15 if cum_volt_vals else 1]),
        legend=_legend_below_xaxis(),
        margin=dict(b=_BOTTOM_MARGIN_FOR_LEGEND),
    )
    fig_volt = _apply_secondary_axis_setting(fig_volt)
    add_watermark(fig_volt)

    return dbc.Tabs(id="compare-subtabs", active_tab="plants", className="sub-tabs-nav", children=[
        dbc.Tab(dcc.Graph(figure=fig_plants), label="Power Plants", tab_id="plants",
                tab_style={"marginTop": "10px"}),
        dbc.Tab(dcc.Graph(figure=fig_lines), label="Transmission Lines", tab_id="lines",
                tab_style={"marginTop": "10px"}),
        dbc.Tab(dcc.Graph(figure=fig_volt), label="By Voltage Class", tab_id="by-voltage",
                tab_style={"marginTop": "10px"}),
    ])


# ── DATA TABLE ──────────────────────────────────────────────────────────────
def render_table(recs, f_crs=None):
    f_crs = f_crs or ct.CRS_WGS84
    cols = ["project", "type", "status", "capacity_mw", "voltage_kv", "line_length_km",
            "district", "province", "promoter", "lat_disp", "lon_disp", "loc_source"]
    data = []
    for r in recs:
        row = {c: r.get(c) for c in cols if c not in ("lat_disp", "lon_disp")}
        lat, lon = r.get("lat"), r.get("lon")
        if lat is not None and lon is not None:
            if f_crs == ct.CRS_EVEREST:
                lat, lon = ct.wgs84_to_everest(lat, lon)
            row["lat_disp"] = round(lat, 6)
            row["lon_disp"] = round(lon, 6)
        else:
            row["lat_disp"] = row["lon_disp"] = None
        data.append(row)
    label_map = {
        "project": "Project Name", "type": "Type", "status": "License Stage",
        "capacity_mw": "Capacity (MW)", "voltage_kv": "Voltage (kV)",
        "line_length_km": "Line Length (KM)", "district": "District",
        "province": "Province", "promoter": "Promoter",
        "lat_disp": f"Latitude ({ct.CRS_LABELS[f_crs]})",
        "lon_disp": f"Longitude ({ct.CRS_LABELS[f_crs]})",
        "loc_source": "Location Source",
    }
    col_widths = {
        "project": "220px", "type": "110px", "status": "160px",
        "capacity_mw": "100px", "voltage_kv": "90px", "line_length_km": "100px",
        "district": "120px", "province": "100px", "promoter": "180px",
        "lat_disp": "110px", "lon_disp": "110px", "loc_source": "130px",
    }
    # Numeric columns are right-aligned (so digits stack on the ones place —
    # left-aligning numbers of different widths is what produces the
    # "zigzag" look); everything else stays left-aligned. Headers match
    # their own column's alignment instead of all being centered, so the
    # header row lines up with the data underneath it rather than looking
    # like a separate, misaligned row.
    numeric_cols = {"capacity_mw", "voltage_kv", "line_length_km", "lat_disp", "lon_disp"}

    return html.Div([
        dash_table.DataTable(
            id="data-table",
            data=data,
            columns=[{"name": label_map.get(c, c.replace("_", " ").title()), "id": c} for c in cols],
            page_size=10,
            page_action="native",
            sort_action="native",
            filter_action="native",
            style_table={
                "overflowX": "auto",
                "border": "1px solid #dee2e6",
                "borderRadius": "6px",
            },
            style_cell={
                "fontFamily": "Helvetica, Arial, sans-serif",
                "fontSize": "13px",
                "padding": "8px 12px",
                "textAlign": "left",
                "whiteSpace": "nowrap",
                "overflow": "hidden",
                "textOverflow": "ellipsis",
                "border": "1px solid #e9ecef",
                "height": "36px",
            },
            style_header={
                "fontWeight": "bold",
                "backgroundColor": "#f8f9fa",
                "color": "#495057",
                "border": "1px solid #dee2e6",
                "textAlign": "left",
                "padding": "10px 12px",
                "fontSize": "13px",
                "whiteSpace": "normal",
                "height": "auto",
            },
            style_data_conditional=[
                {
                    "if": {"row_index": "odd"},
                    "backgroundColor": "#fafbfc",
                },
                {
                    "if": {"state": "active"},
                    "backgroundColor": "#e7f1ff",
                    "border": "1px solid #b8daff",
                },
            ],
            style_cell_conditional=[
                {"if": {"column_id": c}, "minWidth": w, "width": w, "maxWidth": w,
                 "textAlign": "right" if c in numeric_cols else "left"}
                for c, w in col_widths.items()
            ],
            style_header_conditional=[
                {"if": {"column_id": c}, "textAlign": "right" if c in numeric_cols else "left"}
                for c in col_widths
            ],
        ),
        html.Div([
            html.Label("Show entries:", className="me-2 fw-semibold small"),
            dcc.Dropdown(
                id="table-page-size",
                options=[
                    {"label": "25", "value": 25},
                    {"label": "50", "value": 50},
                    {"label": "100", "value": 100},
                    {"label": "All", "value": len(data) if data else 100},
                ],
                value=None,
                placeholder="10 (default)",
                clearable=False,
                style={"width": "140px", "display": "inline-block"},
            ),
            html.Span(className="ms-auto"),
            html.Span("Download filtered data:", className="me-2 fw-semibold small text-muted"),
            dbc.ButtonGroup([
                dbc.Button("⬇ Excel", id="export-table-xlsx", size="sm", color="success", outline=True),
                dbc.Button("⬇ Word", id="export-table-docx", size="sm", color="primary", outline=True),
                dbc.Button("⬇ PDF", id="export-table-pdf", size="sm", color="danger", outline=True),
            ], size="sm"),
            dcc.Download(id="download-table-export"),
        ], className="mt-2 d-flex align-items-center flex-wrap gap-2"),
        html.Div(
            "Exports respect any column filters/sorting currently applied in the table above "
            "(the native filter row and column-header sort), not just the visible page.",
            className="text-muted small mt-1",
        ),
    ])


@app.callback(
    Output("data-table", "page_size"),
    Input("table-page-size", "value"),
    prevent_initial_call=True,
)
def update_table_page_size(page_size):
    return page_size or 10


# ── DATA TABLE — FILTERED EXPORT (Excel / Word / PDF) ────────────────────────
def _export_df_from_table(rows, columns):
    """Build a labeled DataFrame from the DataTable's *currently filtered/
    sorted* rows (dash_table's `derived_virtual_data`), in the same column
    order/labels shown on screen."""
    if not rows:
        return None
    df = pd.DataFrame(rows)
    col_ids = [c["id"] for c in columns if c["id"] in df.columns]
    df = df[col_ids]
    df = df.rename(columns={c["id"]: c["name"] for c in columns})
    return df


def _df_to_docx_bytes(df, title="License Status Data — Filtered Export"):
    from docx import Document
    from docx.shared import Pt
    import io as _io

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(
        f"Exported {datetime.now().strftime('%Y-%m-%d %H:%M')} · {len(df)} row(s) · "
        f"Nepal Power Plants Dashboard (nea.org.np data) — www.neupanesandeep.com.np"
    ).runs[0].italic = True

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            val = row[col]
            cells[i].text = "" if pd.isna(val) else str(val)
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(8)

    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _df_to_pdf_bytes(df, title="License Status Data — Filtered Export"):
    """Paginated table PDF via matplotlib (same PdfPages already used by
    the main 'Generate Report' feature, so no new heavy PDF dependency)."""
    import io as _io
    import matplotlib.pyplot as plt

    rows_per_page = 28
    n_pages = max(1, math.ceil(len(df) / rows_per_page))
    buf = _io.BytesIO()
    with PdfPages(buf) as pdf:
        for page in range(n_pages):
            chunk = df.iloc[page * rows_per_page:(page + 1) * rows_per_page]
            fig_w = max(11, 1.1 * len(df.columns))
            fig, ax = plt.subplots(figsize=(fig_w, 8.5))
            ax.axis("off")
            fig.suptitle(title, fontsize=13, fontweight="bold", y=0.98)
            fig.text(0.01, 0.955,
                      f"Exported {datetime.now().strftime('%Y-%m-%d %H:%M')} · "
                      f"{len(df)} row(s) total · Page {page + 1} of {n_pages}",
                      fontsize=8, color="#555")
            tbl = ax.table(
                cellText=chunk.astype(str).values,
                colLabels=list(df.columns),
                loc="upper center", cellLoc="left",
            )
            tbl.auto_set_font_size(False)
            tbl.set_fontsize(6.5)
            tbl.scale(1, 1.3)
            for (r, c), cell in tbl.get_celld().items():
                if r == 0:
                    cell.set_facecolor("#1565c0")
                    cell.set_text_props(color="white", fontweight="bold")
                elif r % 2 == 0:
                    cell.set_facecolor("#f5f7fa")
            fig.text(0.99, 0.02, "Er. Sandeep Neupane — www.neupanesandeep.com.np",
                      fontsize=6.5, color="#999", ha="right")
            pdf.savefig(fig, bbox_inches="tight")
            plt.close(fig)
    return buf.getvalue()


@app.callback(
    Output("download-table-export", "data"),
    Input("export-table-xlsx", "n_clicks"),
    Input("export-table-docx", "n_clicks"),
    Input("export-table-pdf", "n_clicks"),
    State("data-table", "derived_virtual_data"),
    State("data-table", "columns"),
    prevent_initial_call=True,
)
def export_filtered_table(n_xlsx, n_docx, n_pdf, rows, columns):
    triggered = ctx.triggered_id
    df = _export_df_from_table(rows, columns)
    if df is None or df.empty:
        raise dash.exceptions.PreventUpdate

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    if triggered == "export-table-xlsx":
        return dcc.send_data_frame(
            df.to_excel, f"nepal_power_data_{ts}.xlsx", index=False, sheet_name="Filtered Data"
        )
    if triggered == "export-table-docx":
        return dcc.send_bytes(lambda b: b.write(_df_to_docx_bytes(df)), f"nepal_power_data_{ts}.docx")
    if triggered == "export-table-pdf":
        return dcc.send_bytes(lambda b: b.write(_df_to_pdf_bytes(df)), f"nepal_power_data_{ts}.pdf")
    raise dash.exceptions.PreventUpdate


# ── CUSTOM STYLE TAB ─────────────────────────────────────────────────────────
def render_custom_tab():
    """REQ 12: Custom tab to change chart colors, styles, and types."""
    return html.Div([
        html.H4("🎨 Custom Chart Styling", className="mb-4"),
        dbc.Row([
            dbc.Col(md=6, children=[
                dbc.Card([
                    dbc.CardBody([
                        html.H5("Chart Appearance", className="card-title mb-3"),

                        html.Label("Color Scheme", className="fw-semibold"),
                        dcc.Dropdown(
                            id="custom-color-scheme",
                            options=[
                                {"label": "Default", "value": "default"},
                                {"label": "Pastel", "value": "pastel"},
                                {"label": "Dark", "value": "dark"},
                                {"label": "Vibrant", "value": "vibrant"},
                            ],
                            value=CHART_STYLE_STATE["color_scheme"],
                            clearable=False,
                            className="mb-3",
                        ),

                        html.Label("Bar Mode", className="fw-semibold"),
                        dcc.Dropdown(
                            id="custom-bar-mode",
                            options=[
                                {"label": "Grouped", "value": "group"},
                                {"label": "Stacked", "value": "stack"},
                                {"label": "Relative", "value": "relative"},
                            ],
                            value=CHART_STYLE_STATE["bar_mode"],
                            clearable=False,
                            className="mb-3",
                        ),

                        html.Label("Chart Type", className="fw-semibold"),
                        dcc.Dropdown(
                            id="custom-chart-type",
                            options=[
                                {"label": "Bar Chart", "value": "bar"},
                                {"label": "Line Chart", "value": "line"},
                                {"label": "Area Chart", "value": "area"},
                                {"label": "Scatter Plot", "value": "scatter"},
                            ],
                            value=CHART_STYLE_STATE["chart_type"],
                            clearable=False,
                            className="mb-3",
                        ),

                        html.Label("Font Family", className="fw-semibold"),
                        dcc.Dropdown(
                            id="custom-font-family",
                            options=[
                                {"label": "Arial", "value": "Arial"},
                                {"label": "Helvetica", "value": "Helvetica"},
                                {"label": "Georgia", "value": "Georgia"},
                                {"label": "Times New Roman", "value": "Times New Roman"},
                                {"label": "Courier New", "value": "Courier New"},
                                {"label": "Verdana", "value": "Verdana"},
                            ],
                            value=CHART_STYLE_STATE["font_family"],
                            clearable=False,
                            className="mb-3",
                        ),
                    ]),
                ], className="mb-3 shadow-sm"),
            ]),
            dbc.Col(md=6, children=[
                dbc.Card([
                    dbc.CardBody([
                        html.H5("Advanced Options", className="card-title mb-3"),

                        html.Label("Title Font Size", className="fw-semibold"),
                        dcc.Slider(
                            id="custom-title-size",
                            min=12, max=24, step=1,
                            value=CHART_STYLE_STATE["title_size"],
                            marks={12: "12", 16: "16", 20: "20", 24: "24"},
                            className="mb-4",
                        ),

                        html.Label("Label Font Size", className="fw-semibold"),
                        dcc.Slider(
                            id="custom-label-size",
                            min=10, max=18, step=1,
                            value=CHART_STYLE_STATE["label_size"],
                            marks={10: "10", 12: "12", 14: "14", 16: "16", 18: "18"},
                            className="mb-4",
                        ),

                        dbc.Checklist(
                            id="custom-show-grid",
                            options=[{"label": " Show Grid Lines", "value": "show"}],
                            value=["show"] if CHART_STYLE_STATE["show_grid"] else [],
                            switch=True,
                            className="mb-3",
                        ),

                        dbc.Checklist(
                            id="custom-animation",
                            options=[{"label": " Enable Chart Animations", "value": "animate"}],
                            value=["animate"] if CHART_STYLE_STATE["animation"] else [],
                            switch=True,
                            className="mb-3",
                        ),

                        dbc.Checklist(
                            id="custom-secondary-axis",
                            options=[{"label": " Show Secondary (Cumulative) Axis", "value": "show"}],
                            value=["show"] if CHART_STYLE_STATE["secondary_axis"] else [],
                            switch=True,
                            className="mb-3",
                        ),
                        html.Div(
                            "Turns the cumulative-total line and its right-hand axis "
                            "on or off across every chart in the app.",
                            className="text-muted small mb-3",
                            style={"marginTop": "-8px"},
                        ),

                        html.Hr(),
                        html.Div([
                            html.Strong("Current Settings Preview:"),
                            html.Div(id="custom-style-preview", className="mt-2 small text-muted"),
                        ]),
                    ]),
                ], className="mb-3 shadow-sm"),
            ]),
        ]),
        dbc.Row([
            dbc.Col(md=12, children=[
                dbc.Button("Apply Changes", id="btn-apply-style", color="primary", className="me-2"),
                dbc.Button("Reset to Default", id="btn-reset-style", color="secondary", outline=True),
                html.Div(id="custom-style-feedback", className="mt-2"),
            ]),
        ]),
    ])


@app.callback(
    Output("chart-style-store", "data"),
    Output("custom-style-feedback", "children"),
    Input("btn-apply-style", "n_clicks"),
    State("custom-color-scheme", "value"),
    State("custom-bar-mode", "value"),
    State("custom-chart-type", "value"),
    State("custom-font-family", "value"),
    State("custom-title-size", "value"),
    State("custom-label-size", "value"),
    State("custom-show-grid", "value"),
    State("custom-animation", "value"),
    State("custom-secondary-axis", "value"),
    prevent_initial_call=True,
)
def apply_custom_style(n_clicks, color_scheme, bar_mode, chart_type, font_family,
                        title_size, label_size, show_grid, animation, secondary_axis):
    if not n_clicks:
        return dash.no_update, dash.no_update

    new_style = {
        "color_scheme": color_scheme or "default",
        "bar_mode": bar_mode or "group",
        "chart_type": chart_type or "bar",
        "font_family": font_family or "Arial",
        "title_size": title_size or 16,
        "label_size": label_size or 12,
        "show_grid": bool(show_grid),
        "animation": bool(animation),
        "secondary_axis": bool(secondary_axis),
    }
    global CHART_STYLE_STATE
    CHART_STYLE_STATE.update(new_style)

    feedback = dbc.Alert("✅ Style settings applied! Refresh the page or switch tabs to see changes.",
                          color="success", dismissable=True)
    return new_style, feedback


@app.callback(
    Output("custom-color-scheme", "value"),
    Output("custom-bar-mode", "value"),
    Output("custom-chart-type", "value"),
    Output("custom-font-family", "value"),
    Output("custom-title-size", "value"),
    Output("custom-label-size", "value"),
    Output("custom-show-grid", "value"),
    Output("custom-animation", "value"),
    Output("custom-secondary-axis", "value"),
    Output("custom-style-feedback", "children", allow_duplicate=True),
    Input("btn-reset-style", "n_clicks"),
    prevent_initial_call=True,
)
def reset_custom_style(n_clicks):
    if not n_clicks:
        return [dash.no_update] * 10

    global CHART_STYLE_STATE
    CHART_STYLE_STATE = {
        "bar_mode": "group",
        "chart_type": "bar",
        "color_scheme": "default",
        "show_grid": True,
        "font_family": "Arial",
        "title_size": 16,
        "label_size": 12,
        "animation": True,
        "secondary_axis": True,
    }

    feedback = dbc.Alert("✅ Style settings reset to default!", color="info", dismissable=True)
    return (
        "default", "group", "bar", "Arial", 16, 12, ["show"], ["animate"], ["show"], feedback
    )
# ── REPORTS TAB ──────────────────────────────────────────────────────────────
# Section checkboxes on the Reports tab map straight to these values, which
# download_pdf() below checks before adding each section's charts — so
# picking a subset of sections here is what actually shortens/lengthens the
# generated PDF, not just a cosmetic list.
REPORT_SECTION_OPTIONS = [
    {"label": " Power Plants & Transmission Lines", "value": "plants_tx"},
    {"label": " GoN Studied Projects", "value": "gon_study"},
    {"label": " License Cancelled Projects", "value": "cancelled"},
    {"label": " Growth Trends (by year)", "value": "growth"},
    {"label": " NEA Operational Data (generation, consumers, revenue, forecasts)", "value": "nea"},
]
REPORT_SECTION_DEFAULT = [o["value"] for o in REPORT_SECTION_OPTIONS]


def render_reports_tab():
    """The 'Generate Report' main tab. Previously this was a small button
    tucked into the filter sidebar — which meant it only ever showed up on
    the License Status / License Insights tabs, since the sidebar itself is
    hidden everywhere else. Giving it a proper tab makes it reachable from
    anywhere, gives the report some breathing room to explain what it
    contains, and gives report-generation a real "please wait" status
    instead of a small spinner dot next to the button."""
    return html.Div([
        html.H4("📄 Generate a Report", className="mb-2"),
        html.P(
            "Builds a condensed, presentation-ready PDF — a cover page followed by "
            "roughly a dozen pages of charts, 2–4 to a page — covering whichever "
            "sections you select below.",
            className="text-muted mb-4",
        ),
        dbc.Row([
            dbc.Col(md=7, children=[
                dbc.Card(className="shadow-sm mb-3", children=dbc.CardBody([
                    html.H5("Sections to include", className="card-title mb-3"),
                    dbc.Checklist(
                        id="report-sections",
                        options=REPORT_SECTION_OPTIONS,
                        value=REPORT_SECTION_DEFAULT,
                        className="mb-1",
                    ),
                    html.Div(
                        "Power Plant / Transmission sections respect whatever filters "
                        "are currently set on the License Status and License Insights tabs "
                        "(type, stage, province, dates, search, etc.) — leave those tabs' "
                        "filters at their defaults for a full, unfiltered report.",
                        className="text-muted small mt-2",
                    ),
                ])),
            ]),
            dbc.Col(md=5, children=[
                dbc.Card(className="shadow-sm", children=dbc.CardBody([
                    html.H5("Generate", className="card-title mb-3"),
                    dcc.Loading(
                        id="report-loading", type="circle", color="#c62828",
                        children=[
                            dbc.Button(
                                [html.I(className="bi bi-file-earmark-pdf-fill me-2"),
                                 "Generate PDF Report"],
                                id="btn-pdf", color="danger", size="lg", className="w-100",
                            ),
                            dcc.Download(id="download-pdf"),
                        ],
                    ),
                    html.Div(id="report-status", className="mt-3 small"),
                ])),
            ]),
        ]),
    ])


# Fires the instant "please wait" status the moment the button is clicked —
# client-side, so it appears immediately rather than after the (potentially
# many-second) server-side PDF build finishes. download_pdf() below overwrites
# this same report-status text with the final result once it's done.
app.clientside_callback(
    """
    function(n_clicks) {
        if (!n_clicks) { return window.dash_clientside.no_update; }
        return "⏳ Your report is being generated — this can take up to a minute "
             + "depending on how much data is included. Please don't close this tab.";
    }
    """,
    Output("report-status", "children", allow_duplicate=True),
    Input("btn-pdf", "n_clicks"),
    prevent_initial_call=True,
)


# ── PDF REPORT: PAGE-BUILDING HELPERS ───────────────────────────────────────
_PDF_FOOTER = "Er. Sandeep Neupane  |  www.neupanesandeep.com.np"


def _pdf_page(pdf, panels):
    """panels: list of (draw_fn, kwargs) — each draw_fn(ax, **kwargs) draws
    ONE chart into the given axis. Packs 1-4 panels onto a single
    portrait page (stacked vertically for 1-3, a 2x2 grid for 4) with
    ONE shared footer/watermark per page instead of one full page per
    chart — this is what keeps the whole report to roughly 10 pages
    instead of 25+."""
    n = len(panels)
    if n == 0:
        return
    fig = plt.figure(figsize=(11.69, 8.27))
    if n == 4:
        gs = fig.add_gridspec(2, 2, hspace=0.5, wspace=0.32,
                               top=0.94, bottom=0.08, left=0.08, right=0.96)
        axes = [fig.add_subplot(gs[i // 2, i % 2]) for i in range(4)]
    else:
        gs = fig.add_gridspec(n, 1, hspace=0.65,
                               top=0.94, bottom=0.08, left=0.10, right=0.92)
        axes = [fig.add_subplot(gs[i, 0]) for i in range(n)]
    for ax, (draw_fn, kwargs) in zip(axes, panels):
        draw_fn(ax, **kwargs)
    fig.text(0.98, 0.015, _PDF_FOOTER, fontsize=7.5, color="gray",
              ha="right", va="bottom", alpha=0.6)
    pdf.savefig(fig)
    plt.close(fig)


def _pdf_render_batches(pdf, chart_specs, per_page=3):
    """chart_specs: list of (draw_fn, kwargs). Groups them `per_page` at
    a time and renders each group as one page via _pdf_page()."""
    for i in range(0, len(chart_specs), per_page):
        _pdf_page(pdf, chart_specs[i:i + per_page])


def _draw_bar_cum(ax, title, categories, values, colors, y_label, cum_label, value_fmt="{:,.1f}"):
    if not categories:
        ax.axis("off")
        return
    bars = ax.bar(categories, values, color=colors)
    ax.set_title(title, fontsize=9.5, fontweight="bold", pad=6)
    ax.set_ylabel(y_label, fontsize=8)
    ax.tick_params(axis="both", labelsize=7)
    for b, v in zip(bars, values):
        if v:
            ax.text(b.get_x() + b.get_width() / 2, v, value_fmt.format(v),
                    ha="center", va="bottom", fontsize=6)
    cum = _cumsum(values)
    ax2 = ax.twinx()
    ax2.plot(categories, cum, color="#37474f", marker="o", markersize=3,
              linestyle="--", linewidth=1.5, label=cum_label)
    ax2.set_ylabel(cum_label, fontsize=7)
    ax2.tick_params(axis="y", labelsize=6.5)
    plt.setp(ax.get_xticklabels(), rotation=30, ha="right", fontsize=6.5)
    h1, l1 = ax.get_legend_handles_labels()
    h2, l2 = ax2.get_legend_handles_labels()
    if h1 or h2:
        ax.legend(h1 + h2, l1 + l2, loc="upper left", fontsize=6, framealpha=0.9)


def _draw_pie(ax, title, labels, values, colors):
    if not values:
        ax.axis("off")
        return
    ax.pie(values, labels=labels, autopct="%1.0f%%", colors=colors, textprops={"fontsize": 6.5})
    ax.set_title(title, fontsize=9.5, fontweight="bold", pad=6)


def _draw_multiline_cum(ax, title, x_labels, series, colors, y_label,
                          cum_values=None, cum_label=None, x_label="B.S. Year"):
    if not x_labels:
        ax.axis("off")
        return
    for name, yvals in series.items():
        ax.plot(x_labels, yvals, marker="o", markersize=3, linewidth=1.3,
                 label=name, color=colors.get(name, "#607d8b"))
    ax.set_title(title, fontsize=9.5, fontweight="bold", pad=6)
    ax.set_ylabel(y_label, fontsize=8)
    ax.set_xlabel(x_label, fontsize=7.5)
    ax.tick_params(axis="both", labelsize=6.5)
    plt.setp(ax.get_xticklabels(), rotation=40, ha="right", fontsize=6)
    if cum_values is not None:
        ax2 = ax.twinx()
        ax2.plot(x_labels, cum_values, color="#2e7d32", linestyle="--",
                  linewidth=1.8, marker="o", markersize=3, label=cum_label)
        ax2.set_ylabel(cum_label, fontsize=7)
        ax2.tick_params(axis="y", labelsize=6.5)
    h1, l1 = ax.get_legend_handles_labels()
    if h1:
        ax.legend(h1, l1, loc="upper left", fontsize=5.5, framealpha=0.9, ncol=2)


def _draw_stacked_bar(ax, title, x_labels, series, colors, y_label, x_label="B.S. Year"):
    if not x_labels:
        ax.axis("off")
        return
    bottom = [0.0] * len(x_labels)
    for name, yvals in series.items():
        ax.bar(x_labels, yvals, bottom=bottom, label=name, color=colors.get(name, "#607d8b"))
        bottom = [b + v for b, v in zip(bottom, yvals)]
    ax.set_title(title, fontsize=9.5, fontweight="bold", pad=6)
    ax.set_ylabel(y_label, fontsize=8)
    ax.set_xlabel(x_label, fontsize=7.5)
    ax.tick_params(axis="both", labelsize=6.5)
    plt.setp(ax.get_xticklabels(), rotation=40, ha="right", fontsize=6)
    ax.legend(loc="upper left", fontsize=5.5, framealpha=0.9, ncol=2)


def _pdf_cover_page(pdf, recs, source_label, filter_summary):
    import datetime

    fig = plt.figure(figsize=(11.69, 8.27))
    y = 0.86  # every line below is placed explicitly, top-anchored, so nothing
              # can drift into the line above/below it regardless of length

    fig.text(0.5, y, "Nepal Power Plant & Transmission Line",
             ha="center", va="top", fontsize=22, fontweight="bold")
    y -= 0.07
    fig.text(0.5, y, "License Status Dashboard — Full Report",
             ha="center", va="top", fontsize=16)
    y -= 0.09

    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    fig.text(0.5, y, f"Generated: {now}", ha="center", va="top", fontsize=11, color="#555")
    y -= 0.045

    # Report shows the authoritative public data sources (DoED / NEA),
    # matching the dashboard's own header/footer — not the internal
    # sync mechanism.
    fig.text(0.5, y, "Data sources: www.doed.gov.np  |  www.nea.org.np", ha="center", va="top",
              fontsize=11, color="#555", url="https://www.doed.gov.np")
    y -= 0.045

    wrapped_filters = textwrap.fill(f"Filters applied: {filter_summary}", width=95)
    n_filter_lines = wrapped_filters.count("\n") + 1
    fig.text(0.5, y, wrapped_filters, ha="center", va="top", fontsize=10, color="#555")
    y -= 0.032 * n_filter_lines + 0.035

    plants = [r for r in recs if r["type"] != "Transmission Line" and r["status"] not in de.EXTRA_STATUS_ORDER]
    tx = [r for r in recs if r["type"] == "Transmission Line" and r["status"] not in de.EXTRA_STATUS_ORDER]
    op = [r for r in plants if r["status"] == "Operating"]
    kpi_lines = [
        f"Installed Capacity: {sum(r['capacity_mw'] or 0 for r in op):,.1f} MW ({len(op):,} Operating Plants)",
        f"Active Power Plants: {len(plants):,} Projects ({sum(r['capacity_mw'] or 0 for r in plants):,.1f} MW total)",
        f"Transmission Lines: {len(tx):,} Projects ({sum(r['line_length_km'] or 0 for r in tx):,.1f} KM)",
        f"GoN Studied Projects: {sum(1 for r in recs if r['status'] == 'GoN Study Project'):,}",
        f"License Cancelled: {sum(1 for r in recs if r['status'] == 'Cancelled'):,}",
    ]
    for line in kpi_lines:
        fig.text(0.5, y, line, ha="center", va="top", fontsize=11)
        y -= 0.045

    fig.text(0.98, 0.02, _PDF_FOOTER, fontsize=8, color="gray",
              ha="right", va="bottom", alpha=0.5)
    pdf.savefig(fig)
    plt.close(fig)

def _nan_safe(vals):
    """None -> NaN so a gap in the sheet becomes a break in the line
    (matplotlib handles NaN natively) instead of crashing ax.plot()."""
    return [(float(v) if v is not None else float("nan")) for v in (vals or [])]


def _nea_report_chart_specs():
    """NEA Operational Data section for the PDF report — same
    draw_fn/kwargs shape as the license chart_specs above, so it's
    rendered through the identical 3-per-page layout. Each chart is
    skipped entirely (not left as a blank panel) if its underlying
    series isn't available yet. The whole function is one big
    try/except so a malformed/partial NEA sync can never take down
    the rest of the report — worst case, the NEA section is just
    shorter (or absent) that run."""
    specs = []
    try:
        data = NEA.get_dashboard_data() or {}
        if not data:
            return specs

        ae = data.get("annualEnergy") or {}
        sl = data.get("systemLoss") or {}
        cg = data.get("consumers") or {}
        sr = data.get("sales") or {}
        fin = data.get("financial") or {}

        years_ae = [str(y) for y in (ae.get("years") or [])]
        nat_peak, sys_peak = ae.get("national_peak") or [], ae.get("system_peak") or []
        if years_ae and (nat_peak or sys_peak):
            series = {}
            if nat_peak:
                series["National Peak (MW)"] = _nan_safe(nat_peak)
            if sys_peak:
                series["System Peak (MW)"] = _nan_safe(sys_peak)
            specs.append((_draw_multiline_cum, dict(
                title="NEA — National vs System Peak Demand (MW)",
                x_labels=years_ae, series=series,
                colors={"National Peak (MW)": "#1565c0", "System Peak (MW)": "#c62828"},
                y_label="MW",
            )))

        years_sl = [str(y) for y in (sl.get("years") or [])]
        if years_sl and sl.get("system"):
            series = {}
            if sl.get("transmission"):
                series["Transmission Loss"] = _nan_safe(sl["transmission"])
            if sl.get("distribution"):
                series["Distribution Loss"] = _nan_safe(sl["distribution"])
            series["System Loss"] = _nan_safe(sl["system"])
            specs.append((_draw_multiline_cum, dict(
                title="NEA — System Loss Trend (%)",
                x_labels=years_sl, series=series,
                colors={"Transmission Loss": "#ff8f00", "Distribution Loss": "#8e24aa", "System Loss": "#c62828"},
                y_label="Loss (%)",
            )))

        years_fin = [str(y) for y in (fin.get("years") or [])]
        imp, exp = fin.get("import_mu") or [], fin.get("export_mu") or []
        if years_fin and (imp or exp):
            series = {}
            if imp:
                series["Import (MU)"] = _nan_safe(imp)
            if exp:
                series["Export (MU)"] = _nan_safe(exp)
            specs.append((_draw_multiline_cum, dict(
                title="NEA — Energy Import vs Export (MU)",
                x_labels=years_fin, series=series,
                colors={"Import (MU)": "#6a1b9a", "Export (MU)": "#2e7d32"},
                y_label="MU",
            )))

        years_cg = [str(y) for y in (cg.get("years") or [])]
        totals_cg = cg.get("total") or []
        pairs_cg = [(y, v) for y, v in zip(years_cg, totals_cg) if v is not None]
        if pairs_cg:
            cats, vals = zip(*pairs_cg)
            specs.append((_draw_bar_cum, dict(
                title="NEA — Total Consumers Growth",
                categories=list(cats), values=list(vals),
                colors=["#1565c0"] * len(cats),
                y_label="Consumers", cum_label="Cumulative Consumers", value_fmt="{:,.0f}",
            )))

        years_sr = [str(y) for y in (sr.get("years") or [])]
        profit_loss = fin.get("profit_loss") or []
        if years_sr and sr.get("total"):
            series = {"Total Revenue (Rs. Mn)": _nan_safe(sr["total"])}
            if profit_loss and len(profit_loss) == len(years_sr):
                series["Profit / Loss (Rs. Mn)"] = _nan_safe(profit_loss)
            specs.append((_draw_multiline_cum, dict(
                title="NEA — Revenue & Profit/Loss Trend (Rs. Million)",
                x_labels=years_sr, series=series,
                colors={"Total Revenue (Rs. Mn)": "#2e7d32", "Profit / Loss (Rs. Mn)": "#c62828"},
                y_label="Rs. Million",
            )))

        years_gen = [str(y) for y in (ae.get("years") or [])]
        gen_series = {}
        for key, label in [("nea_own", "NEA Own"), ("nea_sub", "NEA Subsidiary"),
                            ("ipp", "IPP"), ("india", "India Import")]:
            vals = ae.get(key)
            if vals and any(v is not None for v in vals):
                gen_series[label] = [v or 0 for v in vals]
        if years_gen and gen_series:
            specs.append((_draw_stacked_bar, dict(
                title="NEA — Generation Mix, Annual (MU)",
                x_labels=years_gen, series=gen_series,
                colors={"NEA Own": "#1565c0", "NEA Subsidiary": "#00897b",
                         "IPP": "#f9a825", "India Import": "#8e24aa"},
                y_label="MU",
            )))
    except Exception:
        # Any malformed/partial NEA data degrades to "NEA section
        # shorter than usual" rather than breaking the whole report.
        pass

    return specs


# ── PDF REPORT ─────────────────────────────────────────────────────────────
@app.callback(
    Output("download-pdf", "data"),
    Output("report-status", "children", allow_duplicate=True),
    Input("btn-pdf", "n_clicks"),
    State("f-type", "value"), State("f-status", "value"), State("f-province", "value"),
    State("f-capacity", "value"), State("f-tx-length", "value"), State("f-year", "data"),
    State("f-search", "value"),
    State("f-date-from", "value"), State("f-date-to", "value"),
    State("f-cod-from", "value"), State("f-cod-to", "value"),
    State("f-district", "value"), State("f-local", "value"),
    State("report-sections", "value"),
    prevent_initial_call=True,
)
def download_pdf(n_clicks, f_type, f_status, f_province, f_capacity, f_tx_length, f_year, f_search,
                  f_date_from, f_date_to, f_cod_from, f_cod_to, f_district, f_local, sections):
    # Guard against a first click before the checklist has hydrated a value.
    sections = sections if sections is not None else REPORT_SECTION_DEFAULT

    loader = STATE["loader"]
    if loader is None or not loader.records:
        return None, dbc.Alert(
            "⚠️ Project data isn't loaded yet, so there's nothing to put in a report "
            "right now. Please try again in a moment.", color="warning", dismissable=True)
    recs = get_filtered_records(f_type, f_status, f_province, f_capacity, f_year, f_search,
                                 f_date_from, f_date_to, f_cod_from, f_cod_to, f_tx_length,
                                 f_district, f_local)
    if not recs:
        return None, dbc.Alert(
            "⚠️ No projects match the filters currently set on License Status / "
            "License Insights, so the report would be empty. Clear or adjust those "
            "filters and try again.", color="warning", dismissable=True)
    active_recs = [r for r in recs if r["status"] not in de.EXTRA_STATUS_ORDER]
    plant_recs = [r for r in active_recs if r["type"] != "Transmission Line"]
    tx_recs = [r for r in active_recs if r["type"] == "Transmission Line"]

    filter_bits = []
    if f_type: filter_bits.append(f"Type: {', '.join(f_type)}")
    if f_status: filter_bits.append(f"Stage: {', '.join(f_status)}")
    if f_province: filter_bits.append(f"Province: {', '.join(f_province)}")
    if f_district: filter_bits.append(f"District: {', '.join(f_district)}")
    if f_search: filter_bits.append(f'Search: "{f_search}"')
    filter_summary = "; ".join(filter_bits) if filter_bits else "None (showing all records)"

    path = os.path.join(tempfile.gettempdir(), "license_status_report.pdf")
    chart_specs = []  # list of (draw_fn, kwargs) — rendered 3-per-page at the end

    with PdfPages(path) as pdf:
        _pdf_cover_page(pdf, recs, STATE.get("source_label", "—"), filter_summary)

        # ── Power Plants ─────────────────────────────────────────────
        if "plants_tx" in sections:
            by_type = defaultdict(float)
            for r in plant_recs:
                by_type[r["type"]] += r["capacity_mw"] or 0
            types_ = list(by_type.keys())
            chart_specs.append((_draw_bar_cum, dict(
                title="Power Plants — Capacity by Type",
                categories=types_, values=[by_type[t] for t in types_],
                colors=[get_type_colors().get(t, "#607d8b") for t in types_],
                y_label="Capacity (MW)", cum_label="Cumulative Capacity (MW)",
            )))

            by_status_count = defaultdict(int)
            for r in recs:
                by_status_count[r["status"]] += 1
            chart_specs.append((_draw_pie, dict(
                title="License Stage Breakdown — All Records",
                labels=list(by_status_count.keys()), values=list(by_status_count.values()),
                colors=[get_status_colors().get(s, "#90a4ae") for s in by_status_count],
            )))

            stage_totals_p, _ = compute_breakdown(plant_recs, "status")
            stages_p = [s for s in STAGE_DISPLAY_ORDER if s in stage_totals_p]
            chart_specs.append((_draw_bar_cum, dict(
                title="Power Plants — Capacity (MW) by License Stage",
                categories=stages_p, values=[stage_totals_p[s][1] for s in stages_p],
                colors=[get_status_colors().get(s, "#90a4ae") for s in stages_p],
                y_label="Capacity (MW)", cum_label="Cumulative Capacity (MW)",
            )))

            prov_totals_p, _ = compute_breakdown(plant_recs, "province")
            provs_p = [p for p in PROVINCE_DISPLAY_ORDER if p in prov_totals_p] + \
                      [p for p in prov_totals_p if p not in PROVINCE_DISPLAY_ORDER]
            chart_specs.append((_draw_bar_cum, dict(
                title="Power Plants — Capacity (MW) by Province",
                categories=provs_p, values=[prov_totals_p[p][1] for p in provs_p],
                colors=[get_province_colors().get(p, "#455a64") for p in provs_p],
                y_label="Capacity (MW)", cum_label="Cumulative Capacity (MW)",
            )))

            # ── Transmission Lines ───────────────────────────────────
            tx_stage_km = defaultdict(float)
            for r in tx_recs:
                tx_stage_km[r["status"]] += r["line_length_km"] or 0
            stages_tx = [s for s in STAGE_DISPLAY_ORDER if s in tx_stage_km]
            chart_specs.append((_draw_bar_cum, dict(
                title="Transmission Lines — Length (KM) by License Stage",
                categories=stages_tx, values=[tx_stage_km[s] for s in stages_tx],
                colors=[get_status_colors().get(s, "#90a4ae") for s in stages_tx],
                y_label="Length (KM)", cum_label="Cumulative Length (KM)",
            )))

            by_volt = defaultdict(float)
            by_volt_count = defaultdict(int)
            for r in tx_recs:
                if r["voltage_kv"]:
                    by_volt[r["voltage_kv"]] += r["line_length_km"] or 0
                    by_volt_count[r["voltage_kv"]] += 1
            volts = sorted(by_volt)
            chart_specs.append((_draw_bar_cum, dict(
                title="Transmission Lines — Length (KM) by Voltage Class",
                categories=[f"{v:.0f} kV" for v in volts], values=[by_volt[v] for v in volts],
                colors=["#6a1b9a"] * len(volts), y_label="Length (KM)", cum_label="Cumulative Length (KM)",
            )))
            volts_cmp = sorted(by_volt_count)
            chart_specs.append((_draw_bar_cum, dict(
                title="Transmission Lines — Project Count by Voltage Class",
                categories=[f"{v:.0f} kV" for v in volts_cmp], values=[by_volt_count[v] for v in volts_cmp],
                colors=["#6a1b9a"] * len(volts_cmp), y_label="Number of Projects",
                cum_label="Cumulative Projects", value_fmt="{:,.0f}",
            )))

        # ── GoN Studied Projects ─────────────────────────────────────
        gon_recs = [r for r in recs if r["status"] == "GoN Study Project"] if "gon_study" in sections else []
        if gon_recs:
            by_type_g, _ = compute_breakdown(gon_recs, "type")
            types_g = [t for t in de.TYPE_ORDER if t in by_type_g] + \
                      [t for t in by_type_g if t not in de.TYPE_ORDER]
            chart_specs.append((_draw_bar_cum, dict(
                title="GoN Studied Projects — Count by Project Type",
                categories=types_g, values=[by_type_g[t][0] for t in types_g],
                colors=[get_type_colors().get(t, "#607d8b") for t in types_g],
                y_label="Number of Projects", cum_label="Cumulative Projects", value_fmt="{:,.0f}",
            )))
            by_prov_g, _ = compute_breakdown(gon_recs, "province")
            provs_g = [p for p in PROVINCE_DISPLAY_ORDER if p in by_prov_g] + \
                      [p for p in by_prov_g if p not in PROVINCE_DISPLAY_ORDER]
            chart_specs.append((_draw_bar_cum, dict(
                title="GoN Studied Projects — Count by Province",
                categories=provs_g, values=[by_prov_g[p][0] for p in provs_g],
                colors=[get_province_colors().get(p, "#455a64") for p in provs_g],
                y_label="Number of Projects", cum_label="Cumulative Projects", value_fmt="{:,.0f}",
            )))

        # ── License Cancelled ────────────────────────────────────────
        canc_recs = [r for r in recs if r["status"] == "Cancelled"] if "cancelled" in sections else []
        if canc_recs:
            by_type_c, _ = compute_breakdown(canc_recs, "type")
            types_c = [t for t in de.TYPE_ORDER if t in by_type_c] + \
                      [t for t in by_type_c if t not in de.TYPE_ORDER]
            chart_specs.append((_draw_bar_cum, dict(
                title="License Cancelled — Count by Project Type",
                categories=types_c, values=[by_type_c[t][0] for t in types_c],
                colors=[get_type_colors().get(t, "#607d8b") for t in types_c],
                y_label="Number of Projects", cum_label="Cumulative Projects", value_fmt="{:,.0f}",
            )))
            by_prov_c, _ = compute_breakdown(canc_recs, "province")
            provs_c = [p for p in PROVINCE_DISPLAY_ORDER if p in by_prov_c] + \
                      [p for p in by_prov_c if p not in PROVINCE_DISPLAY_ORDER]
            chart_specs.append((_draw_bar_cum, dict(
                title="License Cancelled — Count by Province",
                categories=provs_c, values=[by_prov_c[p][0] for p in provs_c],
                colors=[get_province_colors().get(p, "#455a64") for p in provs_c],
                y_label="Number of Projects", cum_label="Cumulative Projects", value_fmt="{:,.0f}",
            )))

        # ── Growth Trends ────────────────────────────────────────────
        plant_series = loader.yearly_series(plant_recs, key_field="type") if "growth" in sections else {}
        plant_years = sorted(plant_series.keys())
        if plant_years:
            all_plant_types = sorted({k for y in plant_years for k in plant_series[y].keys()})
            series_cap = {t: [plant_series[y].get(t, [0, 0])[1] for y in plant_years]
                          for t in all_plant_types}
            operating_plants = [r for r in plant_recs if r["status"] == "Operating"]
            op_series = loader.yearly_series(operating_plants, key_field="status")
            op_by_year = {y: op_series.get(y, {}).get("Operating", [0, 0.0])[1] for y in plant_years}
            cum, cum_capacity = 0.0, []
            for y in plant_years:
                cum += op_by_year.get(y, 0.0)
                cum_capacity.append(cum)
            chart_specs.append((_draw_multiline_cum, dict(
                title="Power Plants — Licensed Capacity by Year (B.S.)",
                x_labels=[str(y) for y in plant_years], series=series_cap,
                colors={t: get_type_colors().get(t, "#607d8b") for t in all_plant_types},
                y_label="Capacity Licensed This Year (MW)", cum_values=cum_capacity,
                cum_label="Cumulative Installed Capacity — Operating (MW)",
            )))
            series_count = {t: [plant_series[y].get(t, [0, 0])[0] for y in plant_years]
                            for t in all_plant_types}
            chart_specs.append((_draw_stacked_bar, dict(
                title="Power Plants — Project Count by Year",
                x_labels=[str(y) for y in plant_years], series=series_count,
                colors={t: get_type_colors().get(t, "#607d8b") for t in all_plant_types},
                y_label="Number of Projects",
            )))

        tx_series = loader.yearly_series(tx_recs, key_field="status") if "growth" in sections else {}
        tx_years = sorted(tx_series.keys())
        if tx_years:
            all_tx_statuses = sorted({k for y in tx_years for k in tx_series[y].keys()})
            series_tx_cap = {s: [tx_series[y].get(s, [0, 0])[1] for y in tx_years]
                             for s in all_tx_statuses}
            tx_totals_by_year = [sum(tx_series[y].get(s, [0, 0.0])[1] for s in all_tx_statuses)
                                for y in tx_years]
            chart_specs.append((_draw_multiline_cum, dict(
                title="Transmission Lines — Licensed Capacity by Year (B.S.)",
                x_labels=[str(y) for y in tx_years], series=series_tx_cap,
                colors={s: get_status_colors().get(s, "#90a4ae") for s in all_tx_statuses},
                y_label="Capacity (MW)", cum_values=_cumsum(tx_totals_by_year),
                cum_label="Cumulative Total Capacity (MW)",
            )))
            series_tx_count = {s: [tx_series[y].get(s, [0, 0])[0] for y in tx_years]
                               for s in all_tx_statuses}
            chart_specs.append((_draw_stacked_bar, dict(
                title="Transmission Lines — Project Count by Year",
                x_labels=[str(y) for y in tx_years], series=series_tx_count,
                colors={s: get_status_colors().get(s, "#90a4ae") for s in all_tx_statuses},
                y_label="Number of Projects",
            )))

        # ── NEA Operational Data ─────────────────────────────────────
        if "nea" in sections:
            chart_specs.extend(_nea_report_chart_specs())

        _pdf_render_batches(pdf, chart_specs, per_page=3)

    return dcc.send_file(path), dbc.Alert(
        "✅ Your report is ready — the download has started.",
        color="success", dismissable=True)


@server.route("/sitemap.xml")
def serve_sitemap():
    from flask import Response
    xml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">\n'
        '  <url><loc>https://www.neupanesandeep.com.np/</loc></url>\n'
        '</urlset>'
    )
    return Response(xml, mimetype="application/xml")


@server.route("/robots.txt")
def serve_robots():
    from flask import Response
    return Response(
        "User-agent: *\nAllow: /\nSitemap: https://www.neupanesandeep.com.np/sitemap.xml",
        mimetype="text/plain",
    )


if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=int(os.environ.get("PORT", 8050)))
